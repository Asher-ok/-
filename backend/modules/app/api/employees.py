from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Body, Query, BackgroundTasks
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.responses import FileResponse, Response
from sqlalchemy import or_
from sqlalchemy.orm import Session, joinedload
from core.database import get_db
from shared.models import Employee, TrainingRecord, Qualification, EmployeeDocument
from shared.models import BusinessUnread
from shared.models import User
from shared.models.update_notification import touch_business_unread
from ..schemas.employee import EmployeeResponse
from ..dependencies import get_current_employee
from core.auth import decode_access_token
from core.utils.file_utils import (
    get_file_path,
    save_upload_file,
    delete_file,
    build_content_disposition,
    to_ascii_filename,
)
from core.config import settings
from datetime import datetime
import pytz
from pathlib import Path
from typing import Optional
import mimetypes
import logging
import uuid
import base64
import re
import io
import os
import subprocess
import tempfile
import shutil
import time
from PyPDF2 import PdfReader, PdfWriter
from pydantic import BaseModel

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/app/employees", tags=["员工"])
optional_security = HTTPBearer(auto_error=False)


def _parse_data_url(data_url: str):
    """解析base64 data URL"""
    match = re.match(r"^data:(.+?);base64,(.+)$", data_url)
    if not match:
        return None, None
    mime_type = match.group(1)
    data = base64.b64decode(match.group(2))
    return mime_type, data


from PyPDF2.generic import NameObject, NumberObject, DictionaryObject, ArrayObject
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader

def _embed_signature_to_pdf(
    source_pdf_path: Path,
    signature_image_data: bytes,
    position_x: float,
    position_y: float,
    width: float,
    height: float,
    use_percentage: bool = False
) -> Path:
    """
    嵌入签名到PDF（增强版，避免破坏PDF结构）
    :param source_pdf_path: 源PDF路径
    :param signature_image_data: 签名图片二进制数据
    :param position_x/y: 坐标（绝对像素/百分比）
    :param width/height: 签名尺寸
    :param use_percentage: 是否使用百分比定位
    :return: 新PDF路径
    """
    # 1. 验证源PDF
    if not source_pdf_path.exists() or source_pdf_path.stat().st_size == 0:
        raise ValueError(f"源PDF文件为空或不存在：{source_pdf_path}")
    
    # 2. 读取源PDF（兼容加密/损坏的PDF）
    try:
        reader = PdfReader(source_pdf_path)
        # 检查是否加密
        if reader.is_encrypted:
            try:
                reader.decrypt("")  # 尝试空密码解密
            except Exception as e:
                raise RuntimeError(f"PDF已加密，无法嵌入签名：{e}")
    except Exception as e:
        raise RuntimeError(f"读取源PDF失败（结构损坏）：{e}")
    
    # 3. 创建签名层（使用reportlab生成透明签名）
    signature_buffer = io.BytesIO()
    c = canvas.Canvas(signature_buffer)
    
    # 获取第一页尺寸（用于百分比计算）
    first_page = reader.pages[0]
    page_width = float(first_page.mediabox.width)
    page_height = float(first_page.mediabox.height)
    
    # 计算最终坐标（处理百分比）
    if use_percentage:
        x = page_width * (position_x / 100.0)
        y = page_height * (position_y / 100.0)
        sig_width = page_width * (width / 100.0)
        sig_height = page_height * (height / 100.0)
    else:
        x = position_x
        y = position_y
        sig_width = width
        sig_height = height
    
    # 加载签名图片（保留透明通道）
    try:
        img = ImageReader(io.BytesIO(signature_image_data))
        # 绘制图片（透明背景）
        c.drawImage(img, x, y, width=sig_width, height=sig_height, mask='auto')
        c.save()
    except Exception as e:
        raise RuntimeError(f"处理签名图片失败：{e}")
    
    # 4. 合并签名层到PDF（核心：保留原始PDF结构）
    signature_buffer.seek(0)
    signature_reader = PdfReader(signature_buffer)
    signature_page = signature_reader.pages[0]
    
    writer = PdfWriter()
    # 遍历所有页面，仅在第一页嵌入签名
    for idx, page in enumerate(reader.pages):
        if idx == 0:
            # 合并签名层（避免覆盖原始内容）
            page.merge_page(signature_page)
        # 保留原始PDF的元数据/结构
        writer.add_page(page)
    
    # 5. 写入新PDF（避免覆盖原文件）
    output_pdf_path = source_pdf_path.parent / f"{source_pdf_path.stem}_signed_{uuid.uuid4().hex}.pdf"
    try:
        with open(output_pdf_path, "wb") as f:
            writer.write(f)
    except Exception as e:
        raise RuntimeError(f"写入签名PDF失败：{e}")
    
    # 验证生成的PDF
    if not output_pdf_path.exists() or output_pdf_path.stat().st_size == 0:
        raise RuntimeError(f"生成的签名PDF为空：{output_pdf_path}")
    
    return output_pdf_path


def _convert_office_to_pdf(source_path: Path) -> Path:
    """将 doc/docx 转换为 PDF（依赖 LibreOffice）。"""
    soffice = _resolve_soffice_executable()
    if not soffice:
        raise RuntimeError("未找到 LibreOffice（soffice）。请安装 LibreOffice，或将 soffice 加入 PATH（Windows 可检查 LibreOffice\\program\\soffice.com）。")

    soffice_dir = str(Path(soffice).resolve().parent)
    output_dir = Path(tempfile.gettempdir()) / "empowerhub_previews"
    output_dir.mkdir(parents=True, exist_ok=True)
    job_dir = Path(tempfile.mkdtemp(prefix="empowerhub_lo_job_"))
    tmp_dir = job_dir / "tmp"
    tmp_dir.mkdir(parents=True, exist_ok=True)

    def run_convert(isolated_profile: bool):
        env = os.environ.copy()
        env["TMP"] = str(tmp_dir)
        env["TEMP"] = str(tmp_dir)
        env["TMPDIR"] = str(tmp_dir)

        command = [
            soffice,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--norestore",
        ]
        if isolated_profile:
            profile_dir = job_dir / "profile"
            profile_dir.mkdir(parents=True, exist_ok=True)
            command.append(f"-env:UserInstallation={profile_dir.as_uri()}")

        command.extend(
            [
                "--convert-to",
                "pdf:writer_pdf_Export",
                "--outdir",
                str(output_dir),
                str(source_path),
            ]
        )
        return subprocess.run(
            command,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            env=env,
            cwd=soffice_dir,
        )

    def pick_generated_pdf(result: subprocess.CompletedProcess) -> Path | None:
        text = "\n".join([result.stdout or "", result.stderr or ""]).strip()
        match = re.search(r"->\s*(.+?\\.pdf)\\s+using filter", text, flags=re.IGNORECASE)
        if match:
            raw = match.group(1).strip().strip('"')
            p = Path(raw)
            if p.exists() and p.is_file():
                return p
            if not p.is_absolute():
                p2 = (output_dir / p).resolve()
                if p2.exists() and p2.is_file():
                    return p2
        return None
    try:
        start_ts = time.time()
        result = run_convert(isolated_profile=False)
        if result.returncode != 0:
            start_ts = time.time()
            result = run_convert(isolated_profile=True)
        if result.returncode != 0:
            stderr = (result.stderr or "").strip()
            stdout = (result.stdout or "").strip()
            message = stderr or stdout or "docx转PDF失败"
            raise RuntimeError(f"{message} (soffice={soffice})")

        generated_pdf = pick_generated_pdf(result) or (output_dir / f"{source_path.stem}.pdf")
        if not generated_pdf.exists():
            deadline = time.time() + 5
            while time.time() < deadline:
                candidates = list(output_dir.glob("*.pdf"))
                if candidates:
                    break
                time.sleep(0.1)
            candidates = [
                p
                for p in output_dir.glob("*.pdf")
                if p.is_file() and p.stat().st_mtime >= start_ts - 2
            ]
            stem_lower = source_path.stem.lower()
            stem_candidates = [p for p in candidates if p.stem.lower().startswith(stem_lower)]
            generated_pdf = max(stem_candidates or candidates, key=lambda p: p.stat().st_mtime, default=None)
            if not generated_pdf:
                raise RuntimeError("未找到转换后的PDF文件")

        final_pdf = output_dir / f"{source_path.stem}_{uuid.uuid4().hex}.pdf"
        try:
            generated_pdf.replace(final_pdf)
        except Exception:
            final_pdf = generated_pdf

        return final_pdf
    finally:
        shutil.rmtree(job_dir, ignore_errors=True)


def _convert_docx_to_pdf_in_place(source_path: Path) -> Path:
    """将doc/docx转换为PDF并放到原目录，返回新PDF路径。"""
    soffice = _resolve_soffice_executable()
    if not soffice:
        raise RuntimeError("未找到 LibreOffice（soffice）。请安装 LibreOffice，或将 soffice 加入 PATH（Windows 可检查 LibreOffice\\program\\soffice.com）。")

    soffice_dir = str(Path(soffice).resolve().parent)
    final_output_dir = source_path.parent
    job_dir = Path(tempfile.mkdtemp(prefix="empowerhub_lo_job_"))
    output_dir = job_dir / "out"
    output_dir.mkdir(parents=True, exist_ok=True)
    tmp_dir = job_dir / "tmp"
    tmp_dir.mkdir(parents=True, exist_ok=True)

    def run_convert(isolated_profile: bool):
        env = os.environ.copy()
        env["TMP"] = str(tmp_dir)
        env["TEMP"] = str(tmp_dir)
        env["TMPDIR"] = str(tmp_dir)

        command = [
            soffice,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--norestore",
        ]
        if isolated_profile:
            profile_dir = job_dir / "profile"
            profile_dir.mkdir(parents=True, exist_ok=True)
            command.append(f"-env:UserInstallation={profile_dir.as_uri()}")

        command.extend(
            [
                "--convert-to",
                "pdf:writer_pdf_Export",
                "--outdir",
                str(output_dir),
                str(source_path),
            ]
        )
        return subprocess.run(
            command,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            env=env,
            cwd=soffice_dir,
        )

    def pick_generated_pdf(result: subprocess.CompletedProcess) -> Path | None:
        text = "\n".join([result.stdout or "", result.stderr or ""]).strip()
        match = re.search(r"->\s*(.+?\\.pdf)\\s+using filter", text, flags=re.IGNORECASE)
        if match:
            raw = match.group(1).strip().strip('"')
            p = Path(raw)
            if p.exists() and p.is_file():
                return p
            if not p.is_absolute():
                p2 = (output_dir / p).resolve()
                if p2.exists() and p2.is_file():
                    return p2
        return None
    try:
        start_ts = time.time()
        result = run_convert(isolated_profile=False)
        if result.returncode != 0:
            start_ts = time.time()
            result = run_convert(isolated_profile=True)
        if result.returncode != 0:
            stderr = (result.stderr or "").strip()
            stdout = (result.stdout or "").strip()
            message = stderr or stdout or "docx转PDF失败"
            raise RuntimeError(f"{message} (soffice={soffice})")

        generated_pdf = pick_generated_pdf(result) or (output_dir / f"{source_path.stem}.pdf")
        if not generated_pdf.exists():
            deadline = time.time() + 5
            while time.time() < deadline:
                candidates = list(output_dir.glob("*.pdf"))
                if candidates:
                    break
                time.sleep(0.1)
            candidates = [
                p
                for p in output_dir.glob("*.pdf")
                if p.is_file() and p.stat().st_mtime >= start_ts - 2
            ]
            stem_lower = source_path.stem.lower()
            stem_candidates = [p for p in candidates if p.stem.lower().startswith(stem_lower)]
            generated_pdf = max(stem_candidates or candidates, key=lambda p: p.stat().st_mtime, default=None)
            if not generated_pdf:
                raise RuntimeError("未找到转换后的PDF文件")

        final_pdf = final_output_dir / f"{source_path.stem}_{uuid.uuid4().hex}.pdf"
        try:
            generated_pdf.replace(final_pdf)
        except Exception:
            try:
                shutil.copyfile(generated_pdf, final_pdf)
            except Exception:
                final_pdf = generated_pdf

        return final_pdf
    finally:
        shutil.rmtree(job_dir, ignore_errors=True)

# ========== 新增：PDF合法性校验工具函数（放在文件工具函数区，如 _convert_docx_to_pdf_in_place 下方） ==========
def _validate_pdf_file(pdf_path: Path) -> bool:
    """验证PDF文件是否合法（避免返回损坏/加密文件给前端）"""
    # 1. 基础校验：文件存在+大小合理（至少100字节，排除空文件）
    if not pdf_path.exists() or pdf_path.stat().st_size < 100:
        logger.warning(f"PDF文件为空或过小: {pdf_path}")
        return False
    
    try:
        # 2. 校验PDF魔数（文件头必须是 %PDF）
        with open(pdf_path, "rb") as f:
            header = f.read(4)
            if header != b"%PDF":
                logger.warning(f"PDF文件魔数错误: {pdf_path}，文件头={header}")
                return False
        
        # 3. 校验PDF尾部（必须包含 %%EOF）
        with open(pdf_path, "rb") as f:
            # 读取文件最后2048字节（覆盖大多数情况）
            f.seek(max(0, pdf_path.stat().st_size - 2048))
            tail = f.read(2048)
            if b"%%EOF" not in tail:
                logger.warning(f"PDF文件不完整（缺少%%EOF）: {pdf_path}")
                return False
        
        # 4. 尝试用PyPDF2解析（验证结构完整性）
        from PyPDF2 import PdfReader
        reader = PdfReader(str(pdf_path))
        # 检查是否加密（加密PDF无法预览）
        if reader.is_encrypted:
            logger.warning(f"PDF文件已加密: {pdf_path}")
            return False
        # 检查是否有至少1页内容
        if len(reader.pages) == 0:
            logger.warning(f"PDF文件无页面内容: {pdf_path}")
            return False
        
        return True
    except ImportError:
        logger.warning("PyPDF2未安装，跳过PDF结构校验")
        return True  # 无依赖时跳过校验（避免影响其他功能）
    except Exception as e:
        logger.error(f"PDF文件校验失败: {pdf_path}，错误: {str(e)}")
        return False

def _resolve_soffice_executable() -> Optional[str]:
    configured_path = (
        getattr(settings, "soffice_path", None)
        or getattr(settings, "SOFFICE_PATH", None)
        or os.environ.get("SOFFICE_PATH")
    )
    if configured_path:
        candidate = Path(configured_path)
        # 在 Windows 下优先使用 soffice.com
        if candidate.name.lower() == "soffice.exe":
            sibling = candidate.with_name("soffice.com")
            if sibling.exists():
                candidate = sibling
        if candidate.exists() and candidate.is_file():
            return str(candidate)

    configured_dir = (
        getattr(settings, "libreoffice_program_dir", None)
        or getattr(settings, "LIBREOFFICE_PROGRAM_DIR", None)
        or os.environ.get("LIBREOFFICE_PROGRAM_DIR")
    )
    if configured_dir:
        base = Path(configured_dir)
        for exe in ("soffice.com", "soffice.exe"):
            candidate = base / exe
            if candidate.exists() and candidate.is_file():
                return str(candidate)

    for name in ("soffice", "soffice.com", "soffice.exe"):
        resolved = shutil.which(name)
        if resolved:
            return resolved

    if os.name == "nt":
        program_files = [
            os.environ.get("PROGRAMFILES"),
            os.environ.get("PROGRAMFILES(X86)"),
            "C:\\Program Files",
            "C:\\Program Files (x86)",
        ]
        for base in [p for p in program_files if p]:
            for exe in ("soffice.com", "soffice.exe"):
                candidate = Path(base) / "LibreOffice" / "program" / exe
                if candidate.exists():
                    return str(candidate)

    return None


def _embed_signature_to_word(contract_path: Path, signature_blob: bytes, signature_type: str = "employee", x: float = None, y: float = None, width: float = None, height: float = None):
    """将签名嵌入到Word文档的指定坐标位置"""
    try:
        from docx import Document
        from docx.shared import Inches, Mm, Pt, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls, qn
        from docx.oxml import OxmlElement
        from PIL import Image
        import io
        
        doc = Document(str(contract_path))
        
        # 创建签名图片的临时文件
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
            tmp_file.write(signature_blob)
            tmp_signature_path = tmp_file.name
        
        try:
            if x is not None and y is not None and width is not None and height is not None:
                section = doc.sections[0]
                page_width = section.page_width
                page_height = section.page_height

                is_ratio = 0 <= x <= 1 and 0 <= y <= 1 and 0 <= width <= 1 and 0 <= height <= 1
                if is_ratio:
                    box_w_emu = float(page_width) * float(width)
                    box_h_emu = float(page_height) * float(height)
                    box_x_emu = float(page_width) * float(x)
                    box_y_emu = float(page_height) * float(y)
                else:
                    # assume 96 DPI pixel values
                    px_to_emu = 914400.0 / 96.0
                    box_w_emu = float(width) * px_to_emu
                    box_h_emu = float(height) * px_to_emu
                    box_x_emu = float(x) * px_to_emu
                    box_y_emu = float(y) * px_to_emu

                img = Image.open(io.BytesIO(signature_blob)).convert("RGBA")
                w, h = img.size
                px = img.load()
                min_x, min_y = w, h
                max_x, max_y = -1, -1
                for yy in range(h):
                    for xx in range(w):
                        r, g, b, a = px[xx, yy]
                        if a <= 10:
                            continue
                        if r >= 245 and g >= 245 and b >= 245:
                            continue
                        if xx < min_x:
                            min_x = xx
                        if yy < min_y:
                            min_y = yy
                        if xx > max_x:
                            max_x = xx
                        if yy > max_y:
                            max_y = yy
                if max_x >= 0 and max_y >= 0:
                    left = max(min_x, 0)
                    top = max(min_y, 0)
                    right = min(max_x + 1, w)
                    bottom = min(max_y + 1, h)
                    img = img.crop((left, top, right, bottom))
                img.save(tmp_signature_path, format="PNG")

                px_to_emu = 914400.0 / 96.0
                img_w_emu = float(img.size[0]) * px_to_emu
                img_h_emu = float(img.size[1]) * px_to_emu

                pad_emu = 2.0 * px_to_emu
                inner_w_emu = max(box_w_emu - pad_emu * 2, 1.0)
                inner_h_emu = max(box_h_emu - pad_emu * 2, 1.0)
                scale = min(inner_w_emu / max(img_w_emu, 1.0), inner_h_emu / max(img_h_emu, 1.0))
                draw_w_emu = max(img_w_emu * scale, 1.0)
                draw_h_emu = max(img_h_emu * scale, 1.0)

                pos_x_emu = box_x_emu + pad_emu
                pos_y_emu = box_y_emu + box_h_emu - pad_emu - draw_h_emu
                max_x_emu = box_x_emu + box_w_emu - pad_emu - draw_w_emu
                min_y_emu = box_y_emu + pad_emu
                if pos_x_emu > max_x_emu:
                    pos_x_emu = max_x_emu
                if pos_y_emu < min_y_emu:
                    pos_y_emu = min_y_emu

                width_emu = int(draw_w_emu)
                height_emu = int(draw_h_emu)
                pos_x_emu = int(pos_x_emu)
                pos_y_emu = int(pos_y_emu)

                paragraph = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(tmp_signature_path, width=Emu(width_emu), height=Emu(height_emu))

                drawing = run._r.xpath('w:drawing')[0]
                inline = drawing.xpath('wp:inline')[0]

                anchor = OxmlElement('wp:anchor')
                anchor.set(qn('wp:simplePos'), '0')
                anchor.set(qn('wp:relativeHeight'), '251659264')
                anchor.set(qn('wp:behindDoc'), '0')
                anchor.set(qn('wp:locked'), '0')
                anchor.set(qn('wp:layoutInCell'), '1')
                anchor.set(qn('wp:allowOverlap'), '1')

                simple_pos = OxmlElement('wp:simplePos')
                simple_pos.set('x', '0')
                simple_pos.set('y', '0')
                anchor.append(simple_pos)

                position_h = OxmlElement('wp:positionH')
                position_h.set(qn('wp:relativeFrom'), 'page')
                pos_offset_h = OxmlElement('wp:posOffset')
                pos_offset_h.text = str(pos_x_emu)
                position_h.append(pos_offset_h)
                anchor.append(position_h)

                position_v = OxmlElement('wp:positionV')
                position_v.set(qn('wp:relativeFrom'), 'page')
                pos_offset_v = OxmlElement('wp:posOffset')
                pos_offset_v.text = str(pos_y_emu)
                position_v.append(pos_offset_v)
                anchor.append(position_v)

                wrap_none = OxmlElement('wp:wrapNone')
                anchor.append(wrap_none)

                for child in list(inline):
                    anchor.append(child)

                drawing.remove(inline)
                drawing.append(anchor)

                doc.save(str(contract_path))
                return True
            else:
                return _embed_signature_to_word_legacy(contract_path, signature_blob, signature_type)
        finally:
            if os.path.exists(tmp_signature_path):
                os.unlink(tmp_signature_path)
    except Exception as e:
        logger.error(f"嵌入签名到Word文档失败: {e}")
        raise


def _embed_signature_to_word_legacy(contract_path: Path, signature_blob: bytes, signature_type: str = "employee"):
    """将签名嵌入到合同Word文档中（使用占位符方式，向后兼容）"""
    try:
        from docx import Document
        from docx.shared import Inches, Mm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document(str(contract_path))
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
            tmp_file.write(signature_blob)
            tmp_signature_path = tmp_file.name
        
        try:
            placeholder_text = f"[{signature_type.capitalize()} Signature]"
            found_placeholder = False
            
            for paragraph in doc.paragraphs:
                if placeholder_text in paragraph.text:
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.add_picture(tmp_signature_path, width=Mm(100))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    found_placeholder = True
                    break
            
            if not found_placeholder:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if placeholder_text in paragraph.text:
                                    paragraph.clear()
                                    run = paragraph.add_run()
                                    run.add_picture(tmp_signature_path, width=Mm(100))
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                    found_placeholder = True
                                    break
                            if found_placeholder:
                                break
                        if found_placeholder:
                            break
                    if found_placeholder:
                        break
            
            doc.save(str(contract_path))
            return found_placeholder
        finally:
            if os.path.exists(tmp_signature_path):
                os.unlink(tmp_signature_path)
    except Exception as e:
        logger.error(f"嵌入签名到Word文档失败: {e}")
        raise


def _embed_signature_to_pdf(
    contract_path: Path,
    signature_blob: bytes | None,
    signature_type: str = "employee",
    x: float = None,
    y: float = None,
    width: float = None,
    height: float = None,
    page_index: int = 0,
    draw_signature: bool = True,
    date_text: str | None = None,
):
    """将签名嵌入到PDF文档的指定坐标位置"""
    try:
        from PyPDF2 import PdfReader, PdfWriter  # type: ignore
        from reportlab.pdfgen import canvas
        from reportlab.lib.utils import ImageReader
        import io
        
        contract_path = Path(contract_path)
        with open(contract_path, "rb") as f:
            original_pdf_bytes = f.read()
        reader = PdfReader(io.BytesIO(original_pdf_bytes))
        writer = PdfWriter()
        
        # 将签名图片转换为PIL Image
        from PIL import Image

        def _trim_signature_image(img: Image.Image, padding: int = 5) -> Image.Image:
            img = img.convert("RGBA")
            w, h = img.size
            px = img.load()
            min_x, min_y = w, h
            max_x, max_y = -1, -1
            for yy in range(h):
                for xx in range(w):
                    r, g, b, a = px[xx, yy]
                    if a <= 10:
                        continue
                    if r >= 245 and g >= 245 and b >= 245:
                        continue
                    if xx < min_x:
                        min_x = xx
                    if yy < min_y:
                        min_y = yy
                    if xx > max_x:
                        max_x = xx
                    if yy > max_y:
                        max_y = yy
            if max_x < 0 or max_y < 0:
                return img
            left = max(min_x - padding, 0)
            top = max(min_y - padding, 0)
            right = min(max_x + padding + 1, w)
            bottom = min(max_y + padding + 1, h)
            return img.crop((left, top, right, bottom))

        if len(reader.pages) > 0:
            page_index = max(0, min(page_index, len(reader.pages) - 1))
            target_page = reader.pages[page_index]
            page_width = float(target_page.mediabox.width)
            page_height = float(target_page.mediabox.height)
        else:
            page_index = 0
            page_width = 595.0
            page_height = 842.0

        packet = io.BytesIO()
        can = canvas.Canvas(packet, pagesize=(page_width, page_height))
        
        # 计算坐标（PDF坐标系统：左下角为原点，单位是点）
        # 假设输入的坐标是像素，需要转换为点
        signature_bottom_y: float | None = None
        if draw_signature:
            if not signature_blob:
                raise RuntimeError("签名数据为空，无法嵌入签名")
            signature_image = _trim_signature_image(Image.open(io.BytesIO(signature_blob)), padding=0)
            img_w, img_h = signature_image.size

            png_bytes = io.BytesIO()
            signature_image.save(png_bytes, format="PNG")
            png_bytes.seek(0)

            if x is not None and y is not None and width is not None and height is not None:
                is_ratio = 0 <= x <= 1 and 0 <= y <= 1 and 0 <= width <= 1 and 0 <= height <= 1
                if is_ratio:
                    pdf_x = x * page_width
                    pdf_y = (1 - y - height) * page_height
                    pdf_width = width * page_width
                    pdf_height = height * page_height
                else:
                    pdf_x = x
                    pdf_y = page_height - y - height
                    pdf_width = width
                    pdf_height = height

                pdf_width = max(1.0, min(float(pdf_width), page_width))
                pdf_height = max(1.0, min(float(pdf_height), page_height))
                pdf_x = max(0.0, min(float(pdf_x), page_width - pdf_width))
                pdf_y = max(0.0, min(float(pdf_y), page_height - pdf_height))

                pad = 2.0
                inner_w = max(float(pdf_width) - pad * 2, 1.0)
                inner_h = max(float(pdf_height) - pad * 2, 1.0)
                scale = min(inner_w / max(float(img_w), 1.0), inner_h / max(float(img_h), 1.0))
                draw_w = float(img_w) * scale
                draw_h = float(img_h) * scale

                draw_x = float(pdf_x) + pad
                draw_y = float(pdf_y) + pad
                max_x = float(pdf_x) + float(pdf_width) - pad - draw_w
                max_y = float(pdf_y) + float(pdf_height) - pad - draw_h
                if draw_x > max_x:
                    draw_x = max_x
                if draw_y > max_y:
                    draw_y = max_y

                img_reader = ImageReader(png_bytes)
                can.drawImage(img_reader, draw_x, draw_y, width=draw_w, height=draw_h, mask="auto")
                signature_bottom_y = float(draw_y)
            else:
                img_reader = ImageReader(png_bytes)
                box_w, box_h = 200, 80
                scale = min(box_w / max(img_w, 1), box_h / max(img_h, 1))
                draw_w = img_w * scale
                draw_h = img_h * scale
                can.drawImage(img_reader, 50, 50, width=draw_w, height=draw_h, mask="auto")
                signature_bottom_y = 50.0

        if date_text and x is not None and y is not None and width is not None and height is not None:
            is_ratio = 0 <= x <= 1 and 0 <= y <= 1 and 0 <= width <= 1 and 0 <= height <= 1

            # --- 自适应间距算法 ---
            # 计算签名框的基础高度和宽度
            if is_ratio:
                pdf_box_w = float(width) * page_width
                pdf_box_h = float(height) * page_height
                pdf_box_x = float(x) * page_width
                pdf_box_y = (1 - float(y) - float(height)) * page_height
            else:
                pdf_box_w = float(width)
                pdf_box_h = float(height)
                pdf_box_x = float(x)
                pdf_box_y = page_height - float(y) - float(height)

            pdf_box_w = max(1.0, min(float(pdf_box_w), page_width))
            pdf_box_h = max(1.0, min(float(pdf_box_h), page_height))
            pdf_box_x = max(0.0, min(float(pdf_box_x), page_width - pdf_box_w))
            pdf_box_y = max(0.0, min(float(pdf_box_y), page_height - pdf_box_h))

            gap = max(8.0, min(pdf_box_h * 0.2, 20.0))

            font_name = "Helvetica"
            font_size = max(12.0, min(pdf_box_h * 0.35, 16.0))
            can.setFont(font_name, font_size)

            try:
                text_width = can.stringWidth(date_text, font_name, font_size)
            except Exception:
                text_width = len(date_text) * font_size * 0.6

            pdf_date_x = pdf_box_x
            pdf_date_x = max(1.0, min(pdf_date_x, page_width - text_width - 1.0))

            anchor_y = signature_bottom_y if signature_bottom_y is not None else pdf_box_y
            baseline_y = anchor_y - gap - font_size
            baseline_y = max(1.0, min(baseline_y, anchor_y - 1.0))
            baseline_y = min(baseline_y, page_height - font_size - 1.0)

            can.setFillColorRGB(0, 0, 0)
            can.drawString(pdf_date_x, baseline_y, date_text)
        
        can.save()
        packet.seek(0)
        signature_pdf = PdfReader(packet)
        
        # 将签名叠加到指定页
        if len(reader.pages) > 0:
            page_index = max(0, min(page_index, len(reader.pages) - 1))
            target_page = reader.pages[page_index]
            signature_page = signature_pdf.pages[0]
            target_page.merge_page(signature_page)
        
        # 添加所有页面
        for page in reader.pages:
            writer.add_page(page)
        
        tmp_path = contract_path.parent / f"{contract_path.stem}.signed_tmp_{uuid.uuid4().hex}{contract_path.suffix}"
        with open(tmp_path, "wb") as output_file:
            writer.write(output_file)
        os.replace(tmp_path, contract_path)
        
        return True
    except ImportError as e:
        logger.warning("PyPDF2或reportlab未安装，无法嵌入PDF签名")
        raise RuntimeError("PDF签名依赖缺失") from e
    except Exception as e:
        logger.error(f"嵌入签名到PDF失败: {e}")
        return False


def _embed_signature_to_image(contract_path: Path, signature_blob: bytes, signature_type: str = "employee", x: float = None, y: float = None, width: float = None, height: float = None):
    """将签名嵌入到图片的指定坐标位置"""
    try:
        from PIL import Image
        import io
        
        # 打开原图片
        original_image = Image.open(contract_path)
        
        # 打开签名图片
        signature_image = Image.open(io.BytesIO(signature_blob)).convert("RGBA")

        def _trim(img: Image.Image, padding: int = 5) -> Image.Image:
            w, h = img.size
            px = img.load()
            min_x, min_y = w, h
            max_x, max_y = -1, -1
            for yy in range(h):
                for xx in range(w):
                    r, g, b, a = px[xx, yy]
                    if a <= 10:
                        continue
                    if r >= 245 and g >= 245 and b >= 245:
                        continue
                    if xx < min_x:
                        min_x = xx
                    if yy < min_y:
                        min_y = yy
                    if xx > max_x:
                        max_x = xx
                    if yy > max_y:
                        max_y = yy
            if max_x < 0 or max_y < 0:
                return img
            left = max(min_x - padding, 0)
            top = max(min_y - padding, 0)
            right = min(max_x + padding + 1, w)
            bottom = min(max_y + padding + 1, h)
            return img.crop((left, top, right, bottom))

        signature_image = _trim(signature_image, padding=0)

        if width is not None and height is not None and x is not None and y is not None:
            is_ratio = 0 <= x <= 1 and 0 <= y <= 1 and 0 <= width <= 1 and 0 <= height <= 1
            if is_ratio:
                box_x = float(original_image.width) * float(x)
                box_y = float(original_image.height) * float(y)
                box_w = float(original_image.width) * float(width)
                box_h = float(original_image.height) * float(height)
            else:
                box_x = float(x)
                box_y = float(y)
                box_w = float(width)
                box_h = float(height)

            img_w, img_h = signature_image.size

            pad = 2.0
            inner_w = max(box_w - pad * 2, 1.0)
            inner_h = max(box_h - pad * 2, 1.0)
            scale = min(inner_w / max(float(img_w), 1.0), inner_h / max(float(img_h), 1.0))
            draw_w = max(int(round(img_w * scale)), 1)
            draw_h = max(int(round(img_h * scale)), 1)
            signature_image = signature_image.resize((draw_w, draw_h), Image.Resampling.LANCZOS)

            paste_x = int(round(box_x + pad))
            paste_y = int(round(box_y + box_h - pad - draw_h))
            max_x = box_x + box_w - pad - draw_w
            min_y = box_y + pad
            if paste_x > max_x:
                paste_x = int(round(max_x))
            if paste_y < min_y:
                paste_y = int(round(min_y))
        else:
            default_width = original_image.width // 4
            default_height = original_image.height // 8
            signature_image = signature_image.resize((max(default_width, 1), max(default_height, 1)), Image.Resampling.LANCZOS)
            paste_x = original_image.width - signature_image.width - 20
            paste_y = original_image.height - signature_image.height - 20
        
        # 确保坐标在图片范围内
        paste_x = max(0, min(paste_x, original_image.width - signature_image.width))
        paste_y = max(0, min(paste_y, original_image.height - signature_image.height))
        
        if signature_image.mode == 'RGBA':
            original_image.paste(signature_image, (paste_x, paste_y), signature_image)
        else:
            original_image.paste(signature_image, (paste_x, paste_y))
        
        # 保存
        original_image.save(contract_path)
        return True
    except ImportError:
        logger.warning("PIL未安装，无法嵌入图片签名")
        return False
    except Exception as e:
        logger.error(f"嵌入签名到图片失败: {e}")
        return False


def _training_record_to_dict(record):
    certificate_url = None
    if record.has_certificate and (record.certificate_url or record.certificate_number):
        certificate_url = f"/api/app/employees/me/training-records/{record.id}/certificate"
    return {
        "id": record.id,
        "name": record.name,
        "completed_date": record.completed_date or datetime.utcnow(),
        "status": record.status,
        "score": record.score,
        "has_certificate": bool(record.has_certificate),
        "certificate_number": record.certificate_number,
        "certificate_url": certificate_url,
        "training_institution": record.training_institution,
        "notes": record.notes,
        "created_by": getattr(record, 'created_by', None),
    }


@router.get("/me", response_model=EmployeeResponse)
async def get_current_employee_info(
    current_employee: Employee = Depends(get_current_employee),
    db: Session = Depends(get_db)
):
    """获取当前员工信息"""
    employee = db.query(Employee).options(
        joinedload(Employee.qualifications),
        joinedload(Employee.training_records)
    ).filter(Employee.id == current_employee.id).first()
    target = employee or current_employee
    training_records = [
        _training_record_to_dict(record)
        for record in (target.training_records or [])
    ]
    qualifications = []
    logger.info(f"员工资质总数: {len(target.qualifications or [])}")
    for qual in (target.qualifications or []):
        # 如果有 certificate_blob 但没有 certificate_url，使用 API 端点作为 URL
        certificate_url = qual.certificate_url
        if not certificate_url and qual.certificate_blob:
            certificate_url = f"/api/app/employees/me/qualifications/{qual.id}/certificate"
        
        logger.info(f"资质: id={qual.id}, name={qual.name}, certificate_number={qual.certificate_number or '(空)'}, certificate_url={certificate_url or '(空)'}, has_blob={qual.certificate_blob is not None}")
        
        qualifications.append({
            "id": qual.id,
            "name": qual.name,
            "certificate_number": qual.certificate_number,
            "certificate_url": certificate_url,
            "obtained_date": qual.obtained_date,
            "expiry_date": qual.expiry_date,
            "issuing_authority": qual.issuing_authority,
        })
    
    # 从培训记录中提取有证书的数据作为资质证书
    logger.info(f"检查培训记录中的证书数据，培训记录总数: {len(target.training_records or [])}")
    for record in (target.training_records or []):
        if record.has_certificate and (record.certificate_number or record.certificate_url):
            # 构建证书URL - 如果有证书，总是生成API端点
            certificate_url = record.certificate_url
            if not certificate_url or certificate_url.strip() == "":
                # 如果certificate_url为空，生成API端点
                certificate_url = f"/api/app/employees/me/training-records/{record.id}/certificate"
            elif not certificate_url.startswith('http'):
                # 如果是相对路径，转换为API端点
                certificate_url = f"/api/app/employees/me/training-records/{record.id}/certificate"
            
            # 使用培训记录的证书相关字段创建资质证书数据
            qualification_from_training = {
                "id": f"training_{record.id}",  # 使用特殊前缀区分来源
                "name": record.name,
                "certificate_number": record.certificate_number or "",
                "certificate_url": certificate_url,
                "obtained_date": record.certificate_obtained_date or record.completed_date,
                "expiry_date": record.certificate_expiry_date,
                "issuing_authority": record.training_institution,
            }
            qualifications.append(qualification_from_training)
            logger.info(f"从培训记录添加资质: name={record.name}, certificate_number={record.certificate_number or '(空)'}, certificate_url={certificate_url}")
    
    logger.info(f"返回资质数量（包括培训记录）: {len(qualifications)}")
    return {
        "id": target.id,
        "name": target.name,
        "employee_number": target.employee_number,
        "department": target.department,
        "phone": target.phone,
        "email": target.email,
        "avatar_url": target.avatar_url,
        "account_status": getattr(target, "account_status", None),
        "qualifications": qualifications,
        "training_records": training_records,
    }


@router.get("/me/qualifications")
async def get_my_qualifications(
    current_employee: Employee = Depends(get_current_employee),
    db: Session = Depends(get_db)
):
    """获取当前员工的资质列表"""
    employee = db.query(Employee).options(joinedload(Employee.qualifications)).filter(
        Employee.id == current_employee.id
    ).first()
    if not employee:
        return []
    result = []
    for qual in employee.qualifications:
        # 如果有 certificate_blob 但没有 certificate_url，使用 API 端点作为 URL
        certificate_url = qual.certificate_url
        if not certificate_url and qual.certificate_blob:
            certificate_url = f"/api/app/employees/me/qualifications/{qual.id}/certificate"
        
        logger.info(f"资质: name={qual.name}, certificate_number={qual.certificate_number}, certificate_url={certificate_url}, has_blob={qual.certificate_blob is not None}")
        
        result.append({
            "id": qual.id,
            "name": qual.name,
            "certificate_number": qual.certificate_number,
            "certificate_url": certificate_url,
            "obtained_date": qual.obtained_date,
            "expiry_date": qual.expiry_date,
            "issuing_authority": qual.issuing_authority,
        })
    
    # 从培训记录中提取有证书的数据作为资质证书
    training_records = db.query(TrainingRecord).filter(
        TrainingRecord.employee_id == employee.id
    ).all()
    
    for record in training_records:
        if record.has_certificate and (record.certificate_number or record.certificate_url):
            # 构建证书URL - 如果有证书，总是生成API端点
            certificate_url = record.certificate_url
            if not certificate_url or certificate_url.strip() == "":
                # 如果certificate_url为空，生成API端点
                certificate_url = f"/api/app/employees/me/training-records/{record.id}/certificate"
            elif not certificate_url.startswith('http'):
                # 如果是相对路径，转换为API端点
                certificate_url = f"/api/app/employees/me/training-records/{record.id}/certificate"
            
            result.append({
                "id": f"training_{record.id}",
                "name": record.name,
                "certificate_number": record.certificate_number or "",
                "certificate_url": certificate_url,
                "obtained_date": record.certificate_obtained_date or record.completed_date,
                "expiry_date": record.certificate_expiry_date,
                "issuing_authority": record.training_institution,
            })
    
    return result


@router.get("/me/qualifications/{qualification_id}/certificate")
async def get_my_qualification_certificate(
    qualification_id: str,
    current_employee: Employee = Depends(get_current_employee),
    db: Session = Depends(get_db)
):
    """获取当前员工的资质证书文件"""
    logger.info(f"请求资质证书: qualification_id={qualification_id}, employee_id={current_employee.id}")
    
    # 检查是否是来自培训记录的证书（ID以training_开头）
    if qualification_id.startswith("training_"):
        record_id = qualification_id.replace("training_", "")
        # 重定向到培训记录证书端点
        from fastapi.responses import RedirectResponse
        return RedirectResponse(
            url=f"/api/app/employees/me/training-records/{record_id}/certificate",
            status_code=307
        )
    
    qualification = db.query(Qualification).filter(
        Qualification.id == qualification_id,
        Qualification.employee_id == current_employee.id
    ).first()
    
    if not qualification:
        logger.warning(f"资质不存在: qualification_id={qualification_id}, employee_id={current_employee.id}")
        raise HTTPException(status_code=404, detail="资质不存在")
    
    logger.info(f"找到资质: name={qualification.name}, certificate_url={qualification.certificate_url}, has_blob={qualification.certificate_blob is not None}")
    
    # 优先使用 certificate_blob（存储在数据库中的二进制数据）
    if qualification.certificate_blob:
        logger.info(f"从数据库返回证书文件: qualification_id={qualification_id}")
        blob_mime = (qualification.certificate_mime or "application/octet-stream").split(";")[0].strip()
        blob_suffix = mimetypes.guess_extension(blob_mime) or ""
        blob_filename = f"certificate_{qualification_id}{blob_suffix}"
        return Response(
            content=qualification.certificate_blob,
            media_type=blob_mime or "application/octet-stream",
            headers={
                "Content-Disposition": f'inline; filename="{blob_filename}"'
            }
        )
    
    # 如果没有 certificate_blob，尝试从 certificate_url 获取文件
    if not qualification.certificate_url:
        logger.warning(f"证书URL为空: qualification_id={qualification_id}")
        raise HTTPException(status_code=404, detail="证书不存在")
    
    # 获取文件路径 - 尝试多种路径格式
    file_path = None
    certificate_url = qualification.certificate_url.strip()
    logger.info(f"尝试解析证书路径: {certificate_url}")
    
    # 方法1: 如果certificate_url是绝对路径，直接使用
    file_path = Path(certificate_url)
    if file_path.exists() and file_path.is_file():
        logger.info(f"找到证书文件（绝对路径）: {file_path}")
    else:
        # 方法2: 尝试使用get_file_path函数
        file_path = get_file_path(certificate_url)
        if file_path and file_path.exists():
            logger.info(f"找到证书文件（get_file_path）: {file_path}")
        else:
            # 方法3: 尝试基于upload_dir的相对路径
            upload_path = Path(settings.upload_dir).resolve()
            logger.info(f"上传目录: {upload_path}")
            
            # 移除可能的"uploads/"前缀
            relative_path = certificate_url.replace("uploads/", "").replace("uploads\\", "")
            # 移除可能的绝对路径前缀（如果certificate_url包含upload_dir）
            if str(upload_path) in certificate_url:
                relative_path = certificate_url.replace(str(upload_path), "").lstrip("/\\")
            
            # 尝试多种组合
            possible_paths = [
                upload_path / relative_path,
                Path(relative_path),
            ]
            
            logger.info(f"尝试的路径列表: {[str(p) for p in possible_paths]}")
            
            for possible_path in possible_paths:
                logger.info(f"检查路径: {possible_path}, 存在: {possible_path.exists()}, 是文件: {possible_path.is_file() if possible_path.exists() else False}")
                if possible_path.exists() and possible_path.is_file():
                    file_path = possible_path
                    logger.info(f"找到证书文件: {file_path}")
                    break
    
    if not file_path or not file_path.exists() or not file_path.is_file():
        error_msg = f"证书文件不存在。原始URL: {qualification.certificate_url}, 上传目录: {settings.upload_dir}"
        logger.error(error_msg)
        raise HTTPException(status_code=404, detail=error_msg)
    
    # 获取MIME类型
    mime_type, _ = mimetypes.guess_type(str(file_path))
    if not mime_type:
        mime_type = "application/octet-stream"
    
    logger.info(f"返回证书文件: {file_path}, MIME类型: {mime_type}")
    
    return FileResponse(
        path=str(file_path),
        media_type=mime_type,
        filename=f"certificate_{qualification_id}{file_path.suffix}",
        headers={
            "Content-Disposition": f"inline; filename=certificate_{qualification_id}{file_path.suffix}"
        }
    )


@router.get("/me/training-records/{record_id}/certificate")
async def get_my_training_record_certificate(
    record_id: str,
    current_employee: Employee = Depends(get_current_employee),
    db: Session = Depends(get_db)
):
    """获取当前员工的培训记录证书文件"""
    logger.info(f"请求证书: record_id={record_id}, employee_id={current_employee.id}")
    
    record = db.query(TrainingRecord).filter(
        TrainingRecord.id == record_id,
        TrainingRecord.employee_id == current_employee.id
    ).first()
    
    if not record:
        logger.warning(f"培训记录不存在: record_id={record_id}, employee_id={current_employee.id}")
        raise HTTPException(status_code=404, detail="培训记录不存在")
    
    logger.info(f"找到培训记录: name={record.name}, certificate_url={record.certificate_url}, has_certificate={record.has_certificate}")
    
    if not record.certificate_url:
        logger.warning(f"证书URL为空: record_id={record_id}，尝试从资质表回退获取")
        # 回退1：尝试在资质表中查找同员工的对应证书（优先按证书编号匹配，其次按名称匹配）
        from shared.models import Qualification
        fallback_qual = None
        if record.certificate_number:
            fallback_qual = db.query(Qualification).filter(
                Qualification.employee_id == current_employee.id,
                Qualification.certificate_number == record.certificate_number
            ).first()
        if not fallback_qual:
            fallback_qual = db.query(Qualification).filter(
                Qualification.employee_id == current_employee.id,
                Qualification.name == record.name
            ).first()
        if fallback_qual:
            if getattr(fallback_qual, "certificate_blob", None):
                blob_mime = (fallback_qual.certificate_mime or "application/octet-stream").split(";")[0].strip()
                blob_suffix = mimetypes.guess_extension(blob_mime) or ""
                blob_filename = f"certificate_{record_id}{blob_suffix}"
                return Response(
                    content=fallback_qual.certificate_blob,
                    media_type=blob_mime or "application/octet-stream",
                    headers={
                        "Content-Disposition": f'inline; filename="{blob_filename}"'
                    }
                )
            if getattr(fallback_qual, "certificate_url", None):
                # 复用与资质端点一致的解析策略
                certificate_url = fallback_qual.certificate_url.strip()
                upload_path = Path(settings.upload_dir).resolve()
                candidates = [
                    Path(certificate_url),
                    get_file_path(certificate_url) or Path(""),
                    upload_path / certificate_url.replace("uploads/", "").replace("uploads\\", ""),
                    Path(certificate_url.replace(str(upload_path), "").lstrip("/\\")),
                ]
                for p in candidates:
                    try:
                        if p and p.exists() and p.is_file():
                            mime_type, _ = mimetypes.guess_type(str(p))
                            if not mime_type:
                                mime_type = "application/octet-stream"
                            return FileResponse(
                                path=str(p),
                                media_type=mime_type,
                                filename=f"certificate_{record_id}{p.suffix}",
                                headers={
                                    "Content-Disposition": f'inline; filename="certificate_{record_id}{p.suffix}"'
                                }
                            )
                    except Exception:
                        continue
        # 若仍未找到
        raise HTTPException(status_code=404, detail="证书不存在")
    
    # 获取文件路径 - 尝试多种路径格式
    file_path = None
    certificate_url = record.certificate_url.strip()
    logger.info(f"尝试解析证书路径: {certificate_url}")
    
    # 方法1: 如果certificate_url是绝对路径，直接使用
    file_path = Path(certificate_url)
    if file_path.exists() and file_path.is_file():
        logger.info(f"找到证书文件（绝对路径）: {file_path}")
    else:
        # 方法2: 尝试使用get_file_path函数
        file_path = get_file_path(certificate_url)
        if file_path and file_path.exists():
            logger.info(f"找到证书文件（get_file_path）: {file_path}")
        else:
            # 方法3: 尝试基于upload_dir的相对路径
            upload_path = Path(settings.upload_dir).resolve()
            logger.info(f"上传目录: {upload_path}")
            
            # 移除可能的"uploads/"前缀
            relative_path = certificate_url.replace("uploads/", "").replace("uploads\\", "")
            # 移除可能的绝对路径前缀（如果certificate_url包含upload_dir）
            if str(upload_path) in certificate_url:
                relative_path = certificate_url.replace(str(upload_path), "").lstrip("/\\")
            
            # 尝试多种组合
            possible_paths = [
                upload_path / relative_path,
                upload_path / "training_certificates" / relative_path,
                Path(relative_path),
            ]
            
            logger.info(f"尝试的路径列表: {[str(p) for p in possible_paths]}")
            
            for possible_path in possible_paths:
                logger.info(f"检查路径: {possible_path}, 存在: {possible_path.exists()}, 是文件: {possible_path.is_file() if possible_path.exists() else False}")
                if possible_path.exists() and possible_path.is_file():
                    file_path = possible_path
                    logger.info(f"找到证书文件: {file_path}")
                    break
    
    if not file_path or not file_path.exists() or not file_path.is_file():
        error_msg = f"证书文件不存在。原始URL: {record.certificate_url}, 上传目录: {settings.upload_dir}"
        logger.error(error_msg)
        raise HTTPException(status_code=404, detail=error_msg)
    
    # 获取MIME类型
    mime_type, _ = mimetypes.guess_type(str(file_path))
    if not mime_type:
        mime_type = "application/octet-stream"
    
    logger.info(f"返回证书文件: {file_path}, MIME类型: {mime_type}")
    
    return FileResponse(
        path=str(file_path),
        media_type=mime_type,
        filename=f"certificate_{record_id}{file_path.suffix}",
        headers={
            "Content-Disposition": f"inline; filename=certificate_{record_id}{file_path.suffix}"
        }
    )


# ==================== 员工文档管理 API ====================

DOCUMENT_TYPE_FOLDERS = {
    "contract": "employee_contracts",
    "checklist": "employee_checklists",
    "code": "employee_codes",
    "tracker": "employee_trackers",
    "handbook": "employee_handbooks",
    "onboarding": "employee_onboarding",
}


@router.get("/me/documents")
async def get_my_documents(
    document_type: str = None,
    current_employee: Employee = Depends(get_current_employee),
    db: Session = Depends(get_db)
):
    """获取当前员工的文档列表"""
    unread_doc_ids = {
        (r.data_id or "")
        for r in db.query(BusinessUnread)
        .filter(
            BusinessUnread.receiver_user_id == str(current_employee.id),
            BusinessUnread.business_code == "employee_document",
            BusinessUnread.is_unread == 1,
        )
        .all()
    }
    query = db.query(EmployeeDocument).filter(
        EmployeeDocument.employee_id == current_employee.id
    )
    
    if document_type:
        query = query.filter(EmployeeDocument.document_type == document_type)
    
    documents = query.order_by(EmployeeDocument.uploaded_at.desc()).all()
    
    return [
        {
            "id": doc.id,
            "name": doc.name,
            "file_type": doc.file_type,
            "file_url": doc.file_url,
            "document_type": doc.document_type,
            "uploaded_at": doc.uploaded_at.isoformat() if doc.uploaded_at else None,
            "employee_signed_at": doc.employee_signed_at.isoformat() if doc.employee_signed_at else None,
            "employee_signed": bool(doc.employee_signed_at),
            "has_update": str(doc.id) in unread_doc_ids,
        }
        for doc in documents
    ]


@router.post("/me/documents/upload")
async def upload_my_document(
    document_type: str = Form(..., description="文档类型: contract, checklist, code, tracker, handbook, onboarding"),
    name: str = Form(..., description="文档名称"),
    file: UploadFile = File(..., description="上传的文件"),
    current_employee: Employee = Depends(get_current_employee),
    db: Session = Depends(get_db)
):
    """上传员工文档"""
    # 验证文档类型
    if document_type not in DOCUMENT_TYPE_FOLDERS:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid document type. Allowed types: {', '.join(DOCUMENT_TYPE_FOLDERS.keys())}"
        )
    
    # 获取文件扩展名
    file_ext = Path(file.filename).suffix.lower()
    allowed_extensions = {'.pdf', '.doc', '.docx', '.xls', '.xlsx', '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'}
    
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"File type not allowed. Allowed types: {', '.join(allowed_extensions)}"
        )
    
    try:
        # 读取文件内容
        file_content = await file.read()
        
        # 生成文件名
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).rstrip()
        filename = f"{current_employee.id}_{safe_name}_{timestamp}{file_ext}"
        
        # 获取存储子文件夹
        subfolder = DOCUMENT_TYPE_FOLDERS[document_type]
        
        # 保存文件
        file_url = await save_upload_file(file_content, filename, subfolder)
        
        # 创建文档记录
        document = EmployeeDocument(
            id=str(uuid.uuid4()),
            employee_id=current_employee.id,
            name=name,
            file_type=file_ext.lstrip('.'),
            file_url=file_url,
            document_type=document_type,
            uploaded_by=current_employee.id,
        )
        
        db.add(document)
        db.commit()
        db.refresh(document)
        
        logger.info(f"Document uploaded: {document.id}, type: {document_type}, employee: {current_employee.id}")
        
        return {
            "id": document.id,
            "name": document.name,
            "file_type": document.file_type,
            "file_url": document.file_url,
            "document_type": document.document_type,
            "uploaded_at": document.uploaded_at.isoformat() if document.uploaded_at else None,
            "message": "Document uploaded successfully"
        }
        
    except Exception as e:
        logger.error(f"Failed to upload document: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to upload document: {str(e)}")


@router.get("/me/documents/{document_id}")
async def get_my_document_detail(
    document_id: str,
    current_employee: Employee = Depends(get_current_employee),
    db: Session = Depends(get_db)
):
    """获取文档详情"""
    document = db.query(EmployeeDocument).filter(
        EmployeeDocument.id == document_id,
        EmployeeDocument.employee_id == current_employee.id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return {
        "id": document.id,
        "name": document.name,
        "file_type": document.file_type,
        "file_url": document.file_url,
        "document_type": document.document_type,
        "uploaded_at": document.uploaded_at.isoformat() if document.uploaded_at else None,
        "employee_signed_at": document.employee_signed_at.isoformat() if document.employee_signed_at else None,
        "employee_signed": bool(document.employee_signed_at),
        "employee_signature_x": document.employee_signature_x,
        "employee_signature_y": document.employee_signature_y,
        "employee_signature_width": document.employee_signature_width,
        "employee_signature_height": document.employee_signature_height,
        "employee_signature_page": document.employee_signature_page,
    }


@router.delete("/me/documents/{document_id}")
async def delete_my_document(
    document_id: str,
    current_employee: Employee = Depends(get_current_employee),
    db: Session = Depends(get_db)
):
    """删除员工文档"""
    document = db.query(EmployeeDocument).filter(
        EmployeeDocument.id == document_id,
        EmployeeDocument.employee_id == current_employee.id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    try:
        # 删除文件
        if document.file_url:
            delete_file(document.file_url)
        
        # 删除数据库记录
        db.delete(document)
        db.commit()
        
        logger.info(f"Document deleted: {document_id}, employee: {current_employee.id}")
        
        return {"message": "Document deleted successfully"}
        
    except Exception as e:
        logger.error(f"Failed to delete document: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete document: {str(e)}")


@router.get("/me/documents/{document_id}/preview")
async def preview_my_document(
    document_id: str,
    db: Session = Depends(get_db),
    token: Optional[str] = Query(default=None),
    credentials: Optional[HTTPAuthorizationCredentials] = Depends(optional_security),
    format: Optional[str] = Query(default=None),
):
    """预览员工文档（用于内嵌显示）"""
    access_token = token or (credentials.credentials if credentials else None)
    if not access_token:
        raise HTTPException(status_code=401, detail="缺少认证令牌")

    payload = decode_access_token(access_token)
    if payload is None:
        raise HTTPException(status_code=401, detail="无效的认证令牌")

    employee_id = payload.get("sub")
    if not employee_id:
        raise HTTPException(status_code=401, detail="无效的认证令牌")

    current_employee = db.query(Employee).filter(Employee.id == employee_id).first()
    if not current_employee:
        raise HTTPException(status_code=401, detail="员工不存在")

    document = db.query(EmployeeDocument).filter(
        EmployeeDocument.id == document_id,
        EmployeeDocument.employee_id == current_employee.id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if not document.file_url:
        raise HTTPException(status_code=404, detail="File not found")
    
    # 获取文件路径
    file_path = get_file_path(document.file_url)
    
    if not file_path or not file_path.exists():
        # 尝试其他路径
        upload_path = Path(settings.upload_dir).resolve()
        relative_path = document.file_url.replace("uploads/", "").replace("uploads\\", "")
        
        possible_paths = [
            upload_path / relative_path,
            Path(relative_path),
        ]
        
        for possible_path in possible_paths:
            if possible_path.exists() and possible_path.is_file():
                file_path = possible_path
                break
    
    if not file_path or not file_path.exists() or not file_path.is_file():
        raise HTTPException(status_code=404, detail="File not found on server")
    
    preview_format = (format or "").lower()
    if preview_format == "pdf" and file_path.suffix.lower() in [".doc", ".docx"]:
        try:
            file_path = _convert_office_to_pdf(file_path)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"docx转PDF失败: {e}")
        mime_type = "application/pdf"
    else:
        # 获取MIME类型
        mime_type, _ = mimetypes.guess_type(str(file_path))
        if not mime_type:
            mime_type = "application/octet-stream"
    
    display_name = document.name or "document"
    if file_path.suffix and not display_name.lower().endswith(file_path.suffix.lower()):
        display_name = f"{display_name}{file_path.suffix}"
    ascii_filename = to_ascii_filename(display_name)
    return FileResponse(
        path=str(file_path),
        media_type=mime_type,
        filename=ascii_filename,
        headers={
            "Content-Disposition": build_content_disposition(display_name, "inline"),
            "X-Frame-Options": "ALLOWALL"  # 允许在iframe中显示
        }
    )


@router.get("/me/documents/{document_id}/download")
async def download_my_document(
    document_id: str,
    current_employee: Employee = Depends(get_current_employee),
    db: Session = Depends(get_db)
):
    """下载员工文档"""
    document = db.query(EmployeeDocument).filter(
        EmployeeDocument.id == document_id,
        EmployeeDocument.employee_id == current_employee.id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if not document.file_url:
        raise HTTPException(status_code=404, detail="File not found")
    
    # 获取文件路径
    file_path = get_file_path(document.file_url)
    
    if not file_path or not file_path.exists():
        # 尝试其他路径
        upload_path = Path(settings.upload_dir).resolve()
        relative_path = document.file_url.replace("uploads/", "").replace("uploads\\", "")
        
        possible_paths = [
            upload_path / relative_path,
            Path(relative_path),
        ]
        
        for possible_path in possible_paths:
            if possible_path.exists() and possible_path.is_file():
                file_path = possible_path
                break
    
    if not file_path or not file_path.exists() or not file_path.is_file():
        raise HTTPException(status_code=404, detail="File not found on server")
    
    # 获取MIME类型
    mime_type, _ = mimetypes.guess_type(str(file_path))
    if not mime_type:
        mime_type = "application/octet-stream"
    
    display_name = document.name or "document"
    if file_path.suffix and not display_name.lower().endswith(file_path.suffix.lower()):
        display_name = f"{display_name}{file_path.suffix}"
    ascii_filename = to_ascii_filename(display_name)
    return FileResponse(
        path=str(file_path),
        media_type=mime_type,
        filename=ascii_filename,
        headers={
            "Content-Disposition": build_content_disposition(display_name, "attachment")
        }
    )


# ==================== 培训记录管理 API ====================

@router.get("/me/training-records")
async def get_my_training_records(
    category: str = None,
    current_employee: Employee = Depends(get_current_employee),
    db: Session = Depends(get_db)
):
    """获取当前员工的培训记录列表"""
    query = db.query(TrainingRecord).filter(
        TrainingRecord.employee_id == current_employee.id
    )
    
    if category:
        query = query.filter(TrainingRecord.category == category)
    
    records = query.order_by(TrainingRecord.completed_date.desc()).all()
    
    return [
        {
            "id": record.id,
            "name": record.name,
            "category": record.category,
            "status": record.status,
            "completed_date": record.completed_date.isoformat() if record.completed_date else None,
            "score": record.score,
            "has_certificate": record.has_certificate,
            "certificate_number": record.certificate_number,
            "certificate_url": f"/api/app/employees/me/training-records/{record.id}/certificate"
            if record.has_certificate and (record.certificate_url or record.certificate_number)
            else None,
            "certificate_obtained_date": record.certificate_obtained_date.isoformat() if record.certificate_obtained_date else None,
            "certificate_expiry_date": record.certificate_expiry_date.isoformat() if record.certificate_expiry_date else None,
            "training_institution": record.training_institution,
            "notes": record.notes,
            "created_by": getattr(record, 'created_by', None),
        }
        for record in records
    ]


@router.post("/me/training-records")
async def create_my_training_record(
    data: dict,
    current_employee: Employee = Depends(get_current_employee),
    db: Session = Depends(get_db)
):
    """创建培训记录（JSON格式）"""
    try:
        from datetime import datetime as dt
        
        # 解析完成日期
        completed_date_str = data.get("completed_date")
        if completed_date_str:
            if isinstance(completed_date_str, str):
                completed_date = dt.fromisoformat(completed_date_str.replace("Z", "+00:00"))
            else:
                completed_date = completed_date_str
        else:
            completed_date = datetime.utcnow()
        
        # 解析证书日期
        certificate_obtained_date = None
        if data.get("certificate_obtained_date"):
            date_str = data["certificate_obtained_date"]
            if isinstance(date_str, str):
                certificate_obtained_date = dt.fromisoformat(date_str.replace("Z", "+00:00"))
            else:
                certificate_obtained_date = date_str
        
        certificate_expiry_date = None
        if data.get("certificate_expiry_date"):
            date_str = data["certificate_expiry_date"]
            if isinstance(date_str, str):
                certificate_expiry_date = dt.fromisoformat(date_str.replace("Z", "+00:00"))
            else:
                certificate_expiry_date = date_str
        
        # 创建培训记录
        # 员工提交的记录必须为pending状态，需要后台审核
        record = TrainingRecord(
            id=str(uuid.uuid4()),
            employee_id=current_employee.id,
            name=data.get("name", ""),
            category=data.get("category"),
            status="pending",  # 员工提交强制为pending，需要后台审核
            completed_date=completed_date,
            score=data.get("score"),
            has_certificate=data.get("has_certificate", False),
            certificate_number=data.get("certificate_number"),
            certificate_url=data.get("certificate_url"),
            certificate_obtained_date=certificate_obtained_date,
            certificate_expiry_date=certificate_expiry_date,
            training_institution=data.get("training_institution"),
            notes=data.get("notes"),
            created_by="employee",  # 标识为员工提交
        )
        
        db.add(record)
        db.commit()
        db.refresh(record)
        
        logger.info(f"Training record created: {record.id}, employee: {current_employee.id}")

        try:
            admin_users = db.query(User).filter(or_(User.is_active == True, User.is_active.is_(None))).all()
            for u in admin_users:
                touch_business_unread(
                    db,
                    business_code="qualification",
                    receiver_user_id=str(u.id),
                    data_id=f"training_record_created:{record.id}",
                    scope_id=str(current_employee.id),
                    trigger_user_id=str(current_employee.id),
                )
                if record.has_certificate and (record.certificate_url or record.certificate_number):
                    touch_business_unread(
                        db,
                        business_code="employee_qualification",
                        receiver_user_id=str(u.id),
                        data_id=str(current_employee.id),
                        scope_id=str(current_employee.id),
                        trigger_user_id=str(current_employee.id),
                    )
            db.commit()
        except Exception:
            db.rollback()
        
        return {
            "id": record.id,
            "name": record.name,
            "category": record.category,
            "status": record.status,
            "completed_date": record.completed_date.isoformat() if record.completed_date else None,
            "score": record.score,
            "has_certificate": record.has_certificate,
            "certificate_number": record.certificate_number,
            "certificate_url": f"/api/app/employees/me/training-records/{record.id}/certificate"
            if record.has_certificate and (record.certificate_url or record.certificate_number)
            else None,
            "certificate_obtained_date": record.certificate_obtained_date.isoformat() if record.certificate_obtained_date else None,
            "certificate_expiry_date": record.certificate_expiry_date.isoformat() if record.certificate_expiry_date else None,
            "training_institution": record.training_institution,
            "notes": record.notes,
            "created_by": getattr(record, 'created_by', None),
            "message": "Training record created successfully"
        }
        
    except Exception as e:
        logger.error(f"Failed to create training record: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to create training record: {str(e)}")


@router.post("/me/training-records/upload")
async def create_my_training_record_with_file(
    name: str = Form(...),
    completed_date: str = Form(...),
    category: str = Form(None),
    status: str = Form("pending"),
    score: str = Form(None),
    has_certificate: bool = Form(False),
    certificate_number: str = Form(None),
    certificate_obtained_date: str = Form(None),
    certificate_expiry_date: str = Form(None),
    training_institution: str = Form(None),
    notes: str = Form(None),
    file: UploadFile = File(None),
    current_employee: Employee = Depends(get_current_employee),
    db: Session = Depends(get_db)
):
    """创建培训记录（带证书文件上传）"""
    try:
        from datetime import datetime as dt
        
        # 解析日期
        completed_dt = dt.fromisoformat(completed_date.replace("Z", "+00:00"))
        
        cert_obtained_dt = None
        if certificate_obtained_date:
            cert_obtained_dt = dt.fromisoformat(certificate_obtained_date.replace("Z", "+00:00"))
        
        cert_expiry_dt = None
        if certificate_expiry_date:
            cert_expiry_dt = dt.fromisoformat(certificate_expiry_date.replace("Z", "+00:00"))
        
        certificate_url = None
        
        # 处理证书文件上传
        if file and has_certificate and file.filename:
            file_content = await file.read()
            file_ext = Path(file.filename).suffix.lower()
            timestamp = dt.utcnow().strftime("%Y%m%d_%H%M%S")
            filename = f"{current_employee.id}_cert_{timestamp}{file_ext}"
            certificate_url = await save_upload_file(file_content, filename, "training_certificates")
        
        # 创建培训记录
        # 员工提交的记录必须为pending状态，需要后台审核
        record = TrainingRecord(
            id=str(uuid.uuid4()),
            employee_id=current_employee.id,
            name=name,
            category=category,
            status="pending",  # 员工提交强制为pending，需要后台审核
            completed_date=completed_dt,
            score=score,
            has_certificate=has_certificate,
            certificate_number=certificate_number,
            certificate_url=certificate_url,
            certificate_obtained_date=cert_obtained_dt,
            certificate_expiry_date=cert_expiry_dt,
            training_institution=training_institution,
            notes=notes,
            created_by="employee",  # 标识为员工提交
        )
        
        db.add(record)
        db.commit()
        db.refresh(record)
        
        logger.info(f"Training record created with file: {record.id}, employee: {current_employee.id}")

        try:
            admin_users = db.query(User).filter(or_(User.is_active == True, User.is_active.is_(None))).all()
            for u in admin_users:
                touch_business_unread(
                    db,
                    business_code="qualification",
                    receiver_user_id=str(u.id),
                    data_id=f"training_record_created:{record.id}",
                    scope_id=str(current_employee.id),
                    trigger_user_id=str(current_employee.id),
                )
                if record.has_certificate and (record.certificate_url or record.certificate_number):
                    touch_business_unread(
                        db,
                        business_code="employee_qualification",
                        receiver_user_id=str(u.id),
                        data_id=str(current_employee.id),
                        scope_id=str(current_employee.id),
                        trigger_user_id=str(current_employee.id),
                    )
            db.commit()
        except Exception:
            db.rollback()
        
        return {
            "id": record.id,
            "name": record.name,
            "category": record.category,
            "status": record.status,
            "completed_date": record.completed_date.isoformat() if record.completed_date else None,
            "score": record.score,
            "has_certificate": record.has_certificate,
            "certificate_number": record.certificate_number,
            "certificate_url": f"/api/app/employees/me/training-records/{record.id}/certificate"
            if record.has_certificate and (record.certificate_url or record.certificate_number)
            else None,
            "certificate_obtained_date": record.certificate_obtained_date.isoformat() if record.certificate_obtained_date else None,
            "certificate_expiry_date": record.certificate_expiry_date.isoformat() if record.certificate_expiry_date else None,
            "training_institution": record.training_institution,
            "notes": record.notes,
            "created_by": getattr(record, 'created_by', None),
            "message": "Training record created successfully"
        }
        
    except Exception as e:
        logger.error(f"Failed to create training record: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to create training record: {str(e)}")


@router.put("/me/training-records/{record_id}")
async def update_my_training_record(
    record_id: str,
    data: dict,
    current_employee: Employee = Depends(get_current_employee),
    db: Session = Depends(get_db)
):
    """更新培训记录"""
    record = db.query(TrainingRecord).filter(
        TrainingRecord.id == record_id,
        TrainingRecord.employee_id == current_employee.id
    ).first()
    
    if not record:
        raise HTTPException(status_code=404, detail="Training record not found")
    
    try:
        from datetime import datetime as dt
        
        # 更新字段
        if "name" in data:
            record.name = data["name"]
        if "category" in data:
            record.category = data["category"]
        if "status" in data:
            record.status = data["status"]
        if "score" in data:
            record.score = data["score"]
        if "has_certificate" in data:
            record.has_certificate = data["has_certificate"]
        if "certificate_number" in data:
            record.certificate_number = data["certificate_number"]
        if "training_institution" in data:
            record.training_institution = data["training_institution"]
        if "notes" in data:
            record.notes = data["notes"]
        
        # 更新日期字段
        if "completed_date" in data and data["completed_date"]:
            date_str = data["completed_date"]
            if isinstance(date_str, str):
                record.completed_date = dt.fromisoformat(date_str.replace("Z", "+00:00"))
            else:
                record.completed_date = date_str
        
        if "certificate_obtained_date" in data and data["certificate_obtained_date"]:
            date_str = data["certificate_obtained_date"]
            if isinstance(date_str, str):
                record.certificate_obtained_date = dt.fromisoformat(date_str.replace("Z", "+00:00"))
            else:
                record.certificate_obtained_date = date_str
        
        if "certificate_expiry_date" in data and data["certificate_expiry_date"]:
            date_str = data["certificate_expiry_date"]
            if isinstance(date_str, str):
                record.certificate_expiry_date = dt.fromisoformat(date_str.replace("Z", "+00:00"))
            else:
                record.certificate_expiry_date = date_str
        
        db.commit()
        db.refresh(record)
        
        logger.info(f"Training record updated: {record_id}, employee: {current_employee.id}")
        
        return {
            "id": record.id,
            "name": record.name,
            "category": record.category,
            "status": record.status,
            "completed_date": record.completed_date.isoformat() if record.completed_date else None,
            "score": record.score,
            "has_certificate": record.has_certificate,
            "certificate_number": record.certificate_number,
            "certificate_url": f"/api/app/employees/me/training-records/{record.id}/certificate"
            if record.has_certificate and (record.certificate_url or record.certificate_number)
            else None,
            "certificate_obtained_date": record.certificate_obtained_date.isoformat() if record.certificate_obtained_date else None,
            "certificate_expiry_date": record.certificate_expiry_date.isoformat() if record.certificate_expiry_date else None,
            "training_institution": record.training_institution,
            "notes": record.notes,
            "created_by": getattr(record, 'created_by', None),
            "message": "Training record updated successfully"
        }
        
    except Exception as e:
        logger.error(f"Failed to update training record: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to update training record: {str(e)}")


@router.delete("/me/training-records/{record_id}")
async def delete_my_training_record(
    record_id: str,
    current_employee: Employee = Depends(get_current_employee),
    db: Session = Depends(get_db)
):
    """删除培训记录"""
    record = db.query(TrainingRecord).filter(
        TrainingRecord.id == record_id,
        TrainingRecord.employee_id == current_employee.id
    ).first()
    
    if not record:
        raise HTTPException(status_code=404, detail="Training record not found")
    
    try:
        # 删除证书文件
        if record.certificate_url:
            delete_file(record.certificate_url)
        
        db.delete(record)
        db.commit()
        
        logger.info(f"Training record deleted: {record_id}, employee: {current_employee.id}")
        
        return {"message": "Training record deleted successfully"}
        
    except Exception as e:
        logger.error(f"Failed to delete training record: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete training record: {str(e)}")


# ==================== 合同生成 API ====================

@router.post("/me/contracts/generate")
async def generate_my_contract(
    data: dict,
    current_employee: Employee = Depends(get_current_employee),
    db: Session = Depends(get_db)
):
    """生成员工合同"""
    try:
        from docx import Document
        from docx.shared import RGBColor
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx 库未安装，无法生成合同")
    
    # 获取参数
    start_date_str = data.get("start_date")
    employment_type = data.get("employment_type", "full-time")
    position = data.get("position", "support-worker")
    superior_first_name = data.get("superior_first_name", "")
    superior_last_name = data.get("superior_last_name", "")
    superior_title = data.get("superior_title", "")
    hours_per_week = data.get("hours_per_week", 37.5)
    work_hours = data.get("work_hours", "9 am to 5 pm")
    gross_salary = data.get("gross_salary", 0)
    
    # 获取模板文件路径
    template_path = Path(__file__).parent.parent.parent.parent / "static" / "Staff Contract Sample.docx"
    if not template_path.exists():
        raise HTTPException(status_code=404, detail="合同模板文件不存在")
    
    try:
        # 读取模板
        doc = Document(str(template_path))
    except Exception as e:
        logger.error(f"读取合同模板失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"读取合同模板失败: {str(e)}")

    # 格式化日期
    def format_date(date_str):
        if not date_str:
            return ''
        try:
            date_obj = datetime.strptime(date_str, '%Y-%m-%d')
            return date_obj.strftime('%d/%m/%Y')
        except Exception:
            return date_str

    # 格式化薪资（不带 $，避免模板内已有 $ 时出现双 $）
    def format_salary(salary):
        if salary is None:
            return ''
        return f"{float(salary):,.2f}"

    # 准备替换数据
    start_date = format_date(start_date_str)
    hours_per_week_str = str(hours_per_week) if hours_per_week else ''
    work_hours = work_hours or ''
    gross_salary = format_salary(gross_salary)
    superior_first_name = superior_first_name or ''
    superior_last_name = superior_last_name or ''
    superior_title = superior_title or ''

    # 构建替换映射（只替换模板中的红字占位符）
    replacements = []

    # Start Date
    if start_date:
        replacements.append(('Start Date', start_date))

    # permanent full-time（按表单选择的雇佣类型文本替换）
    employment_type_text = ''
    if employment_type == 'full-time':
        employment_type_text = 'Permanent Full-time'
    elif employment_type == 'part-time':
        employment_type_text = 'Part time'
    elif employment_type == 'casual':
        employment_type_text = 'Casual'
    if employment_type_text:
        replacements.append(('permanent full-time', employment_type_text))

    # Position（按表单选择的职位文本替换）
    position_text = ''
    if position == 'support-worker':
        position_text = 'Support Worker'
    elif position == 'admin':
        position_text = 'Admin'
    elif position == 'office-staff':
        position_text = 'Office Staff'
    if position_text:
        replacements.append(('Position', position_text))

    # Superior first name Superior last name, Superior title（包含 fi 连字）
    superior_combined = ''
    if superior_first_name or superior_last_name or superior_title:
        full_name = f"{superior_first_name} {superior_last_name}".strip()
        superior_combined = f"{full_name}, {superior_title}".strip(', ')
    if superior_combined:
        replacements.append(('Superior first name Superior last name, Superior title', superior_combined))
        replacements.append(('Superior ﬁrst name Superior last name, Superior title', superior_combined))

    # 37.5 hours per week (for a full-time employee)
    if hours_per_week_str:
        replacements.append(('37.5 hours per week (for a full-time employee)', f'{hours_per_week_str} hours per week (for a full-time employee)'))

    # 9 am to 5 pm
    if work_hours:
        replacements.append(('9 am to 5 pm', work_hours))

    # Insert gross salary
    if gross_salary:
        replacements.append(('Insert gross salary', gross_salary))
        replacements.append(('$Insert gross salary', f'${gross_salary}'))

    def replace_text_in_paragraph(paragraph):
        runs = list(paragraph.runs)
        punctuation_pattern = re.compile(r'^[\s,.:;]+$')

        def replace_match_in_runs(match_start, match_end, new_text):
            pos = 0
            start_idx = end_idx = None
            start_offset = end_offset = 0
            for idx, run in enumerate(runs):
                run_text = run.text
                run_len = len(run_text)
                if start_idx is None and pos + run_len > match_start:
                    start_idx = idx
                    start_offset = match_start - pos
                if pos + run_len >= match_end:
                    end_idx = idx
                    end_offset = match_end - pos
                    break
                pos += run_len

            if start_idx is None or end_idx is None:
                return

            def set_black(run):
                try:
                    # 尝试直接设置颜色
                    run.font.color.rgb = RGBColor(0, 0, 0)
                except (AttributeError, TypeError):
                    # 如果color为None或无法设置，跳过颜色设置
                    # 这不会影响文本替换功能
                    pass

            if start_idx == end_idx:
                run = runs[start_idx]
                run.text = run.text[:start_offset] + new_text + run.text[end_offset:]
                set_black(run)
                # 邻近标点可能是红色，统一改黑
                if start_idx + 1 < len(runs) and punctuation_pattern.match(runs[start_idx + 1].text):
                    set_black(runs[start_idx + 1])
                if start_idx - 1 >= 0 and punctuation_pattern.match(runs[start_idx - 1].text):
                    set_black(runs[start_idx - 1])
                return

            first = runs[start_idx]
            last = runs[end_idx]
            prefix = first.text[:start_offset]
            suffix = last.text[end_offset:]
            first.text = prefix + new_text
            set_black(first)
            for idx in range(start_idx + 1, end_idx):
                runs[idx].text = ''
                set_black(runs[idx])
            last.text = suffix
            set_black(last)
            # 邻近标点可能是红色，统一改黑
            if end_idx + 1 < len(runs) and punctuation_pattern.match(runs[end_idx + 1].text):
                set_black(runs[end_idx + 1])
            if start_idx - 1 >= 0 and punctuation_pattern.match(runs[start_idx - 1].text):
                set_black(runs[start_idx - 1])

        for old_text, new_text in replacements:
            if not old_text:
                continue
            try:
                pattern = re.compile(re.escape(old_text), re.IGNORECASE)
                max_iterations = 100  # 防止无限循环
                iteration = 0
                while iteration < max_iterations:
                    combined_text = ''.join(run.text for run in runs)
                    match = pattern.search(combined_text)
                    if not match:
                        break
                    replace_match_in_runs(match.start(), match.end(), new_text)
                    iteration += 1
                if iteration >= max_iterations:
                    logger.warning(f"替换文本 '{old_text}' 达到最大迭代次数，可能存在问题")
            except Exception as e:
                logger.error(f"替换文本 '{old_text}' 时出错: {e}", exc_info=True)
                # 继续处理其他替换，不中断整个流程

    try:
        # 处理段落
        for paragraph in doc.paragraphs:
            try:
                replace_text_in_paragraph(paragraph)
            except Exception as e:
                logger.error(f"处理段落时出错: {e}", exc_info=True)
                # 继续处理其他段落

        # 处理表格
        for table in doc.tables:
            try:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            try:
                                replace_text_in_paragraph(paragraph)
                            except Exception as e:
                                logger.error(f"处理表格单元格段落时出错: {e}", exc_info=True)
                                # 继续处理其他单元格
            except Exception as e:
                logger.error(f"处理表格时出错: {e}", exc_info=True)
                # 继续处理其他表格
    except Exception as e:
        logger.error(f"处理Word文档时出错: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"处理Word文档失败: {str(e)}")
    
    # 保存生成的合同
    output_filename = f"contract_{current_employee.id}_{uuid.uuid4()}.docx"
    subfolder = 'employee_contracts'
    tmp_file_path = None
    
    try:
        # 创建临时文件保存生成的合同
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            doc.save(tmp_file.name)
            tmp_file_path = tmp_file.name
        
        # 读取临时文件内容
        with open(tmp_file_path, 'rb') as f:
            file_content = f.read()
        
        # 删除临时文件
        try:
            os.unlink(tmp_file_path)
            tmp_file_path = None
        except Exception as e:
            logger.warning(f"删除临时文件失败: {e}")
        
        # 上传文件
        file_url = await save_upload_file(file_content, output_filename, subfolder=subfolder)
        
        # 创建文档记录
        document = EmployeeDocument(
            id=str(uuid.uuid4()),
            employee_id=current_employee.id,
            name=f"Employment Contract - {start_date}",
            file_type='docx',
            file_url=file_url,
            document_type='contract',
            uploaded_by=current_employee.id
        )
        db.add(document)
        db.commit()
        db.refresh(document)
        
        logger.info(f"Contract generated: {document.id}, employee: {current_employee.id}")
        
        return {
            "id": document.id,
            "name": document.name,
            "file_type": document.file_type,
            "file_url": document.file_url,
            "document_type": document.document_type,
            "uploaded_at": document.uploaded_at.isoformat() if document.uploaded_at else None,
            "message": "Contract generated successfully"
        }
    except Exception as e:
        # 如果出错，回滚数据库事务
        db.rollback()
        logger.error(f"保存合同文件失败: {e}", exc_info=True)
        
        # 清理临时文件
        if tmp_file_path and os.path.exists(tmp_file_path):
            try:
                os.unlink(tmp_file_path)
            except Exception:
                pass
        
        raise HTTPException(status_code=500, detail=f"保存合同文件失败: {str(e)}")


# ==================== 合同签字 API ====================

class EmployeeContractSignPreviewBody(BaseModel):
    signature_data: str
    x: float | None = None
    y: float | None = None
    width: float | None = None
    height: float | None = None
    page: int | None = None


@router.post("/me/contracts/{contract_id}/sign-preview")
async def preview_my_contract_signature(
    contract_id: str,
    body: EmployeeContractSignPreviewBody,
    background_tasks: BackgroundTasks,
    current_employee: Employee = Depends(get_current_employee),
    db: Session = Depends(get_db),
):
    document = db.query(EmployeeDocument).filter(
        EmployeeDocument.id == contract_id,
        EmployeeDocument.employee_id == current_employee.id,
        EmployeeDocument.document_type == "contract",
    ).first()

    if not document or not document.file_url:
        raise HTTPException(status_code=404, detail="合同不存在")

    mime_type, blob = _parse_data_url(body.signature_data)
    if not blob:
        raise HTTPException(status_code=400, detail="无效的签名数据")

    contract_path = get_file_path(document.file_url)
    if not contract_path or not contract_path.exists():
        raise HTTPException(status_code=404, detail="合同文件不存在")

    if contract_path.suffix.lower() in [".doc", ".docx"]:
        try:
            contract_path = _convert_office_to_pdf(contract_path)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"docx转PDF失败: {e}")

    if contract_path.suffix.lower() != ".pdf":
        raise HTTPException(status_code=400, detail="暂不支持该文件类型的预览")

    x = body.x if body.x is not None else document.employee_signature_x
    y = body.y if body.y is not None else document.employee_signature_y
    width = body.width if body.width is not None else document.employee_signature_width
    height = body.height if body.height is not None else document.employee_signature_height
    page = body.page if body.page is not None else (document.employee_signature_page or 0)

    try:
        sydney_tz = pytz.timezone("Australia/Sydney")
        now_sydney = datetime.now(sydney_tz)
        date_text = now_sydney.strftime("%Y/%m/%d")
    except Exception:
        date_text = datetime.now().strftime("%Y/%m/%d")

    try:
        tmp_dir = Path(tempfile.gettempdir()) / "empowerhub_sign_previews"
        tmp_dir.mkdir(parents=True, exist_ok=True)
        tmp_pdf_path = tmp_dir / f"{contract_id}_{uuid.uuid4().hex}.pdf"
        shutil.copyfile(str(contract_path), str(tmp_pdf_path))

        embed_success = _embed_signature_to_pdf(
            tmp_pdf_path,
            blob,
            "employee",
            x=x,
            y=y,
            width=width,
            height=height,
            page_index=page,
            date_text=date_text,
        )
        if embed_success is False:
            raise RuntimeError("签字预览生成失败")

        def _cleanup(path: str):
            try:
                os.remove(path)
            except Exception:
                pass

        background_tasks.add_task(_cleanup, str(tmp_pdf_path))

        display_name = document.name or "contract"
        if not display_name.lower().endswith(".pdf"):
            display_name = f"{display_name}.pdf"
        ascii_filename = to_ascii_filename(display_name)
        return FileResponse(
            path=str(tmp_pdf_path),
            media_type="application/pdf",
            filename=ascii_filename,
            headers={
                "Content-Disposition": build_content_disposition(display_name, "inline"),
                "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
                "Pragma": "no-cache",
            },
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"生成签字预览失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"生成签字预览失败: {e}")


@router.post("/me/contracts/{contract_id}/signature-position")
async def save_signature_position(
    contract_id: str,
    x: float = Body(...),
    y: float = Body(...),
    width: float = Body(...),
    height: float = Body(...),
    page: int = Body(0),
    current_employee: Employee = Depends(get_current_employee),
    db: Session = Depends(get_db)
):
    """保存员工合同签字坐标位置"""
    try:
        # 查找合同文档
        document = db.query(EmployeeDocument).filter(
            EmployeeDocument.id == contract_id,
            EmployeeDocument.employee_id == current_employee.id,
            EmployeeDocument.document_type == "contract"
        ).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="合同不存在")
        
        # 保存坐标
        document.employee_signature_x = x
        document.employee_signature_y = y
        document.employee_signature_width = width
        document.employee_signature_height = height
        document.employee_signature_page = page
        
        db.commit()
        db.refresh(document)
        
        logger.info(f"Signature position saved for contract: {document.id}, employee: {current_employee.id}")
        
        return {
            "id": document.id,
            "employee_signature_x": document.employee_signature_x,
            "employee_signature_y": document.employee_signature_y,
            "employee_signature_width": document.employee_signature_width,
            "employee_signature_height": document.employee_signature_height,
            "employee_signature_page": document.employee_signature_page,
            "message": "签字坐标保存成功"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to save signature position: {e}")
        raise HTTPException(status_code=500, detail=f"保存签字坐标失败: {str(e)}")


@router.post("/me/contracts/{contract_id}/sign")
async def sign_my_contract(
    contract_id: str,
    signature_data: str = Body(..., embed=True),
    lang: str | None = Query(default=None),
    current_employee: Employee = Depends(get_current_employee),
    db: Session = Depends(get_db)
):
    """员工提交合同签名"""
    try:
        resolved_lang = "en" if str(lang or "").strip().lower().startswith("en") else "zh"
        # 查找合同文档
        document = db.query(EmployeeDocument).filter(
            EmployeeDocument.id == contract_id,
            EmployeeDocument.employee_id == current_employee.id,
            EmployeeDocument.document_type == "contract"
        ).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Contract not found" if resolved_lang == "en" else "合同不存在")
        
        # 解析签名数据
        mime_type, blob = _parse_data_url(signature_data)
        if not blob:
            raise HTTPException(status_code=400, detail="Invalid signature data" if resolved_lang == "en" else "无效的签名数据")
        
        # 保存签名到数据库
        document.employee_signature_blob = blob
        document.employee_signature_mime = mime_type or "image/png"
        document.employee_signed_at = datetime.utcnow()
        
        # 将签名嵌入到文档中（使用保存的坐标，如果有）
        if document.file_url:
            contract_path = get_file_path(document.file_url)
            if contract_path and contract_path.exists():
                try:
                    x = document.employee_signature_x
                    y = document.employee_signature_y
                    width = document.employee_signature_width
                    height = document.employee_signature_height
                    page = document.employee_signature_page or 0

                    # doc/docx 多页签字：必须先转为PDF再嵌入，避免总是落在第一页
                    original_path = contract_path
                    if contract_path.suffix.lower() in [".doc", ".docx"]:
                        logger.info(f"检测到docx文件，开始转换为PDF: {contract_path}")
                        try:
                            pdf_path = _convert_docx_to_pdf_in_place(contract_path)
                            if not pdf_path or not pdf_path.exists():
                                raise RuntimeError(f"PDF转换失败：转换后的文件不存在: {pdf_path}")
                            logger.info(f"docx转换为PDF成功: {pdf_path}")
                            contract_path = pdf_path
                            # 更新数据库中的文件URL和类型
                            document.file_url = str(pdf_path)
                            document.file_type = "pdf"
                            logger.info(f"已更新文档file_url为PDF路径: {document.file_url}, file_type: {document.file_type}")
                        except Exception as convert_error:
                            logger.error(f"docx转PDF失败: {convert_error}", exc_info=True)
                            raise HTTPException(status_code=500, detail=f"docx转PDF失败: {str(convert_error)}")
                    
                    # 验证转换后的文件必须是PDF
                    if contract_path.suffix.lower() != ".pdf":
                        error_msg = f"文件类型错误：转换后应该是PDF，但实际是 {contract_path.suffix}。文件路径: {contract_path}"
                        logger.error(error_msg)
                        raise RuntimeError(error_msg)
                    
                    # 抓取澳洲时间并格式化为 YYYY/MM/DD
                    try:
                        sydney_tz = pytz.timezone('Australia/Sydney')
                        now_sydney = datetime.now(sydney_tz)
                        date_text = now_sydney.strftime("%Y/%m/%d")
                    except Exception:
                        date_text = datetime.now().strftime("%Y/%m/%d")

                    # 嵌入签名到PDF
                    logger.info(f"开始嵌入签名到PDF: {contract_path}, 坐标: x={x}, y={y}, width={width}, height={height}, page={page}")
                    embed_success = _embed_signature_to_pdf(
                        contract_path,
                        blob,
                        "employee",
                        x=x,
                        y=y,
                        width=width,
                        height=height,
                        page_index=page,
                        date_text=date_text,
                    )
                    
                    if embed_success is False:
                        raise RuntimeError("嵌入签名失败")
                    logger.info(f"签名嵌入成功: {contract_path}")
                except Exception as e:
                    logger.error(f"嵌入签名到合同失败: {e}", exc_info=True)
                    raise HTTPException(status_code=500, detail=f"嵌入签名失败: {e}")
        
        db.commit()
        db.refresh(document)
        
        logger.info(f"Contract signed by employee: {document.id}, employee: {current_employee.id}")

        try:
            admin_users = db.query(User).filter(or_(User.is_active == True, User.is_active.is_(None))).all()
            for u in admin_users:
                touch_business_unread(
                    db,
                    business_code="employee",
                    receiver_user_id=str(u.id),
                    data_id=str(current_employee.id),
                    scope_id=str(current_employee.id),
                    trigger_user_id=str(current_employee.id),
                )
            db.commit()
        except Exception:
            db.rollback()
        
        return {
            "id": document.id,
            "employee_signed_at": document.employee_signed_at.isoformat() if document.employee_signed_at else None,
            "lang": resolved_lang,
            "message": "Contract signing submitted successfully" if resolved_lang == "en" else "合同签名提交成功"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to sign contract: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to submit contract signing: {str(e)}" if (str(lang or '').strip().lower().startswith('en')) else f"提交合同签名失败: {str(e)}")


@router.get("/me/contracts/{contract_id}/employee-signature/image")
async def get_my_contract_employee_signature_image(
    contract_id: str,
    current_employee: Employee = Depends(get_current_employee),
    db: Session = Depends(get_db)
):
    """获取员工合同签名图片"""
    document = db.query(EmployeeDocument).filter(
        EmployeeDocument.id == contract_id,
        EmployeeDocument.employee_id == current_employee.id,
        EmployeeDocument.document_type == "contract"
    ).first()
    
    if not document or not document.employee_signature_blob:
        raise HTTPException(status_code=404, detail="签名不存在")
    
    return Response(
        content=document.employee_signature_blob,
        media_type=document.employee_signature_mime or "image/png"
    )
