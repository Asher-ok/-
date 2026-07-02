from datetime import datetime
import re
import uuid
import shutil
import tempfile
from pathlib import Path

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Query
from fastapi.responses import FileResponse, Response
from sqlalchemy.orm import Session

from core.database import get_db
from core.utils.file_utils import (
    save_upload_file,
    get_file_path,
    ensure_upload_dir,
    build_content_disposition,
    to_ascii_filename,
)
from shared.models import TemplateFile
from ..dependencies import get_current_user


router = APIRouter(prefix="/api/houtai/template-files", tags=["管理-模板文件"])


PLACEHOLDER_RE = re.compile(r"\$\{([^}]+)\}|\{([^}]+)\}")


def _extract_placeholders_from_docx(path: Path) -> list[str]:
    try:
        from docx import Document
    except Exception:
        return []
    try:
        doc = Document(str(path))
    except Exception:
        return []
    values: set[str] = set()

    def scan_text(text: str | None) -> None:
        if not text:
            return
        for m in PLACEHOLDER_RE.finditer(text):
            key = (m.group(1) or m.group(2) or "").strip()
            if key:
                values.add(key)

    for p in doc.paragraphs:
        scan_text(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                scan_text(cell.text)
    return sorted(values)


def _resolve_soffice_executable() -> str | None:
    try:
        from modules.houtai.api.employees import _resolve_soffice_executable as _resolve
        return _resolve()
    except Exception:
        return None


def _convert_office_file(source: Path, target_ext: str, out_dir: Path) -> Path | None:
    soffice = _resolve_soffice_executable()
    if not soffice:
        return None
    out_dir.mkdir(parents=True, exist_ok=True)
    soffice_dir = str(Path(soffice).resolve().parent)
    import subprocess

    cmd = [
        soffice,
        "--headless",
        "--nologo",
        "--nolockcheck",
        "--norestore",
        "--convert-to",
        target_ext.lstrip("."),
        "--outdir",
        str(out_dir),
        str(source),
    ]
    proc = subprocess.run(cmd, cwd=soffice_dir, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    if proc.returncode != 0:
        return None
    expected = out_dir / f"{source.stem}.{target_ext.lstrip('.')}"
    if expected.exists():
        return expected
    candidates = list(out_dir.glob(f"{source.stem}.*"))
    for c in candidates:
        if c.suffix.lower() == f".{target_ext.lstrip('.').lower()}":
            return c
    return None


def _extract_placeholders_from_pdf(path: Path) -> list[str]:
    try:
        from PyPDF2 import PdfReader
    except Exception:
        return []
    try:
        reader = PdfReader(str(path))
    except Exception:
        return []
    try:
        fields = reader.get_fields() or {}
    except Exception:
        fields = {}
    values: set[str] = set()
    for k in fields.keys():
        raw = (k or "").strip()
        if not raw:
            continue
        m = re.fullmatch(r"\$\{([^}]+)\}", raw)
        if m:
            values.add(m.group(1).strip())
        else:
            values.add(raw)
    return sorted(values)


@router.get("")
async def list_template_files(
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    rows = db.query(TemplateFile).order_by(TemplateFile.create_time.desc()).all()
    return [
        {
            "id": r.id,
            "template_name": r.template_name,
            "file_name": r.file_name,
            "file_type": r.file_type,
            "created_at": r.create_time.isoformat() if r.create_time else None,
            "updated_at": r.update_time.isoformat() if r.update_time else None,
        }
        for r in rows
    ]


@router.post("")
async def create_template_file(
    template_name: str = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    original_name = (file.filename or "").strip() or "template"
    ext = Path(original_name).suffix.lower().lstrip(".")
    if ext not in ("doc", "docx", "pdf"):
        raise HTTPException(status_code=400, detail="仅支持 doc/docx/pdf")
    content = await file.read()
    row = TemplateFile(
        id=str(uuid.uuid4()),
        template_name=template_name.strip(),
        file_name=original_name,
        file_type=ext,
        file_url="",
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    stored_name = f"{row.id}.{ext}"
    file_url = await save_upload_file(content, stored_name, subfolder="template_files")
    row.file_url = file_url
    db.commit()
    return {"id": row.id}


@router.put("/{template_id}")
async def update_template_file(
    template_id: str,
    template_name: str | None = Form(default=None),
    file: UploadFile | None = File(default=None),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    row = db.query(TemplateFile).filter(TemplateFile.id == template_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="模板不存在")

    if template_name is not None and template_name.strip():
        row.template_name = template_name.strip()

    if file is not None:
        original_name = (file.filename or "").strip() or row.file_name
        ext = Path(original_name).suffix.lower().lstrip(".")
        if ext not in ("doc", "docx", "pdf"):
            raise HTTPException(status_code=400, detail="仅支持 doc/docx/pdf")
        content = await file.read()
        stored_name = f"{row.id}.{ext}"
        file_url = await save_upload_file(content, stored_name, subfolder="template_files")
        row.file_name = original_name
        row.file_type = ext
        row.file_url = file_url
        row.update_time = datetime.utcnow()

    db.commit()
    return {"ok": True}


@router.delete("/{template_id}")
async def delete_template_file(
    template_id: str,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    row = db.query(TemplateFile).filter(TemplateFile.id == template_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="模板不存在")
    p = get_file_path(row.file_url)
    if p and p.exists():
        try:
            p.unlink(missing_ok=True)
        except Exception:
            pass
    db.delete(row)
    db.commit()
    return {"ok": True}


@router.get("/{template_id}/placeholders")
async def get_template_placeholders(
    template_id: str,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    row = db.query(TemplateFile).filter(TemplateFile.id == template_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="模板不存在")
    p = get_file_path(row.file_url)
    if not p or not p.exists():
        raise HTTPException(status_code=404, detail="文件不存在")
    ext = (row.file_type or p.suffix.lstrip(".")).lower()
    if ext == "docx":
        placeholders = _extract_placeholders_from_docx(p)
    elif ext == "pdf":
        placeholders = _extract_placeholders_from_pdf(p)
    else:
        with tempfile.TemporaryDirectory() as td:
            out_dir = Path(td)
            converted = _convert_office_file(p, "docx", out_dir)
            placeholders = _extract_placeholders_from_docx(converted) if converted else []
    return {"placeholders": placeholders}


@router.get("/{template_id}/download")
async def download_template_file(
    template_id: str,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    row = db.query(TemplateFile).filter(TemplateFile.id == template_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="模板不存在")
    p = get_file_path(row.file_url)
    if not p or not p.exists():
        raise HTTPException(status_code=404, detail="文件不存在")
    name = row.file_name or p.name
    return FileResponse(
        str(p),
        filename=to_ascii_filename(name),
        media_type="application/octet-stream",
        headers={"Content-Disposition": build_content_disposition(name, "attachment")},
    )


@router.get("/{template_id}/preview")
async def preview_template_file(
    template_id: str,
    format: str | None = Query(default="pdf", description="预览格式：pdf / origin"),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    row = db.query(TemplateFile).filter(TemplateFile.id == template_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="模板不存在")
    p = get_file_path(row.file_url)
    if not p or not p.exists():
        raise HTTPException(status_code=404, detail="文件不存在")

    ext = (row.file_type or p.suffix.lstrip(".")).lower()
    preview_name = row.file_name or p.name

    if (format or "").lower() == "origin":
        media_type = "application/pdf" if ext == "pdf" else "application/octet-stream"
        return FileResponse(
            str(p),
            filename=to_ascii_filename(preview_name),
            media_type=media_type,
            headers={"Content-Disposition": build_content_disposition(preview_name, "inline")},
        )

    if ext == "pdf":
        return FileResponse(
            str(p),
            filename=to_ascii_filename(preview_name),
            media_type="application/pdf",
            headers={"Content-Disposition": build_content_disposition(preview_name, "inline")},
        )

    if ext in ("doc", "docx"):
        uploads_dir = ensure_upload_dir()
        cache_dir = (uploads_dir / "template_files_previews").resolve()
        cache_dir.mkdir(parents=True, exist_ok=True)
        stem = f"{template_id}_{uuid.uuid4().hex}"
        source_copy = cache_dir / f"{stem}.{ext}"
        try:
            shutil.copyfile(str(p), str(source_copy))
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"预览失败：复制文件失败: {e}")

        converted = _convert_office_file(source_copy, "pdf", cache_dir)
        if not converted or not converted.exists():
            raise HTTPException(status_code=500, detail="doc/docx 转 PDF 预览失败（请检查 LibreOffice/soffice 配置）")

        pdf_name = f"{Path(preview_name).stem}.pdf"
        try:
            pdf_bytes = converted.read_bytes()
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"预览失败：读取转换后的PDF失败: {e}")
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": build_content_disposition(pdf_name, "inline")},
        )

    raise HTTPException(status_code=400, detail="该文件类型不支持预览")
