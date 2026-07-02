from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Query, Request, Form, Body
from fastapi.responses import FileResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from sqlalchemy.orm import Session
from sqlalchemy import func, or_
from typing import List
from datetime import datetime, timedelta
import json
import uuid
import re
import tempfile
from pathlib import Path
from core.database import get_db
from core.utils.file_utils import save_upload_file, get_file_path
from core.utils.email import send_contact_email
from core.config import settings
from shared.models import Customer, CustomerDocument, Task, TaskStatus, InvoiceServiceLevel1, BusinessUnread, TemplateFile
from shared.models.customer import CUSTOMER_STATUS_ARCHIVED, CUSTOMER_STATUS_UNARCHIVED, CUSTOMER_STATUS_PENDING_ARCHIVE
from shared.models.customer_document import STATUS_SIGNED
from shared.models.update_notification import touch_business_unread
from ..schemas.customer import CustomerCreate, CustomerUpdate, CustomerResponse, CustomerAttachment
from ..schemas.customer_document import CustomerDocumentResponse
from ..dependencies import get_current_user
from core.auth import decode_access_token
from shared.models import User

security_optional = HTTPBearer(auto_error=False)

router = APIRouter(prefix="/api/houtai/customers", tags=["管理-客户"])

PLACEHOLDER_RE = re.compile(r"\$\{([^}]+)\}|\{([^}]+)\}")


def _customer_placeholder_values(customer: Customer) -> dict[str, str]:
    accepted_names: list[str] = []
    try:
        accepted_names = [
            str(svc.name).strip()
            for svc in (getattr(customer, "accepted_service_level1", None) or [])
            if getattr(svc, "name", None)
        ]
    except Exception:
        accepted_names = []
    accepted_services = ", ".join([n for n in accepted_names if n])

    raw = {
        "name": customer.name,
        "customer_name": customer.name,
        "客户姓名": customer.name,
        "姓名": customer.name,

        "phone": customer.phone,
        "mobile": customer.phone,
        "手机号": customer.phone,
        "电话": customer.phone,

        "address": customer.address,
        "地址": customer.address,

        "email": customer.email or "",
        "邮箱": customer.email or "",

        "customer_code": customer.customer_code or "",
        "客户编号": customer.customer_code or "",

        "customer_type": customer.customer_type or "",
        "客户类型": customer.customer_type or "",
        "类型": customer.customer_type or "",

        "gender": customer.gender or "",
        "性别": customer.gender or "",

        "age": str(customer.age) if customer.age is not None else "",
        "年龄": str(customer.age) if customer.age is not None else "",

        "disability_type": customer.disability_type or "",
        "残疾类型": customer.disability_type or "",

        "weekly_service_hours": f"{customer.weekly_service_hours:.2f}" if customer.weekly_service_hours is not None else "",
        "每周服务时长": f"{customer.weekly_service_hours:.2f}" if customer.weekly_service_hours is not None else "",
        "每周服务时长(小时)": f"{customer.weekly_service_hours:.2f}" if customer.weekly_service_hours is not None else "",

        "accepted_services": accepted_services,
        "accepted_service_level1": accepted_services,
        "accepted_service_level1_names": accepted_services,
        "可接受的服务": accepted_services,
        "可接受服务": accepted_services,

        "ndis_number": customer.ndis_number,
        "NDIS号码": customer.ndis_number,
    }
    values: dict[str, str] = {}
    for k, v in raw.items():
        if v is None:
            continue
        s = str(v).strip()
        if s:
            values[str(k).strip()] = s
    for k in list(values.keys()):
        values[k.strip().lower()] = values[k]
    return values


def _replace_placeholders_in_text(text: str, mapping: dict[str, str]) -> str:
    if not text:
        return text

    def repl(m: re.Match) -> str:
        key = (m.group(1) or m.group(2) or "").strip()
        if not key:
            return m.group(0)
        value = mapping.get(key)
        if value is None:
            value = mapping.get(key.strip().lower())
        return value if value is not None else m.group(0)

    return PLACEHOLDER_RE.sub(repl, text)


def _render_docx_template(source_docx: Path, mapping: dict[str, str], output_path: Path) -> Path:
    from docx import Document

    doc = Document(str(source_docx))
    for p in doc.paragraphs:
        for run in p.runs:
            if run.text:
                run.text = _replace_placeholders_in_text(run.text, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.text:
                            run.text = _replace_placeholders_in_text(run.text, mapping)
    doc.save(str(output_path))
    return output_path


def _resolve_soffice_executable() -> str | None:
    try:
        from modules.houtai.api.employees import _resolve_soffice_executable as _resolve
        return _resolve()
    except Exception:
        return None


def _convert_office_file(source: Path, target_ext: str, out_dir: Path) -> Path | None:
    soffice = _resolve_soffice_executable()
    if not soffice:
        return None
    out_dir.mkdir(parents=True, exist_ok=True)
    soffice_dir = str(Path(soffice).resolve().parent)
    import subprocess

    cmd = [
        soffice,
        "--headless",
        "--nologo",
        "--nolockcheck",
        "--norestore",
        "--convert-to",
        target_ext.lstrip("."),
        "--outdir",
        str(out_dir),
        str(source),
    ]
    proc = subprocess.run(cmd, cwd=soffice_dir, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    if proc.returncode != 0:
        return None
    expected = out_dir / f"{source.stem}.{target_ext.lstrip('.')}"
    if expected.exists():
        return expected
    for c in out_dir.glob(f"{source.stem}.*"):
        if c.suffix.lower() == f".{target_ext.lstrip('.').lower()}":
            return c
    return None


def _render_pdf_template(source_pdf: Path, mapping: dict[str, str]) -> bytes:
    from PyPDF2 import PdfReader, PdfWriter

    reader = PdfReader(str(source_pdf))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    try:
        fields = reader.get_fields() or {}
    except Exception:
        fields = {}
    if not fields:
        with open(source_pdf, "rb") as f:
            return f.read()
    values: dict[str, str] = {}
    for field_name in fields.keys():
        raw = (field_name or "").strip()
        if not raw:
            continue
        m = re.fullmatch(r"\$\{([^}]+)\}|\{([^}]+)\}", raw)
        key = (m.group(1) or m.group(2) or "").strip() if m else raw
        v = mapping.get(key)
        if v is None:
            v = mapping.get(key.strip().lower())
        if v is None:
            continue
        values[field_name] = v
    if values:
        for page in writer.pages:
            writer.update_page_form_field_values(page, values)
    buf = tempfile.SpooledTemporaryFile(max_size=8 * 1024 * 1024)
    writer.write(buf)
    buf.seek(0)
    return buf.read()


def _render_template_for_customer(db: Session, customer: Customer, template_id: str) -> tuple[bytes, str, str]:
    row = db.query(TemplateFile).filter(TemplateFile.id == template_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="模板不存在")
    p = get_file_path(row.file_url)
    if not p or not p.exists():
        raise HTTPException(status_code=404, detail="模板文件不存在")
    ext = (row.file_type or p.suffix.lstrip(".")).lower()
    mapping = _customer_placeholder_values(customer)
    display_name = row.file_name or p.name

    if ext == "pdf":
        return _render_pdf_template(p, mapping), "pdf", display_name

    with tempfile.TemporaryDirectory() as td:
        td_path = Path(td)
        source_docx = p
        if ext == "doc":
            converted = _convert_office_file(p, "docx", td_path)
            if not converted:
                with open(p, "rb") as f:
                    return f.read(), ext, display_name
            source_docx = converted

        rendered_docx = td_path / f"rendered_{uuid.uuid4().hex}.docx"
        _render_docx_template(source_docx, mapping, rendered_docx)

        if ext == "docx":
            return rendered_docx.read_bytes(), "docx", display_name

        converted_back = _convert_office_file(rendered_docx, "doc", td_path)
        if converted_back and converted_back.exists():
            return converted_back.read_bytes(), "doc", display_name
        return rendered_docx.read_bytes(), "docx", display_name


def _generate_customer_code(db: Session) -> str:
    while True:
        code = uuid.uuid4().hex[:6].upper()
        existing = db.query(Customer).filter(Customer.customer_code == code).first()
        if not existing:
            return code


def _notify_admins_customer_status_changed(
    db: Session,
    *,
    customer_id: str,
    new_status: str | None,
    trigger_user_id: str | None,
) -> None:
    try:
        admin_users = db.query(User).filter(or_(User.is_active == True, User.is_active.is_(None))).all()
        for u in admin_users:
            if trigger_user_id and str(u.id) == str(trigger_user_id):
                continue
            if str(new_status or "") == CUSTOMER_STATUS_PENDING_ARCHIVE:
                touch_business_unread(
                    db,
                    business_code="customer_pending",
                    receiver_user_id=str(u.id),
                    data_id=str(customer_id),
                    scope_id=str(new_status) if new_status is not None else None,
                    trigger_user_id=str(trigger_user_id) if trigger_user_id else None,
                )
            else:
                touch_business_unread(
                    db,
                    business_code="customer",
                    receiver_user_id=str(u.id),
                    data_id=str(customer_id),
                    scope_id=str(new_status) if new_status is not None else None,
                    trigger_user_id=str(trigger_user_id) if trigger_user_id else None,
                )
        db.commit()
    except Exception:
        db.rollback()


def _load_attachments(raw_value: str):
    if not raw_value:
        return []
    try:
        data = json.loads(raw_value)
        return data if isinstance(data, list) else []
    except Exception:
        return []


def _dump_attachments(items) -> str:
    try:
        return json.dumps(items or [], ensure_ascii=False)
    except Exception:
        return "[]"


def _build_attachment_response(customer: Customer) -> List[CustomerAttachment]:
    raw = _load_attachments(customer.attachments)
    results = []
    for idx, item in enumerate(raw):
        name = item.get("name") if isinstance(item, dict) else None
        path = item.get("path") if isinstance(item, dict) else None
        if not name and isinstance(item, str):
            name = item.split("/")[-1]
        results.append(
            CustomerAttachment(
                name=name or f"附件{idx + 1}",
                url=f"/api/houtai/customers/{customer.id}/attachments/{idx}/download",
                path=path
            )
        )
    return results


def _normalize_attachments_for_save(existing_raw, incoming):
    existing = [item for item in (existing_raw or []) if isinstance(item, dict)]
    results = []
    for item in incoming or []:
        if isinstance(item, dict):
            name = item.get("name")
            path = item.get("path")
            if not path and name:
                matched = next((x for x in existing if x.get("name") == name), None)
                path = matched.get("path") if matched else None
            if name and path:
                results.append({"name": name, "path": path})
        elif isinstance(item, str):
            name = item.split("/")[-1]
            matched = next((x for x in existing if x.get("name") == name), None)
            path = matched.get("path") if matched else None
            if name and path:
                results.append({"name": name, "path": path})
    return results


def _get_user_from_token(db: Session, token: str) -> User:
    payload = decode_access_token(token)
    if payload is None:
        raise HTTPException(status_code=401, detail="无效的认证令牌")
    user_id: str = payload.get("sub")
    if not user_id:
        raise HTTPException(status_code=401, detail="无效的认证令牌")
    user = db.query(User).filter(User.id == user_id).first()
    if user is None or not user.is_active:
        raise HTTPException(status_code=401, detail="用户不存在或已被禁用")
    return user


def _calc_service_stats(db: Session, customer_id: str):
    tasks = db.query(Task).filter(
        Task.customer_id == customer_id,
        Task.status.in_([TaskStatus.completed, TaskStatus.approved])
    ).all()
    count = len(tasks)
    last_time = None
    for task in tasks:
        candidate = (
            task.service_end_time
            or task.completed_at
            or task.service_time
            or task.service_start_time
            or task.created_at
        )
        if candidate and (last_time is None or candidate > last_time):
            last_time = candidate
    return count, last_time


def _build_customer_response(
    customer: Customer,
    service_count: int,
    last_service_time,
    has_update: bool = False,
):
    accepted_ids = [svc.id for svc in getattr(customer, "accepted_service_level1", []) or []]
    accepted_names = [svc.name for svc in getattr(customer, "accepted_service_level1", []) or []]
    return CustomerResponse(
        id=customer.id,
        has_update=bool(has_update),
        customer_code=customer.customer_code,
        customer_status=getattr(customer, "customer_status", None),
        name=customer.name,
        phone=customer.phone,
        address=customer.address,
        email=customer.email,
        introduction=customer.introduction,
        notes=customer.notes,
        customer_type=customer.customer_type,
        gender=customer.gender,
        age=customer.age,
        disability_type=customer.disability_type,
        ndis_number=customer.ndis_number,
        ndis_plan_copy_path=getattr(customer, "ndis_plan_copy_path", None),
        aboriginal_torres_strait=getattr(customer, "aboriginal_torres_strait", None),
        ndis_funding_type=getattr(customer, "ndis_funding_type", None),
        medicare_number=getattr(customer, "medicare_number", None),
        medicare_expiry=getattr(customer, "medicare_expiry", None),
        has_medical_card=getattr(customer, "has_medical_card", None),
        medical_card_number=getattr(customer, "medical_card_number", None),
        private_health_fund=getattr(customer, "private_health_fund", None),
        private_policy_number=getattr(customer, "private_policy_number", None),
        invoice_receiver_name=getattr(customer, "invoice_receiver_name", None),
        invoice_receiver_phone=getattr(customer, "invoice_receiver_phone", None),
        invoice_receiver_email=getattr(customer, "invoice_receiver_email", None),
        invoice_receiver_address=getattr(customer, "invoice_receiver_address", None),
        emergency_contact1_name=getattr(customer, "emergency_contact1_name", None),
        emergency_contact1_phone=getattr(customer, "emergency_contact1_phone", None),
        emergency_contact1_email=getattr(customer, "emergency_contact1_email", None),
        emergency_contact2_name=getattr(customer, "emergency_contact2_name", None),
        emergency_contact2_phone=getattr(customer, "emergency_contact2_phone", None),
        emergency_contact2_email=getattr(customer, "emergency_contact2_email", None),
        weekly_service_hours=getattr(customer, "weekly_service_hours", None),
        weekly_served_hours=float(getattr(customer, "weekly_served_hours", 0) or 0),
        attachments=_build_attachment_response(customer),
        accepted_service_level1_ids=accepted_ids or None,
        accepted_service_level1_names=accepted_names or None,
        service_count=service_count,
        last_service_time=last_service_time,
        created_at=customer.created_at
    )


@router.get("", response_model=List[CustomerResponse])
async def get_customers(
    customer_status: str | None = Query(default=None),
    customer_type: str | None = Query(default=None),
    include_all: bool = Query(default=False),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """获取所有客户列表"""
    unread_customer_ids = {
        (r.data_id or "")
        for r in db.query(BusinessUnread)
        .filter(
            BusinessUnread.receiver_user_id == str(current_user.id),
            BusinessUnread.business_code == "customer",
            BusinessUnread.is_unread == 1,
        )
        .all()
    }
    unread_pending_customer_ids = {
        (r.data_id or "")
        for r in db.query(BusinessUnread)
        .filter(
            BusinessUnread.receiver_user_id == str(current_user.id),
            BusinessUnread.business_code == "customer_pending",
            BusinessUnread.is_unread == 1,
        )
        .all()
    }
    query = db.query(Customer)
    if not include_all and not customer_status:
        customer_status = CUSTOMER_STATUS_ARCHIVED
    if customer_status:
        query = query.filter(getattr(Customer, "customer_status") == customer_status)
    if customer_type:
        query = query.filter(func.trim(getattr(Customer, "customer_type")) == customer_type)
    customers = query.order_by(Customer.created_at.desc()).all()
    results = []
    for customer in customers:
        service_count, last_service_time = _calc_service_stats(db, customer.id)
        results.append(
            _build_customer_response(
                customer,
                service_count,
                last_service_time,
                has_update=(str(customer.id) in unread_customer_ids) or (str(customer.id) in unread_pending_customer_ids),
            )
        )
    return results


@router.post("", response_model=CustomerResponse, status_code=status.HTTP_201_CREATED)
async def create_customer(
    customer_data: CustomerCreate,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """创建客户"""
    payload = customer_data.dict()
    attachments = payload.pop("attachments", None)
    accepted_ids = payload.pop("accepted_service_level1_ids", None)
    customer = Customer(**payload)
    if getattr(customer, "customer_status", None) is None:
        customer.customer_status = CUSTOMER_STATUS_UNARCHIVED
    if not customer.customer_code:
        customer.customer_code = _generate_customer_code(db)
    if attachments is not None:
        customer.attachments = _dump_attachments(attachments)
    db.add(customer)
    if accepted_ids:
        services = db.query(InvoiceServiceLevel1).filter(InvoiceServiceLevel1.id.in_(accepted_ids)).all()
        if len(services) != len(set(accepted_ids)):
            raise HTTPException(status_code=400, detail="存在无效的一级服务ID")
        customer.accepted_service_level1 = services
    db.commit()
    db.refresh(customer)
    service_count, last_service_time = _calc_service_stats(db, customer.id)
    return _build_customer_response(customer, service_count, last_service_time)


@router.post("/{customer_id}/archive/approve", response_model=CustomerResponse)
async def approve_customer_archive(
    customer_id: str,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")
    old_status = getattr(customer, "customer_status", None)

    if getattr(customer, "customer_status", None) == CUSTOMER_STATUS_ARCHIVED:
        service_count, last_service_time = _calc_service_stats(db, customer.id)
        return _build_customer_response(customer, service_count, last_service_time)

    contract = db.query(CustomerDocument).filter(
        CustomerDocument.customer_id == customer_id,
        CustomerDocument.document_type == "service_agreement",
        CustomerDocument.status == STATUS_SIGNED,
    ).first()
    if not contract:
        raise HTTPException(status_code=400, detail="客户合同未签署，无法审核归档")

    customer.customer_status = CUSTOMER_STATUS_ARCHIVED
    db.commit()
    db.refresh(customer)
    new_status = getattr(customer, "customer_status", None)
    if str(old_status or "") != str(new_status or ""):
        _notify_admins_customer_status_changed(
            db,
            customer_id=str(customer.id),
            new_status=new_status,
            trigger_user_id=str(current_user.id) if getattr(current_user, "id", None) else None,
        )
    service_count, last_service_time = _calc_service_stats(db, customer.id)
    return _build_customer_response(customer, service_count, last_service_time)


@router.post("/{customer_id}/archive/start-review", response_model=CustomerResponse)
async def start_customer_archive_review(
    customer_id: str,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")

    old_status = getattr(customer, "customer_status", None)
    if old_status == CUSTOMER_STATUS_ARCHIVED:
        raise HTTPException(status_code=400, detail="客户已建档，无需开始审核")
    if old_status == CUSTOMER_STATUS_PENDING_ARCHIVE:
        service_count, last_service_time = _calc_service_stats(db, customer.id)
        return _build_customer_response(customer, service_count, last_service_time)
    if old_status != CUSTOMER_STATUS_UNARCHIVED:
        raise HTTPException(status_code=400, detail="仅未建档客户可开始审核")

    doc = db.query(CustomerDocument).filter(
        CustomerDocument.customer_id == customer_id,
        CustomerDocument.document_type == "service_agreement",
    ).first()
    if not doc or not doc.file_url:
        raise HTTPException(status_code=400, detail="请先上传合同文件")

    # 线下已签文件直接进入审核：复用当前合同文件作为已签文件。
    doc.status = STATUS_SIGNED
    doc.signed_file_url = doc.signed_file_url or doc.file_url
    doc.signed_at = doc.signed_at or datetime.utcnow()

    from shared.models import DocumentSignRequest
    db.query(DocumentSignRequest).filter(
        DocumentSignRequest.customer_id == customer_id,
        DocumentSignRequest.document_id == doc.id,
        DocumentSignRequest.status == "pending",
    ).update({"status": "expired"})

    customer.customer_status = CUSTOMER_STATUS_PENDING_ARCHIVE
    db.commit()
    db.refresh(customer)

    new_status = getattr(customer, "customer_status", None)
    if str(old_status or "") != str(new_status or ""):
        _notify_admins_customer_status_changed(
            db,
            customer_id=str(customer.id),
            new_status=new_status,
            trigger_user_id=str(current_user.id) if getattr(current_user, "id", None) else None,
        )
    service_count, last_service_time = _calc_service_stats(db, customer.id)
    return _build_customer_response(customer, service_count, last_service_time)


def _get_or_create_contract_doc(db: Session, customer_id: str) -> CustomerDocument:
    doc = db.query(CustomerDocument).filter(
        CustomerDocument.customer_id == customer_id,
        CustomerDocument.document_type == "service_agreement",
    ).first()
    if not doc:
        doc = CustomerDocument(
            customer_id=customer_id,
            document_type="service_agreement",
            name="Service Agreement",
            status="draft",
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)
    return doc

@router.post("/{customer_id}/contract/upload", response_model=CustomerDocumentResponse)
async def upload_customer_contract_pdf(
    customer_id: str,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user),
):
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")
    if getattr(customer, "customer_status", None) != CUSTOMER_STATUS_UNARCHIVED:
        raise HTTPException(status_code=400, detail="仅未建档客户可上传合同")
    ext = (file.filename or "").rsplit(".", 1)[-1].lower() if "." in (file.filename or "") else "pdf"
    if ext != "pdf":
        raise HTTPException(status_code=400, detail="请上传 PDF 格式的合同")
    content = await file.read()
    filename = f"contract_{uuid.uuid4().hex[:8]}_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}.pdf"
    file_url = await save_upload_file(content, filename, subfolder=f"customers/{customer_id}/documents")

    doc = _get_or_create_contract_doc(db, customer_id)
    doc.file_url = file_url
    doc.file_type = "pdf"
    doc.status = "draft"
    db.commit()
    db.refresh(doc)
    return CustomerDocumentResponse(
        id=doc.id,
        customer_id=doc.customer_id,
        document_type=doc.document_type,
        name=doc.name,
        file_type=doc.file_type,
        file_url=doc.file_url,
        form_data=None,
        status=doc.status,
        signed_at=doc.signed_at,
        signed_file_url=doc.signed_file_url,
        created_at=doc.created_at,
        updated_at=doc.updated_at,
    )

@router.post("/{customer_id}/contract/send")
async def send_customer_contract_sign_link(
    customer_id: str,
    request: Request,
    language: str = Query("zh", description="邮件语言：zh / en"),
    payload: dict | None = Body(default=None),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user),
):
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")
    if not customer.email:
        raise HTTPException(status_code=400, detail="客户邮箱不存在，无法发送签署链接")
    if getattr(customer, "customer_status", None) != CUSTOMER_STATUS_UNARCHIVED:
        raise HTTPException(status_code=400, detail="仅未建档客户可发送合同")

    doc = _get_or_create_contract_doc(db, customer_id)
    if not doc.file_url:
        raise HTTPException(status_code=400, detail="请先上传合同 PDF")

    from shared.models import DocumentSignRequest
    db.query(DocumentSignRequest).filter(
        DocumentSignRequest.customer_id == customer_id,
        DocumentSignRequest.document_id == doc.id,
        DocumentSignRequest.status == "pending",
    ).update({"status": "expired"})

    token = str(uuid.uuid4())
    expires_at = datetime.utcnow() + timedelta(days=7)
    req = DocumentSignRequest(
        token=token,
        customer_id=customer_id,
        document_id=doc.id,
        status="pending",
        expires_at=expires_at,
    )
    doc.status = "pending_sign"
    db.add(req)
    db.commit()

    if settings.sign_frontend_base_url:
        raw_base_url = settings.sign_frontend_base_url.strip()
        base_url = raw_base_url.split("#", 1)[0].rstrip("/")
        sign_path = f"/sign/document/{token}" if base_url.lower().endswith("/admin") else f"/admin/sign/document/{token}"
        sign_url = f"{base_url}{sign_path}"
    else:
        raw_origin = (request.headers.get("origin") or str(request.base_url)).strip()
        origin = raw_origin.split("#", 1)[0].rstrip("/")
        sign_path = f"/admin/sign/document/{token}"
        sign_url = f"{origin}{sign_path}"

    def _normalize_lang(value: str | None, accept: str | None) -> str:
        raw = (value or "").strip().lower()
        if not raw and accept:
            raw = str(accept).split(",", 1)[0].strip().lower()
        if raw in ("zh", "zh-cn", "zh_hans", "cn", "chinese", "中文", "简体中文"):
            return "zh"
        if raw in ("en", "en-us", "en-gb", "english", "英文"):
            return "en"
        if raw.startswith("zh"):
            return "zh"
        if raw.startswith("en"):
            return "en"
        return "zh"

    body_lang = None
    if isinstance(payload, dict):
        body_lang = payload.get("language") or payload.get("lang")
    lang = _normalize_lang(language or body_lang, request.headers.get("accept-language"))

    sep = "&" if "?" in sign_url else "?"
    sign_url_with_lang = f"{sign_url}{sep}lang={lang}"

    if lang == "en":
        subject = "Service Agreement Signing Link"
        plain_body = (
            f"Dear {customer.name},\n\n"
            f"Please click the link below to sign your service agreement (valid until: {expires_at.isoformat()}):\n"
            f"{sign_url_with_lang}\n\n"
            "If the link has expired, please contact the administrator to resend it.\n"
        )
        html_body = f"""
<p>Dear {customer.name},</p>
<p>Please click the link below to sign your service agreement (valid until: {expires_at.isoformat()}):</p>
<p><a href="{sign_url_with_lang}" target="_blank" rel="noopener noreferrer">{sign_url_with_lang}</a></p>
<p>If the link has expired, please contact the administrator to resend it.</p>
""".strip()
    else:
        subject = "客户合同签署链接"
        plain_body = (
            f"{customer.name}，您好：\n\n"
            f"请点击以下链接完成合同签署（有效期至：{expires_at.isoformat()}）：\n"
            f"{sign_url_with_lang}\n\n"
            "如链接已过期，请联系管理员重新发送。"
        )
        html_body = f"""
<p>{customer.name}，您好：</p>
<p>请点击以下链接完成合同签署（有效期至：{expires_at.isoformat()}）：</p>
<p><a href="{sign_url_with_lang}" target="_blank" rel="noopener noreferrer">{sign_url_with_lang}</a></p>
<p>如链接已过期，请联系管理员重新发送。</p>
""".strip()

    try:
        send_contact_email(subject=subject, html_body=html_body, plain_body=plain_body, to_emails=[customer.email])
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"发送签署链接邮件失败: {str(e)}")

    return {
        "token": token,
        "sign_path": sign_path,
        "sign_url": sign_url_with_lang,
        "expires_at": expires_at.isoformat(),
        "document_id": doc.id,
    }

@router.post("/{customer_id}/documents/upload", response_model=CustomerDocumentResponse)
async def upload_customer_document_like_employee(
    customer_id: str,
    document_type: str = Form(...),
    template_id: str | None = Form(default=None),
    file: UploadFile | None = File(default=None),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """上传客户文档（签约合同），接口风格与员工文档上传一致"""
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")
    if getattr(customer, "customer_status", None) != CUSTOMER_STATUS_UNARCHIVED:
        raise HTTPException(status_code=400, detail="仅未建档客户可上传文件")
    valid_types = ["service_agreement"]
    if document_type not in valid_types:
        raise HTTPException(status_code=400, detail=f"无效的文档类型，必须是: {', '.join(valid_types)}")

    if template_id:
        file_content, file_type, display_name = _render_template_for_customer(db, customer, template_id)
        output_filename = f"{uuid.uuid4().hex}.{file_type}"
        file_url = await save_upload_file(file_content, output_filename, subfolder=f"customers/{customer_id}/documents")
        file_ext = file_type
        original_filename = display_name
    else:
        if file is None:
            raise HTTPException(status_code=400, detail="请选择模板或上传文件")
        file_ext = (file.filename or "").rsplit(".", 1)[-1] if "." in (file.filename or "") else "pdf"
        if file_ext.lower() not in ("doc", "docx", "pdf"):
            raise HTTPException(status_code=400, detail="仅支持 doc/docx/pdf")
        output_filename = f"{uuid.uuid4().hex}.{file_ext}"
        file_content = await file.read()
        file_url = await save_upload_file(file_content, output_filename, subfolder=f"customers/{customer_id}/documents")
        original_filename = file.filename or output_filename
        file_type = file_ext.lower() or "unknown"

    # 创建或更新合同文档记录
    doc = _get_or_create_contract_doc(db, customer_id)
    doc.name = original_filename or doc.name
    doc.file_type = file_type
    doc.file_url = file_url
    doc.status = "draft"
    db.commit()
    db.refresh(doc)

    return CustomerDocumentResponse(
        id=doc.id,
        customer_id=doc.customer_id,
        document_type=doc.document_type,
        name=doc.name,
        file_type=doc.file_type,
        file_url=doc.file_url,
        form_data=None,
        status=doc.status,
        signed_at=doc.signed_at,
        signed_file_url=doc.signed_file_url,
        created_at=doc.created_at,
        updated_at=doc.updated_at,
    )

@router.get("/{customer_id}/contract/view")
async def view_signed_customer_contract(
    customer_id: str,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user),
):
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")
    if getattr(customer, "customer_status", None) != "待建档":
        raise HTTPException(status_code=400, detail="仅待建档客户可查看已签署合同")
    doc = db.query(CustomerDocument).filter(
        CustomerDocument.customer_id == customer_id,
        CustomerDocument.document_type == "service_agreement",
    ).first()
    if not doc or not doc.signed_file_url:
        raise HTTPException(status_code=404, detail="未找到已签署的合同")
    file_path = get_file_path(doc.signed_file_url)
    if not file_path or not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在")
    return FileResponse(str(file_path), filename=file_path.name, media_type="application/pdf")

@router.post("/{customer_id}/archive/reject", response_model=CustomerResponse)
async def reject_customer_archive(
    customer_id: str,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user),
):
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")
    old_status = getattr(customer, "customer_status", None)
    if getattr(customer, "customer_status", None) != "待建档":
        raise HTTPException(status_code=400, detail="仅待建档客户可拒绝")

    doc = db.query(CustomerDocument).filter(
        CustomerDocument.customer_id == customer_id,
        CustomerDocument.document_type == "service_agreement",
    ).first()
    if doc:
        doc.status = "draft"
        doc.signed_file_url = None
        doc.signed_at = None
        from shared.models import DocumentSignRequest
        db.query(DocumentSignRequest).filter(
            DocumentSignRequest.customer_id == customer_id,
            DocumentSignRequest.document_id == doc.id,
            DocumentSignRequest.status == "pending",
        ).update({"status": "expired"})

    customer.customer_status = CUSTOMER_STATUS_UNARCHIVED
    db.commit()
    db.refresh(customer)
    new_status = getattr(customer, "customer_status", None)
    if str(old_status or "") != str(new_status or ""):
        _notify_admins_customer_status_changed(
            db,
            customer_id=str(customer.id),
            new_status=new_status,
            trigger_user_id=str(current_user.id) if getattr(current_user, "id", None) else None,
        )
    service_count, last_service_time = _calc_service_stats(db, customer.id)
    return _build_customer_response(customer, service_count, last_service_time)

@router.delete("/{customer_id}/contract", status_code=status.HTTP_204_NO_CONTENT)
async def delete_customer_contract(
    customer_id: str,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user),
):
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")
    doc = db.query(CustomerDocument).filter(
        CustomerDocument.customer_id == customer_id,
        CustomerDocument.document_type == "service_agreement",
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="客户合同不存在")
    for path_value in [doc.signed_file_url, doc.file_url]:
        p = get_file_path(path_value) if path_value else None
        if p and p.exists():
            try:
                p.unlink(missing_ok=True)
            except Exception:
                pass
    from shared.models import DocumentSignRequest
    db.query(DocumentSignRequest).filter(
        DocumentSignRequest.customer_id == customer_id,
        DocumentSignRequest.document_id == doc.id,
    ).delete()
    db.delete(doc)
    db.commit()
    return None

@router.get("/{customer_id}", response_model=CustomerResponse)
async def get_customer(
    customer_id: str,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """获取客户详情"""
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")
    service_count, last_service_time = _calc_service_stats(db, customer.id)
    return _build_customer_response(customer, service_count, last_service_time)


@router.put("/{customer_id}", response_model=CustomerResponse)
async def update_customer(
    customer_id: str,
    customer_data: CustomerUpdate,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """更新客户信息"""
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")
    old_status = getattr(customer, "customer_status", None)
    
    update_data = customer_data.dict(exclude_unset=True)
    attachments = update_data.pop("attachments", None)
    accepted_ids = update_data.pop("accepted_service_level1_ids", None)
    for key, value in update_data.items():
        setattr(customer, key, value)
    if attachments is not None:
        existing = _load_attachments(customer.attachments)
        normalized = _normalize_attachments_for_save(existing, attachments)
        customer.attachments = _dump_attachments(normalized)
    if accepted_ids is not None:
        services = db.query(InvoiceServiceLevel1).filter(InvoiceServiceLevel1.id.in_(accepted_ids)).all()
        if len(services) != len(set(accepted_ids)):
            raise HTTPException(status_code=400, detail="存在无效的一级服务ID")
        customer.accepted_service_level1 = services
    
    db.commit()
    db.refresh(customer)
    new_status = getattr(customer, "customer_status", None)
    if str(old_status or "") != str(new_status or ""):
        _notify_admins_customer_status_changed(
            db,
            customer_id=str(customer.id),
            new_status=new_status,
            trigger_user_id=str(current_user.id) if getattr(current_user, "id", None) else None,
        )
    service_count, last_service_time = _calc_service_stats(db, customer.id)
    return _build_customer_response(customer, service_count, last_service_time)


@router.delete("/{customer_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_customer(
    customer_id: str,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """删除客户"""
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")
    
    db.delete(customer)
    db.commit()
    return None


@router.post("/{customer_id}/attachments", response_model=CustomerResponse)
async def upload_customer_attachment(
    customer_id: str,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """上传客户附件"""
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")

    file_content = await file.read()
    filename = f"{uuid.uuid4().hex}_{file.filename}"
    file_url = await save_upload_file(file_content, filename, subfolder=f"customers/{customer_id}")

    attachments = _load_attachments(customer.attachments)
    attachments.append({"name": file.filename, "path": file_url})
    customer.attachments = _dump_attachments(attachments)
    db.commit()
    db.refresh(customer)

    service_count, last_service_time = _calc_service_stats(db, customer.id)
    return _build_customer_response(customer, service_count, last_service_time)


@router.post("/{customer_id}/ndis-plan", response_model=CustomerResponse)
async def upload_ndis_plan(
    customer_id: str,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """上传 NDIS 计划文件（客户类型为 NDIS 时使用）"""
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")
    if customer.customer_type != "NDIS":
        raise HTTPException(status_code=400, detail="仅 NDIS 类型客户可上传 NDIS 计划文件")

    file_content = await file.read()
    ext = (file.filename or "").rsplit(".", 1)[-1] if "." in (file.filename or "") else "pdf"
    if ext.lower() not in ("pdf", "doc", "docx"):
        ext = "pdf"
    filename = f"ndis_plan_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.{ext}"
    file_url = await save_upload_file(file_content, filename, subfolder=f"customers/{customer_id}")

    customer.ndis_plan_copy_path = file_url
    db.commit()
    db.refresh(customer)

    service_count, last_service_time = _calc_service_stats(db, customer.id)
    return _build_customer_response(customer, service_count, last_service_time)


@router.get("/{customer_id}/ndis-plan/download")
async def download_ndis_plan(
    customer_id: str,
    token: str | None = Query(default=None),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_optional),
    db: Session = Depends(get_db)
):
    """下载 NDIS 计划文件"""
    auth_token = token or (credentials.credentials if credentials else None)
    if not auth_token:
        raise HTTPException(status_code=401, detail="无效的认证令牌")
    _get_user_from_token(db, auth_token)
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")
    path_value = getattr(customer, "ndis_plan_copy_path", None)
    if not path_value:
        raise HTTPException(status_code=404, detail="NDIS 计划文件未上传")
    file_path = get_file_path(path_value)
    if not file_path or not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在")
    return FileResponse(
        str(file_path),
        filename=f"ndis_plan_{customer.name or customer_id}.pdf",
        media_type="application/octet-stream"
    )


@router.get("/{customer_id}/attachments/{attachment_index}/download")
async def download_customer_attachment(
    customer_id: str,
    attachment_index: int,
    token: str | None = Query(default=None),
    credentials: HTTPAuthorizationCredentials | None = Depends(security_optional),
    db: Session = Depends(get_db)
):
    """下载客户附件"""
    auth_token = token or (credentials.credentials if credentials else None)
    if not auth_token:
        raise HTTPException(status_code=401, detail="无效的认证令牌")
    _get_user_from_token(db, auth_token)
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="客户不存在")
    attachments = _load_attachments(customer.attachments)
    if attachment_index < 0 or attachment_index >= len(attachments):
        raise HTTPException(status_code=404, detail="附件不存在")
    item = attachments[attachment_index]
    path_value = item.get("path") if isinstance(item, dict) else None
    if not path_value:
        raise HTTPException(status_code=404, detail="附件不存在")
    file_path = get_file_path(path_value)
    if not file_path:
        raise HTTPException(status_code=404, detail="附件不存在")
    filename = item.get("name") if isinstance(item, dict) else file_path.name
    return FileResponse(
        str(file_path),
        filename=filename or file_path.name,
        media_type="application/octet-stream"
    )
