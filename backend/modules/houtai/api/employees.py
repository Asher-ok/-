from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form, Query, Body, Request
from fastapi.responses import FileResponse, Response
from sqlalchemy.orm import Session
from sqlalchemy import or_, func
from sqlalchemy.exc import IntegrityError, SQLAlchemyError
from typing import List, Optional
from pathlib import Path
from core.database import get_db
from shared.models import Employee, Qualification, TrainingRecord, SystemSetting, EmployeeDocument, EmployeeContractSignRequest, BusinessUnread, Task, TaskServiceItem
from shared.models.update_notification import touch_business_unread
from ..schemas.employee import EmployeeCreate, EmployeeUpdate, EmployeeAccountStatusUpdate, EmployeeResponse, QualificationCreate, TrainingRecordCreate, TrainingRecordUpdate, TrainingRecordResponse, ExpiringTrainingRecordItem, ReminderSettingPayload, EmployeeDocumentResponse, ContractGenerateRequest, TrainingStatus
from ..dependencies import get_current_user
from core.auth import get_password_hash
from core.utils.file_utils import save_upload_file, get_file_path, build_content_disposition, to_ascii_filename
from core.utils.email import send_contact_email
from core.config import settings
from datetime import datetime, timedelta, timezone
from dateutil.relativedelta import relativedelta
import os
import uuid
import mimetypes
import logging
import base64
import re
import tempfile
import subprocess
import shutil
import time

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/houtai/employees", tags=["管理-员工"])


def _normalize_training_status(raw: Optional[str]) -> Optional[str]:
    if raw is None:
        return None
    value = str(raw).strip()
    if not value:
        return None

    direct = {
        "in_progress": TrainingStatus.in_progress.value,
        "completed": TrainingStatus.completed.value,
        "rejected": TrainingStatus.rejected.value,
        "pending": TrainingStatus.pending.value,
    }
    lowered = value.lower()
    if lowered in direct:
        return direct[lowered]

    aliases = {
        "已通过": TrainingStatus.completed.value,
        "通过": TrainingStatus.completed.value,
        "approved": TrainingStatus.completed.value,
        "审核通过": TrainingStatus.completed.value,
        "已完成": TrainingStatus.completed.value,
        "未通过": TrainingStatus.rejected.value,
        "驳回": TrainingStatus.rejected.value,
        "审核驳回": TrainingStatus.rejected.value,
        "待审核": TrainingStatus.pending.value,
        "审核中": TrainingStatus.pending.value,
        "进行中": TrainingStatus.in_progress.value,
    }
    if value in aliases:
        return aliases[value]

    raise HTTPException(status_code=400, detail=f"无效的培训状态: {value}")


def _parse_data_url(data_url: str):
    """解析base64 data URL"""
    match = re.match(r"^data:(.+?);base64,(.+)$", data_url)
    if not match:
        return None, None
    mime_type = match.group(1)
    data = base64.b64decode(match.group(2))
    return mime_type, data


def _embed_signature_to_contract(
    contract_path: Path,
    signature_blob: bytes,
    signature_type: str = "admin",
    x: float = None,
    y: float = None,
    width: float = None,
    height: float = None,
    page: int = 0,
):
    """将签名嵌入到合同文档中（支持PDF、图片）
    
    注意：docx文件应该先转换为PDF再调用此函数，不要直接对docx文件调用此函数
    
    Args:
        contract_path: 合同文件路径
        signature_blob: 签名图片二进制数据
        signature_type: 签名类型，'employee' 或 'admin'
        x, y, width, height: 签字坐标和尺寸（像素或点）
    """
    file_ext = contract_path.suffix.lower()
    
    if file_ext == '.pdf':
        return _embed_signature_to_pdf(
            contract_path,
            signature_blob,
            signature_type,
            x,
            y,
            width,
            height,
            page,
        )
    elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
        return _embed_signature_to_image(contract_path, signature_blob, signature_type, x, y, width, height)
    else:
        # 不支持的文件类型
        logger.error(f"不支持的文件类型: {file_ext}，请先将docx转换为PDF")
        return False


def _convert_docx_to_pdf_in_place(source_path: Path) -> Path:
    """将doc/docx转换为PDF并放到原目录，返回新PDF路径（增强版）"""
    soffice = _resolve_soffice_executable()
    if not soffice:
        raise RuntimeError("未找到 LibreOffice（soffice）。请安装 LibreOffice，或将 soffice 加入 PATH（Windows 可检查 LibreOffice\\program\\soffice.com）。")

    # 强制使用绝对路径，避免相对路径问题
    source_path = source_path.resolve()
    if not source_path.exists():
        raise FileNotFoundError(f"源文件不存在: {source_path}")
    
    soffice_dir = str(Path(soffice).resolve().parent)
    final_output_dir = source_path.parent
    job_dir = Path(tempfile.mkdtemp(prefix="empowerhub_lo_job_"))
    output_dir = job_dir / "out"
    output_dir.mkdir(parents=True, exist_ok=True)
    tmp_dir = job_dir / "tmp"
    tmp_dir.mkdir(parents=True, exist_ok=True)

    def run_convert(isolated_profile: bool):
        env = os.environ.copy()
        env["TMP"] = str(tmp_dir)
        env["TEMP"] = str(tmp_dir)
        env["TMPDIR"] = str(tmp_dir)

        command = [
            soffice,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--norestore",
            "--nolockcheck",  # 新增：避免锁文件导致失败
            "--nodefault",    # 新增：不加载默认配置
        ]
        if isolated_profile:
            profile_dir = job_dir / "profile"
            profile_dir.mkdir(parents=True, exist_ok=True)
            command.append(f"-env:UserInstallation={profile_dir.as_uri()}")

        command.extend(
            [
                "--convert-to",
                "pdf:writer_pdf_Export",
                "--outdir",
                str(output_dir),
                str(source_path),
            ]
        )
        return subprocess.run(
            command,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            env=env,
            cwd=soffice_dir,
            timeout=60,  # 新增：超时保护（60秒）
        )

    try:
        # 第一次尝试（非隔离模式）
        result = run_convert(isolated_profile=False)
        if result.returncode != 0:
            # 失败后重试隔离模式
            result = run_convert(isolated_profile=True)
        if result.returncode != 0:
            stderr = (result.stderr or "").strip()
            stdout = (result.stdout or "").strip()
            raise RuntimeError(
                f"docx转PDF失败，返回码：{result.returncode}\nSTDOUT: {stdout}\nSTDERR: {stderr}"
            )

        # ========== 核心优化：增强的PDF文件查找逻辑 ==========
        # 1. 先等待文件生成（最多10秒，轮询检查）
        start_ts = time.time()
        deadline = start_ts + 10
        generated_pdf = None
        
        while time.time() < deadline and not generated_pdf:
            # 遍历输出目录下的所有PDF文件（兼容大小写扩展名）
            pdf_files = [p for p in output_dir.iterdir() if p.is_file() and p.suffix.lower() == ".pdf"]
            if not pdf_files:
                time.sleep(0.5)
                continue
            
            # 2. 匹配规则：优先按文件名（忽略大小写）
            source_stem = source_path.stem.lower()
            candidate_files = []
            
            for p in pdf_files:
                # 检查文件修改时间（只取转换后生成的）
                if p.stat().st_mtime >= start_ts - 2:
                    candidate_files.append(p)
                    # 完全匹配文件名（忽略大小写）
                    if p.stem.lower() == source_stem:
                        generated_pdf = p
                        break
            
            # 3. 无完全匹配时，取最新生成的PDF
            if not generated_pdf and candidate_files:
                generated_pdf = max(candidate_files, key=lambda x: x.stat().st_mtime)
            
            time.sleep(0.5)
        
        if not generated_pdf or not generated_pdf.exists():
            # 终极兜底：直接按源文件名拼接（兼容部分环境的命名规则）
            fallback_pdf = output_dir / f"{source_path.stem}.pdf"
            if fallback_pdf.exists():
                generated_pdf = fallback_pdf
            else:
                raise RuntimeError(
                    f"未找到转换后的PDF文件！\n输出目录文件列表：{os.listdir(output_dir)}"
                )

        # ========== 复制/移动到最终目录 ==========
        # 生成唯一文件名（避免冲突）
        final_pdf_name = f"{source_path.stem}_{uuid.uuid4().hex}.pdf"
        final_pdf = final_output_dir / final_pdf_name
        
        # 优先移动，失败则复制
        try:
            generated_pdf.replace(final_pdf)
        except Exception as e:
            logger.warning(f"移动PDF文件失败，降级为复制：{e}")
            shutil.copy2(generated_pdf, final_pdf)  # copy2 保留元数据
        
        # 验证最终文件
        if not final_pdf.exists() or final_pdf.stat().st_size == 0:
            raise RuntimeError(f"生成的PDF文件为空或不存在：{final_pdf}")
        
        return final_pdf

    finally:
        # 清理临时目录（忽略错误）
        shutil.rmtree(job_dir, ignore_errors=True)

def _resolve_soffice_executable() -> Optional[str]:
    configured_path = (
        getattr(settings, "soffice_path", None)
        or getattr(settings, "SOFFICE_PATH", None)
        or os.environ.get("SOFFICE_PATH")
    )
    if configured_path:
        candidate = Path(configured_path)
        # 在 Windows 下优先使用 soffice.com
        if candidate.name.lower() == "soffice.exe":
            sibling = candidate.with_name("soffice.com")
            if sibling.exists():
                candidate = sibling
        if candidate.exists() and candidate.is_file():
            return str(candidate)

    configured_dir = (
        getattr(settings, "libreoffice_program_dir", None)
        or getattr(settings, "LIBREOFFICE_PROGRAM_DIR", None)
        or os.environ.get("LIBREOFFICE_PROGRAM_DIR")
    )
    if configured_dir:
        base = Path(configured_dir)
        for exe in ("soffice.com", "soffice.exe"):
            candidate = base / exe
            if candidate.exists() and candidate.is_file():
                return str(candidate)

    for name in ("soffice", "soffice.com", "soffice.exe"):
        resolved = shutil.which(name)
        if resolved:
            return resolved

    if os.name == "nt":
        program_files = [
            os.environ.get("PROGRAMFILES"),
            os.environ.get("PROGRAMFILES(X86)"),
            "C:\\Program Files",
            "C:\\Program Files (x86)",
        ]
        for base in [p for p in program_files if p]:
            for exe in ("soffice.com", "soffice.exe"):
                candidate = Path(base) / "LibreOffice" / "program" / exe
                if candidate.exists():
                    return str(candidate)

    return None


def _embed_signature_to_word(contract_path: Path, signature_blob: bytes, signature_type: str = "admin", x: float = None, y: float = None, width: float = None, height: float = None):
    """将签名嵌入到Word文档的指定坐标位置（不支持绝对坐标时回退到占位符）"""
    try:
        from docx import Document
        from docx.shared import Mm, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from PIL import Image
        import io
        
        doc = Document(str(contract_path))
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
            tmp_file.write(signature_blob)
            tmp_signature_path = tmp_file.name
        
        try:
            if x is not None and y is not None and width is not None and height is not None:
                section = doc.sections[0]
                page_width = section.page_width
                page_height = section.page_height

                is_ratio = 0 <= x <= 1 and 0 <= y <= 1 and 0 <= width <= 1 and 0 <= height <= 1
                if is_ratio:
                    box_w_emu = float(page_width) * float(width)
                    box_h_emu = float(page_height) * float(height)
                    box_x_emu = float(page_width) * float(x)
                    box_y_emu = float(page_height) * float(y)
                else:
                    px_to_emu = 914400.0 / 96.0
                    box_w_emu = float(width) * px_to_emu
                    box_h_emu = float(height) * px_to_emu
                    box_x_emu = float(x) * px_to_emu
                    box_y_emu = float(y) * px_to_emu

                img = Image.open(io.BytesIO(signature_blob)).convert("RGBA")
                w, h = img.size
                px = img.load()
                min_x, min_y = w, h
                max_x, max_y = -1, -1
                for yy in range(h):
                    for xx in range(w):
                        r, g, b, a = px[xx, yy]
                        if a <= 10:
                            continue
                        if r >= 245 and g >= 245 and b >= 245:
                            continue
                        if xx < min_x:
                            min_x = xx
                        if yy < min_y:
                            min_y = yy
                        if xx > max_x:
                            max_x = xx
                        if yy > max_y:
                            max_y = yy
                if max_x >= 0 and max_y >= 0:
                    left = max(min_x, 0)
                    top = max(min_y, 0)
                    right = min(max_x + 1, w)
                    bottom = min(max_y + 1, h)
                    img = img.crop((left, top, right, bottom))
                img.save(tmp_signature_path, format="PNG")

                px_to_emu = 914400.0 / 96.0
                img_w_emu = float(img.size[0]) * px_to_emu
                img_h_emu = float(img.size[1]) * px_to_emu

                pad_emu = 2.0 * px_to_emu
                inner_w_emu = max(box_w_emu - pad_emu * 2, 1.0)
                inner_h_emu = max(box_h_emu - pad_emu * 2, 1.0)
                scale = min(inner_w_emu / max(img_w_emu, 1.0), inner_h_emu / max(img_h_emu, 1.0))
                draw_w_emu = max(img_w_emu * scale, 1.0)
                draw_h_emu = max(img_h_emu * scale, 1.0)

                pos_x_emu = box_x_emu + pad_emu
                pos_y_emu = box_y_emu + box_h_emu - pad_emu - draw_h_emu
                max_x_emu = box_x_emu + box_w_emu - pad_emu - draw_w_emu
                min_y_emu = box_y_emu + pad_emu
                if pos_x_emu > max_x_emu:
                    pos_x_emu = max_x_emu
                if pos_y_emu < min_y_emu:
                    pos_y_emu = min_y_emu
                width_emu = draw_w_emu
                height_emu = draw_h_emu

                paragraph = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(tmp_signature_path, width=Emu(int(width_emu)), height=Emu(int(height_emu)))

                drawing = run._r.xpath('w:drawing')[0]
                inline = drawing.xpath('wp:inline')[0]

                anchor = OxmlElement('wp:anchor')
                anchor.set(qn('wp:simplePos'), '0')
                anchor.set(qn('wp:relativeHeight'), '251659264')
                anchor.set(qn('wp:behindDoc'), '0')
                anchor.set(qn('wp:locked'), '0')
                anchor.set(qn('wp:layoutInCell'), '1')
                anchor.set(qn('wp:allowOverlap'), '1')

                simple_pos = OxmlElement('wp:simplePos')
                simple_pos.set('x', '0')
                simple_pos.set('y', '0')
                anchor.append(simple_pos)

                position_h = OxmlElement('wp:positionH')
                position_h.set(qn('wp:relativeFrom'), 'page')
                pos_offset_h = OxmlElement('wp:posOffset')
                pos_offset_h.text = str(int(pos_x_emu))
                position_h.append(pos_offset_h)
                anchor.append(position_h)

                position_v = OxmlElement('wp:positionV')
                position_v.set(qn('wp:relativeFrom'), 'page')
                pos_offset_v = OxmlElement('wp:posOffset')
                pos_offset_v.text = str(int(pos_y_emu))
                position_v.append(pos_offset_v)
                anchor.append(position_v)

                wrap_none = OxmlElement('wp:wrapNone')
                anchor.append(wrap_none)

                for child in list(inline):
                    anchor.append(child)

                drawing.remove(inline)
                drawing.append(anchor)

                doc.save(str(contract_path))
                return True
            else:
                return _embed_signature_to_word_legacy(contract_path, signature_blob, signature_type)
        finally:
            if os.path.exists(tmp_signature_path):
                os.unlink(tmp_signature_path)
    except Exception as e:
        logger.error(f"嵌入签名到Word文档失败: {e}")
        raise


def _embed_signature_to_word_legacy(contract_path: Path, signature_blob: bytes, signature_type: str = "admin"):
    """将签名嵌入到合同Word文档中（使用占位符方式）"""
    try:
        from docx import Document
        from docx.shared import Mm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document(str(contract_path))
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
            tmp_file.write(signature_blob)
            tmp_signature_path = tmp_file.name
        
        try:
            placeholder_text = f"[{signature_type.capitalize()} Signature]"
            found_placeholder = False
            
            for paragraph in doc.paragraphs:
                if placeholder_text in paragraph.text:
                    paragraph.clear()
                    run = paragraph.add_run()
                    run.add_picture(tmp_signature_path, width=Mm(100))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    found_placeholder = True
                    break
            
            if not found_placeholder:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if placeholder_text in paragraph.text:
                                    paragraph.clear()
                                    run = paragraph.add_run()
                                    run.add_picture(tmp_signature_path, width=Mm(100))
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                    found_placeholder = True
                                    break
                            if found_placeholder:
                                break
                        if found_placeholder:
                            break
                    if found_placeholder:
                        break
            
            doc.save(str(contract_path))
            return found_placeholder
        finally:
            if os.path.exists(tmp_signature_path):
                os.unlink(tmp_signature_path)
    except Exception as e:
        logger.error(f"嵌入签名到Word文档失败: {e}")
        raise


def _embed_signature_to_pdf(
    contract_path: Path,
    signature_blob: bytes | None,
    signature_type: str = "admin",
    x: float = None,
    y: float = None,
    width: float = None,
    height: float = None,
    page_index: int = 0,
    draw_signature: bool = True,
    date_text: str | None = None,
):
    """将签名嵌入到PDF文档的指定坐标位置"""
    try:
        from PyPDF2 import PdfReader, PdfWriter
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.utils import ImageReader
        import io
        
        contract_path = Path(contract_path)
        with open(contract_path, "rb") as f:
            original_pdf_bytes = f.read()
        reader = PdfReader(io.BytesIO(original_pdf_bytes))
        writer = PdfWriter()
        
        # 先获取目标页面的实际尺寸
        page_index = max(0, min(page_index, len(reader.pages) - 1))
        target_page = reader.pages[page_index]
        page_width = float(target_page.mediabox.width)
        page_height = float(target_page.mediabox.height)
        
        from PIL import Image

        def _trim_signature_image(img: Image.Image, padding: int = 5) -> Image.Image:
            img = img.convert("RGBA")
            w, h = img.size
            px = img.load()
            min_x, min_y = w, h
            max_x, max_y = -1, -1
            for yy in range(h):
                for xx in range(w):
                    r, g, b, a = px[xx, yy]
                    if a <= 10:
                        continue
                    if r >= 245 and g >= 245 and b >= 245:
                        continue
                    if xx < min_x:
                        min_x = xx
                    if yy < min_y:
                        min_y = yy
                    if xx > max_x:
                        max_x = xx
                    if yy > max_y:
                        max_y = yy
            if max_x < 0 or max_y < 0:
                return img
            left = max(min_x - padding, 0)
            top = max(min_y - padding, 0)
            right = min(max_x + padding + 1, w)
            bottom = min(max_y + padding + 1, h)
            return img.crop((left, top, right, bottom))

        # 使用原PDF页面的实际尺寸创建canvas，而不是固定的A4
        packet = io.BytesIO()
        can = canvas.Canvas(packet, pagesize=(page_width, page_height))

        signature_bottom_y: float | None = None
        if draw_signature:
            if not signature_blob:
                raise RuntimeError("签名数据为空，无法嵌入签名")

            signature_image = _trim_signature_image(Image.open(io.BytesIO(signature_blob)), padding=0)
            img_w, img_h = signature_image.size

            png_bytes = io.BytesIO()
            signature_image.save(png_bytes, format="PNG")
            png_bytes.seek(0)

            if x is not None and y is not None and width is not None and height is not None:
                is_ratio = 0 <= x <= 1 and 0 <= y <= 1 and 0 <= width <= 1 and 0 <= height <= 1
                if is_ratio:
                    pdf_x = x * page_width
                    pdf_y = (1 - y - height) * page_height
                    pdf_width = width * page_width
                    pdf_height = height * page_height
                else:
                    pdf_x = x
                    pdf_y = page_height - y - height
                    pdf_width = width
                    pdf_height = height

                pdf_width = max(1.0, min(float(pdf_width), page_width))
                pdf_height = max(1.0, min(float(pdf_height), page_height))
                pdf_x = max(0.0, min(float(pdf_x), page_width - pdf_width))
                pdf_y = max(0.0, min(float(pdf_y), page_height - pdf_height))

                pad = 2.0
                inner_w = max(float(pdf_width) - pad * 2, 1.0)
                inner_h = max(float(pdf_height) - pad * 2, 1.0)
                scale = min(inner_w / max(float(img_w), 1.0), inner_h / max(float(img_h), 1.0))
                draw_w = float(img_w) * scale
                draw_h = float(img_h) * scale

                draw_x = float(pdf_x) + pad
                draw_y = float(pdf_y) + pad

                max_x = float(pdf_x) + float(pdf_width) - pad - draw_w
                max_y = float(pdf_y) + float(pdf_height) - pad - draw_h
                if draw_x > max_x:
                    draw_x = max_x
                if draw_y > max_y:
                    draw_y = max_y

                img_reader = ImageReader(png_bytes)
                can.drawImage(img_reader, draw_x, draw_y, width=draw_w, height=draw_h, mask="auto")
                signature_bottom_y = float(draw_y)
            else:
                img_reader = ImageReader(png_bytes)
                box_w, box_h = 200, 80
                scale = min(box_w / max(img_w, 1), box_h / max(img_h, 1))
                draw_w = img_w * scale
                draw_h = img_h * scale
                can.drawImage(img_reader, 50, 50, width=draw_w, height=draw_h, mask="auto")
                signature_bottom_y = 50.0

        if date_text and x is not None and y is not None and width is not None and height is not None:
            is_ratio = 0 <= x <= 1 and 0 <= y <= 1 and 0 <= width <= 1 and 0 <= height <= 1

            # --- 自适应间距算法 ---
            # 计算签名框的基础高度和宽度
            if is_ratio:
                pdf_box_w = float(width) * page_width
                pdf_box_h = float(height) * page_height
                pdf_box_x = float(x) * page_width
                pdf_box_y = (1 - float(y) - float(height)) * page_height
            else:
                pdf_box_w = float(width)
                pdf_box_h = float(height)
                pdf_box_x = float(x)
                pdf_box_y = page_height - float(y) - float(height)

            pdf_box_w = max(1.0, min(float(pdf_box_w), page_width))
            pdf_box_h = max(1.0, min(float(pdf_box_h), page_height))
            pdf_box_x = max(0.0, min(float(pdf_box_x), page_width - pdf_box_w))
            pdf_box_y = max(0.0, min(float(pdf_box_y), page_height - pdf_box_h))

            gap = max(10.0, min(pdf_box_h * 0.22, 26.0))

            font_name = "Helvetica"
            font_size = max(14.0, min(pdf_box_h * 0.45, 22.0))
            can.setFont(font_name, font_size)

            try:
                text_width = can.stringWidth(date_text, font_name, font_size)
            except Exception:
                text_width = len(date_text) * font_size * 0.6

            pdf_date_x = pdf_box_x
            pdf_date_x = max(1.0, min(pdf_date_x, page_width - text_width - 1.0))

            anchor_y = signature_bottom_y if signature_bottom_y is not None else pdf_box_y
            baseline_y = anchor_y - gap - font_size
            baseline_y = max(1.0, min(baseline_y, anchor_y - 1.0))
            baseline_y = min(baseline_y, page_height - font_size - 1.0)

            can.setFillColorRGB(0, 0, 0)
            can.drawString(pdf_date_x, baseline_y, date_text)
        
        can.save()
        packet.seek(0)
        signature_pdf = PdfReader(packet)
        
        if len(reader.pages) > 0:
            page_index = max(0, min(page_index, len(reader.pages) - 1))
            target_page = reader.pages[page_index]
            signature_page = signature_pdf.pages[0]
            target_page.merge_page(signature_page)
        
        for page in reader.pages:
            writer.add_page(page)

        tmp_path = contract_path.parent / f".{contract_path.stem}.{uuid.uuid4().hex}.tmp.pdf"
        with open(tmp_path, "wb") as output_file:
            writer.write(output_file)

        with open(tmp_path, "rb") as f:
            header = f.read(4)
        if header != b"%PDF":
            raise RuntimeError("生成的PDF文件格式无效")
        size = tmp_path.stat().st_size
        with open(tmp_path, "rb") as f:
            f.seek(max(0, size - 2048))
            tail = f.read(2048)
        if b"%%EOF" not in tail:
            raise RuntimeError("生成的PDF文件不完整")

        tmp_path.replace(contract_path)

        return True
    except ImportError as e:
        logger.warning("PyPDF2或reportlab未安装，无法嵌入PDF签名")
        raise RuntimeError("PDF签名依赖缺失") from e
    except Exception as e:
        logger.error(f"嵌入签名到PDF失败: {e}")
        return False


def _embed_signature_to_image(contract_path: Path, signature_blob: bytes, signature_type: str = "admin", x: float = None, y: float = None, width: float = None, height: float = None):
    """将签名嵌入到图片的指定坐标位置"""
    try:
        from PIL import Image
        import io
        
        original_image = Image.open(contract_path)
        signature_image = Image.open(io.BytesIO(signature_blob))
        signature_image = signature_image.convert("RGBA")

        def _trim(img: Image.Image, padding: int = 5) -> Image.Image:
            w, h = img.size
            px = img.load()
            min_x, min_y = w, h
            max_x, max_y = -1, -1
            for yy in range(h):
                for xx in range(w):
                    r, g, b, a = px[xx, yy]
                    if a <= 10:
                        continue
                    if r >= 245 and g >= 245 and b >= 245:
                        continue
                    if xx < min_x:
                        min_x = xx
                    if yy < min_y:
                        min_y = yy
                    if xx > max_x:
                        max_x = xx
                    if yy > max_y:
                        max_y = yy
            if max_x < 0 or max_y < 0:
                return img
            left = max(min_x - padding, 0)
            top = max(min_y - padding, 0)
            right = min(max_x + padding + 1, w)
            bottom = min(max_y + padding + 1, h)
            return img.crop((left, top, right, bottom))

        signature_image = _trim(signature_image, padding=0)
        
        if width is not None and height is not None and x is not None and y is not None:
            is_ratio = 0 <= x <= 1 and 0 <= y <= 1 and 0 <= width <= 1 and 0 <= height <= 1
            if is_ratio:
                box_x = float(original_image.width) * float(x)
                box_y = float(original_image.height) * float(y)
                box_w = float(original_image.width) * float(width)
                box_h = float(original_image.height) * float(height)
            else:
                box_x = float(x)
                box_y = float(y)
                box_w = float(width)
                box_h = float(height)

            img_w, img_h = signature_image.size

            pad = 2.0
            inner_w = max(box_w - pad * 2, 1.0)
            inner_h = max(box_h - pad * 2, 1.0)
            scale = min(inner_w / max(float(img_w), 1.0), inner_h / max(float(img_h), 1.0))
            draw_w = max(int(round(img_w * scale)), 1)
            draw_h = max(int(round(img_h * scale)), 1)
            signature_image = signature_image.resize((draw_w, draw_h), Image.Resampling.LANCZOS)

            paste_x = int(round(box_x + pad))
            paste_y = int(round(box_y + box_h - pad - draw_h))
            max_x = box_x + box_w - pad - draw_w
            min_y = box_y + pad
            if paste_x > max_x:
                paste_x = int(round(max_x))
            if paste_y < min_y:
                paste_y = int(round(min_y))
        else:
            default_width = original_image.width // 4
            default_height = original_image.height // 8
            signature_image = signature_image.resize((max(default_width, 1), max(default_height, 1)), Image.Resampling.LANCZOS)
            paste_x = original_image.width - signature_image.width - 20
            paste_y = original_image.height - signature_image.height - 20

        paste_x = max(0, min(paste_x, original_image.width - signature_image.width))
        paste_y = max(0, min(paste_y, original_image.height - signature_image.height))
        
        if signature_image.mode == 'RGBA':
            original_image.paste(signature_image, (paste_x, paste_y), signature_image)
        else:
            original_image.paste(signature_image, (paste_x, paste_y))
        
        original_image.save(contract_path)
        return True
    except ImportError:
        logger.warning("PIL未安装，无法嵌入图片签名")
        return False
    except Exception as e:
        logger.error(f"嵌入签名到图片失败: {e}")
        return False

# 培训记录到期提醒相关常量和函数（需要在路由之前定义）
TRAINING_RECORD_REMINDER_SETTING_KEY = "training_record_reminder_days"
QUALIFICATION_EXPIRING_SETTING_KEY = "qualification_expiring_days"
DEFAULT_REMINDER_DAYS = 90


def get_training_record_reminder_days(db: Session) -> int:
    """获取培训记录提醒天数设置"""
    setting = db.query(SystemSetting).filter(SystemSetting.key == TRAINING_RECORD_REMINDER_SETTING_KEY).first()
    if setting and setting.value.isdigit():
        return max(int(setting.value), 1)
    qualification_setting = db.query(SystemSetting).filter(SystemSetting.key == QUALIFICATION_EXPIRING_SETTING_KEY).first()
    if qualification_setting and qualification_setting.value.isdigit():
        return max(int(qualification_setting.value), 1)
    return DEFAULT_REMINDER_DAYS


def calculate_expiry_date(completed_date: datetime) -> datetime:
    """计算到期日期：completed_date + 12个月"""
    expiry_date = completed_date + relativedelta(months=12)
    # 确保返回的日期是 aware datetime
    if expiry_date.tzinfo is None and completed_date.tzinfo is not None:
        expiry_date = expiry_date.replace(tzinfo=completed_date.tzinfo)
    elif expiry_date.tzinfo is None:
        expiry_date = expiry_date.replace(tzinfo=timezone.utc)
    return expiry_date


def calculate_reminder_status(days_until_expiry: int) -> str:
    """计算提醒状态"""
    if days_until_expiry < 0:
        return "expired"
    elif 6 <= days_until_expiry <= 8:
        return "1_week"
    elif 28 <= days_until_expiry <= 31:
        return "1_month"
    elif 90 <= days_until_expiry <= 91:
        return "3_months"
    else:
        return "normal"


def _qualification_certificate_url(qualification: Qualification) -> Optional[str]:
    certificate_url = qualification.certificate_url
    if not certificate_url and getattr(qualification, "certificate_blob", None):
        certificate_url = f"/api/houtai/qualifications/{qualification.id}/certificate"
    return certificate_url


def _qualification_to_training_record_dict(qualification: Qualification) -> dict:
    obtained_date = qualification.obtained_date
    expiry_date = qualification.expiry_date
    if obtained_date and obtained_date.tzinfo is None:
        obtained_date = obtained_date.replace(tzinfo=timezone.utc)
    if expiry_date and expiry_date.tzinfo is None:
        expiry_date = expiry_date.replace(tzinfo=timezone.utc)
    certificate_url = _qualification_certificate_url(qualification)
    return {
        "id": qualification.id,
        "name": qualification.name,
        "category": "certificate",
        "completed_date": obtained_date,
        "status": "completed",
        "score": None,
        "has_certificate": bool(certificate_url or getattr(qualification, "certificate_blob", None)),
        "certificate_number": qualification.certificate_number,
        "certificate_url": certificate_url,
        "certificate_obtained_date": obtained_date,
        "certificate_expiry_date": expiry_date,
        "training_institution": qualification.issuing_authority,
        "notes": None,
        "created_by": "admin",
    }


def _training_record_certificate_url(employee_id: str, record_id: str) -> str:
    return f"/api/houtai/employees/{employee_id}/training-records/{record_id}/certificate"


@router.get("", response_model=List[EmployeeResponse])
async def get_employees(
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """获取所有员工列表"""
    from sqlalchemy.orm import joinedload
    reminder_days = get_training_record_reminder_days(db)
    now = datetime.now(timezone.utc)
    today = now.date()
    unread_employee_ids = {
        (r.data_id or "")
        for r in db.query(BusinessUnread)
        .filter(
            BusinessUnread.receiver_user_id == str(current_user.id),
            BusinessUnread.business_code == "employee",
            BusinessUnread.is_unread == 1,
        )
        .all()
    }
    unread_employee_qualification_ids = {
        (r.data_id or "")
        for r in db.query(BusinessUnread)
        .filter(
            BusinessUnread.receiver_user_id == str(current_user.id),
            BusinessUnread.business_code == "employee_qualification",
            BusinessUnread.is_unread == 1,
        )
        .all()
    }
    employees = db.query(Employee).options(
        joinedload(Employee.qualifications),
        joinedload(Employee.training_records),
    ).order_by(Employee.created_at.desc()).all()

    week_start_date = today - timedelta(days=today.weekday())
    week_start = datetime(week_start_date.year, week_start_date.month, week_start_date.day)
    week_end = week_start + timedelta(days=7)
    counted_statuses = ["approved", "completed"]

    weekly_hours_by_employee: dict[str, float] = {}
    try:
        rows = (
            db.query(Task.assigned_employee_id, func.sum(TaskServiceItem.quantity))
            .join(TaskServiceItem, TaskServiceItem.task_id == Task.id)
            .filter(
                Task.assigned_employee_id.isnot(None),
                Task.status.in_(counted_statuses),
                Task.service_time >= week_start,
                Task.service_time < week_end,
            )
            .group_by(Task.assigned_employee_id)
            .all()
        )
        weekly_hours_by_employee = {str(emp_id): float(total or 0) for emp_id, total in rows if emp_id}

        tasks_without_items = (
            db.query(Task)
            .filter(
                Task.assigned_employee_id.isnot(None),
                Task.status.in_(counted_statuses),
                Task.service_time >= week_start,
                Task.service_time < week_end,
            )
            .filter(~db.query(TaskServiceItem.id).filter(TaskServiceItem.task_id == Task.id).exists())
            .all()
        )
        for t in tasks_without_items:
            emp_id = getattr(t, "assigned_employee_id", None)
            if not emp_id:
                continue
            raw = getattr(t, "service_duration_hours", None)
            try:
                hours = float(str(raw).strip().replace(",", ".")) if raw is not None else 0.0
            except Exception:
                hours = 0.0
            weekly_hours_by_employee[str(emp_id)] = weekly_hours_by_employee.get(str(emp_id), 0.0) + hours
    except Exception:
        weekly_hours_by_employee = {}

    result = []
    for emp in employees:
        quals = emp.qualifications or []
        records = emp.training_records or []

        expiring_list = []
        primary_item = None
        primary_rank = 2
        primary_expiry_ord = 10**12
        primary_days_until_expiry = None

        def _normalize_expiry_dt(value: Optional[datetime]) -> Optional[datetime]:
            if not value:
                return None
            if value.tzinfo is None:
                return value.replace(tzinfo=timezone.utc)
            return value

        candidate_items = []

        for record in records:
            if record.status != "completed":
                continue
            if not record.has_certificate and not (record.certificate_url and str(record.certificate_url).strip()):
                continue

            expiry_dt = record.certificate_expiry_date
            if not expiry_dt and record.category in ["first-aid", "manual-handling"] and record.completed_date:
                expiry_dt = calculate_expiry_date(record.completed_date)
            expiry_dt = _normalize_expiry_dt(expiry_dt)
            if not expiry_dt:
                continue

            obtained_dt = record.certificate_obtained_date or record.completed_date
            obtained_dt = _normalize_expiry_dt(obtained_dt)

            candidate_items.append({
                "id": record.id,
                "name": record.name,
                "certificate_number": record.certificate_number,
                "certificate_url": _training_record_certificate_url(emp.id, record.id),
                "obtained_date": obtained_dt or expiry_dt,
                "expiry_date": expiry_dt,
                "issuing_authority": record.training_institution,
            })

        if not candidate_items:
            for q in quals:
                expiry_dt = _normalize_expiry_dt(getattr(q, "expiry_date", None))
                if not expiry_dt:
                    continue
                obtained_dt = _normalize_expiry_dt(getattr(q, "obtained_date", None))
                candidate_items.append({
                    "id": q.id,
                    "name": getattr(q, "name", None),
                    "certificate_number": getattr(q, "certificate_number", None),
                    "certificate_url": _qualification_certificate_url(q),
                    "obtained_date": obtained_dt or expiry_dt,
                    "expiry_date": expiry_dt,
                    "issuing_authority": getattr(q, "issuing_authority", None),
                })

        def _candidate_sort_key(item: dict):
            expiry_dt = item.get("expiry_date")
            if not expiry_dt:
                return (2, 10**12)
            expiry_date_only = expiry_dt.date()
            days_until_expiry = (expiry_date_only - today).days
            rank = 0 if days_until_expiry < 0 else 1
            return (rank, expiry_date_only.toordinal())

        candidate_items.sort(key=_candidate_sort_key)
        expiring_list = candidate_items
        expiring_count = len(candidate_items)

        if candidate_items:
            first = candidate_items[0]
            expiry_dt = first.get("expiry_date")
            if expiry_dt:
                expiry_date_only = expiry_dt.date()
                primary_days_until_expiry = (expiry_date_only - today).days
                primary_rank = 0 if primary_days_until_expiry < 0 else 1
                primary_expiry_ord = expiry_date_only.toordinal()
                primary_item = first

        primary_certificate_number = None
        primary_expiry_date = None
        primary_qualification_name = None
        if primary_item is not None:
            primary_certificate_number = primary_item.get("certificate_number")
            primary_expiry_date = primary_item.get("expiry_date")
            primary_qualification_name = primary_item.get("name")

        def _normalize_cert_no(value: Optional[str]) -> str:
            if value is None:
                return "~~~~"
            s = str(value).strip()
            if not s:
                return "~~~~"
            return s.upper()

        if primary_item is None:
            employee_sort_key = (2, "~~~~", 10**12, emp.employee_number)
        else:
            employee_sort_key = (primary_rank, _normalize_cert_no(primary_certificate_number), primary_expiry_ord, emp.employee_number)

        result.append({
            "employee_number": emp.employee_number,
            "name": emp.name,
            "id": emp.id,
            "has_update": str(emp.id) in unread_employee_ids,
            "has_qualification_update": str(emp.id) in unread_employee_qualification_ids,
            "department": emp.department,
            "phone": emp.phone,
            "email": emp.email,
            "avatar_url": emp.avatar_url,
            "weekly_served_hours": float(round(weekly_hours_by_employee.get(str(emp.id), 0.0), 2)),
            "created_at": emp.created_at,
            "qualifications": quals,
            "training_records": [],
            "expiring_qualifications": expiring_list,
            "expiring_count": expiring_count,
            "expiring_primary_certificate_number": primary_certificate_number,
            "expiring_primary_expiry_date": primary_expiry_date,
            "expiring_primary_qualification_name": primary_qualification_name,
            "expiring_primary_days_until_expiry": primary_days_until_expiry,
            "_sort_key": employee_sort_key,
        })
    result.sort(key=lambda x: x.get("_sort_key"))
    for item in result:
        item.pop("_sort_key", None)
    return result


# 培训记录到期提醒相关API（必须在参数化路由之前定义）
@router.get("/training-records/expiring", response_model=List[ExpiringTrainingRecordItem])
async def get_expiring_training_records(
    advance_days: Optional[int] = Query(None, description="提前提醒天数"),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """获取即将到期的培训记录"""
    # 获取提醒天数设置
    reminder_days = advance_days if advance_days is not None else get_training_record_reminder_days(db)
    
    # 获取所有已完成且有到期日期的培训记录
    now = datetime.now(timezone.utc)
    # 只使用日期部分进行比较，忽略时间部分
    today = now.date()
    
    result = []

    qualification_rows = db.query(Qualification, Employee).join(
        Employee, Qualification.employee_id == Employee.id
    ).filter(
        Qualification.expiry_date.isnot(None)
    ).all()
    for qualification, employee in qualification_rows:
        expiry_date = qualification.expiry_date
        if not expiry_date:
            continue
        expiry_date_only = expiry_date.date()
        days_until_expiry = (expiry_date_only - today).days
        if 0 <= days_until_expiry <= reminder_days:
            if expiry_date.tzinfo is None:
                expiry_date = expiry_date.replace(tzinfo=timezone.utc)
            completed_date = qualification.obtained_date
            if completed_date and completed_date.tzinfo is None:
                completed_date = completed_date.replace(tzinfo=timezone.utc)
            result.append(ExpiringTrainingRecordItem(
                id=qualification.id,
                employee_id=employee.id,
                employee_name=employee.name,
                employee_number=employee.employee_number,
                name=qualification.name,
                category="certificate",
                completed_date=completed_date or expiry_date,
                expiry_date=expiry_date,
                days_until_expiry=days_until_expiry,
                reminder_status=calculate_reminder_status(days_until_expiry),
                certificate_url=_qualification_certificate_url(qualification),
                certificate_number=qualification.certificate_number,
                training_institution=qualification.issuing_authority
            ))

    records = db.query(TrainingRecord, Employee).join(
        Employee, TrainingRecord.employee_id == Employee.id
    ).filter(
        TrainingRecord.status == "completed",
        TrainingRecord.completed_date.isnot(None),
        or_(
            TrainingRecord.certificate_expiry_date.isnot(None),
            TrainingRecord.category.in_(["first-aid", "manual-handling"])
        )
    ).all()
    for record, employee in records:
        # 优先使用证书的到期日期，如果没有则计算（completed_date + 12个月，仅适用于first-aid和manual-handling）
        if record.certificate_expiry_date:
            expiry_date = record.certificate_expiry_date
            # 确保 expiry_date 是 aware datetime
            if expiry_date.tzinfo is None:
                expiry_date = expiry_date.replace(tzinfo=timezone.utc)
        elif record.category in ["first-aid", "manual-handling"]:
            # 对于first-aid和manual-handling，如果没有certificate_expiry_date，则计算（completed_date + 12个月）
            expiry_date = calculate_expiry_date(record.completed_date)
            # 确保计算出的 expiry_date 也是 aware datetime
            if expiry_date.tzinfo is None:
                expiry_date = expiry_date.replace(tzinfo=timezone.utc)
        else:
            # 其他类型如果没有certificate_expiry_date，跳过
            continue
        
        # 只使用日期部分计算距离到期的天数（忽略时间部分）
        expiry_date_only = expiry_date.date()
        days_until_expiry = (expiry_date_only - today).days
        
        # 包含在提醒范围内的记录：
        # 1. 未过期且在提醒范围内：0 <= days_until_expiry <= reminder_days
        # 2. 已过期但在提醒范围内：-reminder_days <= days_until_expiry < 0（已过期但刚过期不久，仍需要提醒）
        # 但"即将到期"标签页应该只显示未过期的，已过期的应该在"已过期"标签页
        if 0 <= days_until_expiry <= reminder_days:
            reminder_status = calculate_reminder_status(days_until_expiry)
            
            result.append(ExpiringTrainingRecordItem(
                id=record.id,
                employee_id=employee.id,
                employee_name=employee.name,
                employee_number=employee.employee_number,
                name=record.name,
                category=record.category,
                completed_date=record.completed_date,
                expiry_date=expiry_date,
                days_until_expiry=days_until_expiry,
                reminder_status=reminder_status,
                certificate_url=record.certificate_url,
                certificate_number=record.certificate_number,
                training_institution=record.training_institution
            ))
    
    result.sort(key=lambda x: x.expiry_date)
    return result


@router.get("/training-records/expired", response_model=List[ExpiringTrainingRecordItem])
async def get_expired_training_records(
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """获取已过期的培训记录"""
    # 获取所有已完成且有到期日期的培训记录
    now = datetime.now(timezone.utc)
    # 只使用日期部分进行比较，忽略时间部分
    today = now.date()
    
    result = []

    qualification_rows = db.query(Qualification, Employee).join(
        Employee, Qualification.employee_id == Employee.id
    ).filter(
        Qualification.expiry_date.isnot(None)
    ).all()
    for qualification, employee in qualification_rows:
        expiry_date = qualification.expiry_date
        if not expiry_date:
            continue
        expiry_date_only = expiry_date.date()
        days_until_expiry = (expiry_date_only - today).days
        if days_until_expiry < 0:
            if expiry_date.tzinfo is None:
                expiry_date = expiry_date.replace(tzinfo=timezone.utc)
            completed_date = qualification.obtained_date
            if completed_date and completed_date.tzinfo is None:
                completed_date = completed_date.replace(tzinfo=timezone.utc)
            result.append(ExpiringTrainingRecordItem(
                id=qualification.id,
                employee_id=employee.id,
                employee_name=employee.name,
                employee_number=employee.employee_number,
                name=qualification.name,
                category="certificate",
                completed_date=completed_date or expiry_date,
                expiry_date=expiry_date,
                days_until_expiry=days_until_expiry,
                reminder_status=calculate_reminder_status(days_until_expiry),
                certificate_url=_qualification_certificate_url(qualification),
                certificate_number=qualification.certificate_number,
                training_institution=qualification.issuing_authority
            ))

    records = db.query(TrainingRecord, Employee).join(
        Employee, TrainingRecord.employee_id == Employee.id
    ).filter(
        TrainingRecord.status == "completed",
        TrainingRecord.completed_date.isnot(None),
        or_(
            TrainingRecord.certificate_expiry_date.isnot(None),
            TrainingRecord.category.in_(["first-aid", "manual-handling"])
        )
    ).all()
    for record, employee in records:
        # 优先使用证书的到期日期，如果没有则计算（completed_date + 12个月，仅适用于first-aid和manual-handling）
        if record.certificate_expiry_date:
            expiry_date = record.certificate_expiry_date
            # 确保 expiry_date 是 aware datetime
            if expiry_date.tzinfo is None:
                expiry_date = expiry_date.replace(tzinfo=timezone.utc)
        elif record.category in ["first-aid", "manual-handling"]:
            # 对于first-aid和manual-handling，如果没有certificate_expiry_date，则计算（completed_date + 12个月）
            expiry_date = calculate_expiry_date(record.completed_date)
            # 确保计算出的 expiry_date 也是 aware datetime
            if expiry_date.tzinfo is None:
                expiry_date = expiry_date.replace(tzinfo=timezone.utc)
        else:
            # 其他类型如果没有certificate_expiry_date，跳过
            continue
        
        # 只使用日期部分计算距离到期的天数（忽略时间部分）
        expiry_date_only = expiry_date.date()
        days_until_expiry = (expiry_date_only - today).days
        
        if days_until_expiry < 0:
            reminder_status = calculate_reminder_status(days_until_expiry)
            
            result.append(ExpiringTrainingRecordItem(
                id=record.id,
                employee_id=employee.id,
                employee_name=employee.name,
                employee_number=employee.employee_number,
                name=record.name,
                category=record.category,
                completed_date=record.completed_date,
                expiry_date=expiry_date,
                days_until_expiry=days_until_expiry,
                reminder_status=reminder_status,
                certificate_url=record.certificate_url,
                certificate_number=record.certificate_number,
                training_institution=record.training_institution
            ))
    
    result.sort(key=lambda x: x.expiry_date, reverse=True)
    return result


@router.get("/training-records/reminder-settings")
async def get_training_record_reminder_settings(
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """获取培训记录提醒设置"""
    return {"days": get_training_record_reminder_days(db)}


@router.put("/training-records/reminder-settings")
async def update_training_record_reminder_settings(
    payload: ReminderSettingPayload,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """更新培训记录提醒设置"""
    if payload.days <= 0:
        raise HTTPException(status_code=400, detail="提醒天数必须大于0")
    
    setting = db.query(SystemSetting).filter(
        SystemSetting.key == TRAINING_RECORD_REMINDER_SETTING_KEY
    ).first()
    
    if setting:
        setting.value = str(payload.days)
    else:
        setting = SystemSetting(key=TRAINING_RECORD_REMINDER_SETTING_KEY, value=str(payload.days))
        db.add(setting)

    qualification_setting = db.query(SystemSetting).filter(
        SystemSetting.key == QUALIFICATION_EXPIRING_SETTING_KEY
    ).first()
    if qualification_setting:
        qualification_setting.value = str(payload.days)
    else:
        db.add(SystemSetting(key=QUALIFICATION_EXPIRING_SETTING_KEY, value=str(payload.days)))
    
    db.commit()
    return {"days": payload.days}


@router.post("", response_model=EmployeeResponse, status_code=status.HTTP_201_CREATED)
async def create_employee(
    employee_data: EmployeeCreate,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """创建员工"""
    incoming_email = (employee_data.email or "").strip().lower()
    if incoming_email:
        existed = db.query(Employee).filter(func.lower(func.trim(Employee.email)) == incoming_email).first()
        if existed:
            raise HTTPException(status_code=400, detail="该邮箱号已绑定员工，请重新输入")

    def _next_employee_number() -> str:
        rows = db.query(Employee.employee_number).all()
        max_value = 0
        for (value,) in rows:
            if not value:
                continue
            if re.fullmatch(r"\d+", value):
                max_value = max(max_value, int(value))
        return f"{max_value + 1:06d}"

    employee_dict = employee_data.dict()
    if incoming_email:
        employee_dict["email"] = incoming_email
    employee_dict["password_hash"] = get_password_hash(employee_dict.pop("password"))
    employee_dict["account_status"] = "normal"

    for _ in range(10):
        employee_dict["employee_number"] = _next_employee_number()
        employee = Employee(**employee_dict)
        db.add(employee)
        try:
            db.commit()
        except IntegrityError:
            db.rollback()
            continue

        db.refresh(employee)

        # 为新员工自动分配员工手册
        try:
            # 找到员工手册源文件
            # 这里的 static 目录是在 aozhou-backend/static
            # settings.upload_dir 可能是 aozhou-backend/uploads
            project_root = Path(__file__).resolve().parents[3]
            handbook_src = project_root / "static" / "Employee Handbook.docx"
            
            if handbook_src.exists():
                with open(handbook_src, "rb") as f:
                    content = f.read()
                
                # 保存到员工手册目录
                filename = f"Employee Handbook - {employee.name}.docx"
                # save_upload_file 是异步函数，这里在 async def create_employee 中调用是正确的
                file_url = await save_upload_file(content, filename, "employee_handbooks")
                
                # 创建文档记录
                handbook_doc = EmployeeDocument(
                    employee_id=employee.id,
                    name="Employee Handbook.docx",
                    file_type="docx",
                    file_url=file_url,
                    document_type="handbook",
                    uploaded_by=str(current_user.id)
                )
                db.add(handbook_doc)
                db.commit()
                logger.info(f"Automatically assigned Employee Handbook to new employee {employee.id}")
            else:
                logger.warning(f"Employee Handbook source file not found at {handbook_src}")
        except Exception as e:
            db.rollback()
            logger.error(f"Failed to assign Employee Handbook to new employee {employee.id}: {str(e)}")
            # 不因为手册分配失败而导致创建员工失败

        return employee

    raise HTTPException(status_code=500, detail="生成员工号失败，请重试")


@router.get("/{employee_id}", response_model=EmployeeResponse)
async def get_employee(
    employee_id: str,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """获取员工详情"""
    from sqlalchemy.orm import joinedload
    unread_employee_qualification = (
        db.query(BusinessUnread)
        .filter(
            BusinessUnread.receiver_user_id == str(current_user.id),
            BusinessUnread.business_code == "employee_qualification",
            BusinessUnread.data_id == str(employee_id),
            BusinessUnread.is_unread == 1,
        )
        .first()
        is not None
    )
    employee = db.query(Employee).options(
        joinedload(Employee.qualifications)
    ).filter(Employee.id == employee_id).first()
    if not employee:
        raise HTTPException(status_code=404, detail="员工不存在")
    training_records = []
    for qual in (employee.qualifications or []):
        certificate_url = qual.certificate_url
        if not certificate_url and getattr(qual, "certificate_blob", None):
            certificate_url = f"/api/houtai/qualifications/{qual.id}/certificate"
        training_records.append({
            "id": qual.id,
            "name": qual.name,
            "category": "certificate",
            "completed_date": qual.obtained_date,
            "status": "completed",
            "score": None,
            "has_certificate": bool(certificate_url or getattr(qual, "certificate_blob", None)),
            "certificate_number": qual.certificate_number,
            "certificate_url": certificate_url,
            "certificate_obtained_date": qual.obtained_date,
            "certificate_expiry_date": qual.expiry_date,
            "training_institution": qual.issuing_authority,
            "notes": None,
            "created_by": "admin",
        })
    return {
        "employee_number": employee.employee_number,
        "name": employee.name,
        "id": employee.id,
        "has_qualification_update": bool(unread_employee_qualification),
        "department": employee.department,
        "phone": employee.phone,
        "email": employee.email,
        "avatar_url": employee.avatar_url,
        "account_status": getattr(employee, "account_status", None),
        "created_at": employee.created_at,
        "qualifications": employee.qualifications or [],
        "training_records": training_records,
    }


@router.put("/{employee_id}", response_model=EmployeeResponse)
async def update_employee(
    employee_id: str,
    employee_data: EmployeeUpdate,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """更新员工信息"""
    employee = db.query(Employee).filter(Employee.id == employee_id).first()
    if not employee:
        raise HTTPException(status_code=404, detail="员工不存在")

    update_data = employee_data.dict(exclude_unset=True)

    if "email" in update_data:
        normalized_email = (update_data.get("email") or "").strip().lower()
        if normalized_email:
            existed = db.query(Employee).filter(
                Employee.id != employee_id,
                func.lower(func.trim(Employee.email)) == normalized_email
            ).first()
            if existed:
                raise HTTPException(status_code=400, detail="该邮箱号已绑定员工，请重新输入")
            update_data["email"] = normalized_email
        else:
            raise HTTPException(status_code=400, detail="邮箱必填")

    # 如果更新了密码，需要进行哈希
    if "password" in update_data:
        update_data["password_hash"] = get_password_hash(update_data.pop("password"))

    for key, value in update_data.items():
        setattr(employee, key, value)

    db.commit()
    db.refresh(employee)
    return employee


@router.put("/{employee_id}/account-status", response_model=EmployeeResponse)
async def update_employee_account_status(
    employee_id: str,
    payload: EmployeeAccountStatusUpdate,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user),
):
    employee = db.query(Employee).filter(Employee.id == employee_id).first()
    if not employee:
        raise HTTPException(status_code=404, detail="员工不存在")

    normalized_status = (payload.account_status or "").strip().lower()
    if normalized_status not in {"normal", "disabled"}:
        raise HTTPException(status_code=400, detail="账号状态无效")

    employee.account_status = normalized_status
    db.commit()
    db.refresh(employee)
    return employee


@router.delete("/{employee_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_employee(
    employee_id: str,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """删除员工"""
    employee = db.query(Employee).filter(Employee.id == employee_id).first()
    if not employee:
        raise HTTPException(status_code=404, detail="员工不存在")
    
    db.delete(employee)
    db.commit()
    return None


@router.post("/{employee_id}/qualifications", status_code=status.HTTP_201_CREATED)
async def add_qualification(
    employee_id: str,
    qualification_data: QualificationCreate,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """添加员工资质"""
    employee = db.query(Employee).filter(Employee.id == employee_id).first()
    if not employee:
        raise HTTPException(status_code=404, detail="员工不存在")
    
    qualification = Qualification(employee_id=employee_id, **qualification_data.dict())
    db.add(qualification)
    db.commit()
    db.refresh(qualification)
    return qualification


@router.delete("/{employee_id}/qualifications/{qualification_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_qualification(
    employee_id: str,
    qualification_id: str,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """删除员工资质"""
    qualification = db.query(Qualification).filter(
        Qualification.id == qualification_id,
        Qualification.employee_id == employee_id
    ).first()
    if not qualification:
        raise HTTPException(status_code=404, detail="资质不存在")
    
    db.delete(qualification)
    db.commit()
    return None


@router.post("/{employee_id}/training-records", status_code=status.HTTP_201_CREATED)
async def add_training_record(
    employee_id: str,
    training_data: TrainingRecordCreate,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """添加培训记录"""
    employee = db.query(Employee).filter(Employee.id == employee_id).first()
    if not employee:
        raise HTTPException(status_code=404, detail="员工不存在")
    
    training_payload = training_data.model_dump()

    def _strip_tz(value):
        if value is None:
            return None
        if isinstance(value, datetime) and value.tzinfo is not None:
            return value.replace(tzinfo=None)
        return value

    raw_status = training_payload.get("status")
    if raw_status is None or str(raw_status).strip() == "":
        training_payload["status"] = TrainingStatus.completed.value
    else:
        normalized = _normalize_training_status(getattr(raw_status, "value", raw_status))
        training_payload["status"] = normalized

    if training_payload.get("has_certificate") is None:
        training_payload["has_certificate"] = False

    training_payload["completed_date"] = _strip_tz(training_payload.get("completed_date"))
    training_payload["certificate_obtained_date"] = _strip_tz(training_payload.get("certificate_obtained_date"))
    training_payload["certificate_expiry_date"] = _strip_tz(training_payload.get("certificate_expiry_date"))

    if not training_payload["has_certificate"]:
        training_payload["certificate_number"] = None
        training_payload["certificate_url"] = None
        training_payload["certificate_obtained_date"] = None
        training_payload["certificate_expiry_date"] = None

    training_payload["created_by"] = "admin"

    training_record = TrainingRecord(employee_id=employee_id, **training_payload)
    db.add(training_record)
    try:
        db.commit()
    except SQLAlchemyError as e:
        db.rollback()
        err = str(getattr(e, "orig", e))
        raise HTTPException(status_code=500, detail=f"数据库操作失败: {err[:500]}")
    db.refresh(training_record)
    return training_record


@router.get("/{employee_id}/training-records", response_model=List[TrainingRecordResponse])
async def get_training_records(
    employee_id: str,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """获取员工培训记录"""
    qualifications = db.query(Qualification).filter(Qualification.employee_id == employee_id).all()
    qualification_records = [_qualification_to_training_record_dict(q) for q in (qualifications or [])]
    records = db.query(TrainingRecord).filter(TrainingRecord.employee_id == employee_id).all()
    record_dicts = []
    for record in (records or []):
        cert_url = record.certificate_url
        if record.has_certificate or (record.certificate_url and str(record.certificate_url).strip()):
            cert_url = _training_record_certificate_url(employee_id, record.id)
        record_dicts.append({
            "id": record.id,
            "name": record.name,
            "category": record.category,
            "completed_date": record.completed_date,
            "status": record.status,
            "score": record.score,
            "has_certificate": record.has_certificate,
            "certificate_number": record.certificate_number,
            "certificate_url": cert_url,
            "certificate_obtained_date": record.certificate_obtained_date,
            "certificate_expiry_date": record.certificate_expiry_date,
            "training_institution": record.training_institution,
            "notes": record.notes,
            "created_by": getattr(record, "created_by", None),
        })
    result = qualification_records + record_dicts
    def _sort_key(item):
        value = item.get("completed_date") if isinstance(item, dict) else getattr(item, "completed_date", None)
        if isinstance(value, datetime):
            return value.replace(tzinfo=value.tzinfo or timezone.utc)
        return datetime.min.replace(tzinfo=timezone.utc)
    result.sort(key=_sort_key, reverse=True)
    return result

@router.post("/{employee_id}/training-records/batch-upload", response_model=List[TrainingRecordResponse], status_code=status.HTTP_201_CREATED)
async def batch_upload_training_certificates(
    employee_id: str,
    files: List[UploadFile] = File(...),
    names: Optional[List[str]] = Form(None),
    certificate_numbers: Optional[List[str]] = Form(None),
    obtained_dates: Optional[List[str]] = Form(None),
    expiry_dates: Optional[List[str]] = Form(None),
    issuing_authorities: Optional[List[str]] = Form(None),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """批量上传证书并创建资质（以培训记录形式展示）"""
    employee = db.query(Employee).filter(Employee.id == employee_id).first()
    if not employee:
        raise HTTPException(status_code=404, detail="员工不存在")

    created = []
    now = datetime.now(timezone.utc)

    def _get_value(values: Optional[List[str]], index: int) -> Optional[str]:
        if not values:
            return None
        if index < 0 or index >= len(values):
            return None
        value = values[index]
        return value if value is not None and str(value).strip() else None

    for i, file in enumerate(files):
        file_content = await file.read()
        file_mime = file.content_type or "application/octet-stream"
        name = _get_value(names, i) or Path(file.filename).stem
        certificate_number = _get_value(certificate_numbers, i)
        obtained_date_str = _get_value(obtained_dates, i)
        expiry_date_str = _get_value(expiry_dates, i)
        issuing_authority = _get_value(issuing_authorities, i)

        obtained_date = datetime.fromisoformat(obtained_date_str) if obtained_date_str else now
        expiry_date = datetime.fromisoformat(expiry_date_str) if expiry_date_str else None

        qualification_id = str(uuid.uuid4())
        qualification = Qualification(
            id=qualification_id,
            employee_id=employee_id,
            name=name,
            certificate_number=certificate_number,
            certificate_blob=file_content,
            certificate_mime=file_mime,
            obtained_date=obtained_date,
            expiry_date=expiry_date,
            issuing_authority=issuing_authority
        )
        qualification.certificate_url = f"/api/houtai/qualifications/{qualification_id}/certificate"
        db.add(qualification)
        created.append(qualification)

    db.commit()
    return [_qualification_to_training_record_dict(q) for q in created]


@router.put("/{employee_id}/training-records/{record_id}", response_model=TrainingRecordResponse)
async def update_training_record(
    employee_id: str,
    record_id: str,
    training_data: TrainingRecordUpdate,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """更新培训记录"""
    record = db.query(TrainingRecord).filter(
        TrainingRecord.id == record_id,
        TrainingRecord.employee_id == employee_id
    ).first()
    if record:
        update_payload = training_data.model_dump(exclude_unset=True)

        def _strip_tz(value):
            if value is None:
                return None
            if isinstance(value, datetime) and value.tzinfo is not None:
                return value.replace(tzinfo=None)
            return value

        if "status" in update_payload:
            normalized = _normalize_training_status(getattr(update_payload["status"], "value", update_payload["status"]))
            update_payload["status"] = normalized

        if "completed_date" in update_payload:
            update_payload["completed_date"] = _strip_tz(update_payload.get("completed_date"))
        if "certificate_obtained_date" in update_payload:
            update_payload["certificate_obtained_date"] = _strip_tz(update_payload.get("certificate_obtained_date"))
        if "certificate_expiry_date" in update_payload:
            update_payload["certificate_expiry_date"] = _strip_tz(update_payload.get("certificate_expiry_date"))

        if update_payload.get("has_certificate") is False:
            update_payload["certificate_number"] = None
            update_payload["certificate_url"] = None
            update_payload["certificate_obtained_date"] = None
            update_payload["certificate_expiry_date"] = None

        if "certificate_url" in update_payload and update_payload.get("certificate_url") is not None:
            if str(update_payload["certificate_url"]).strip() == "":
                update_payload["certificate_url"] = None

        for key, value in update_payload.items():
            setattr(record, key, value)

        try:
            db.commit()
        except SQLAlchemyError as e:
            db.rollback()
            err = str(getattr(e, "orig", e))
            raise HTTPException(status_code=500, detail=f"数据库操作失败: {err[:500]}")
        db.refresh(record)

        response = TrainingRecordResponse.model_validate(record).model_dump()
        if record.has_certificate or (record.certificate_url and str(record.certificate_url).strip()):
            response["certificate_url"] = _training_record_certificate_url(employee_id, record.id)
        return response

    qualification = db.query(Qualification).filter(
        Qualification.id == record_id,
        Qualification.employee_id == employee_id
    ).first()
    if not qualification:
        raise HTTPException(status_code=404, detail="培训记录不存在")

    update_payload = training_data.model_dump(exclude_unset=True)
    if "name" in update_payload:
        qualification.name = update_payload["name"]
    if "certificate_number" in update_payload:
        qualification.certificate_number = update_payload["certificate_number"]
    if "completed_date" in update_payload:
        qualification.obtained_date = update_payload["completed_date"]
    if "certificate_obtained_date" in update_payload:
        qualification.obtained_date = update_payload["certificate_obtained_date"]
    if "certificate_expiry_date" in update_payload:
        qualification.expiry_date = update_payload["certificate_expiry_date"]
    if "training_institution" in update_payload:
        qualification.issuing_authority = update_payload["training_institution"]

    try:
        db.commit()
    except SQLAlchemyError as e:
        db.rollback()
        err = str(getattr(e, "orig", e))
        raise HTTPException(status_code=500, detail=f"数据库操作失败: {err[:500]}")

    response = _qualification_to_training_record_dict(qualification)
    return response


@router.put("/{employee_id}/training-records/{record_id}/upload", response_model=TrainingRecordResponse)
async def update_training_record_with_file(
    employee_id: str,
    record_id: str,
    name: str = Form(None),
    category: str = Form(None),
    completed_date: str = Form(None),
    status: str = Form(None),
    score: str = Form(None),
    has_certificate: bool = Form(None),
    certificate_number: str = Form(None),
    certificate_obtained_date: str = Form(None),
    certificate_expiry_date: str = Form(None),
    training_institution: str = Form(None),
    notes: str = Form(None),
    file: UploadFile = File(None),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """更新培训记录（支持文件上传）"""
    record = db.query(TrainingRecord).filter(
        TrainingRecord.id == record_id,
        TrainingRecord.employee_id == employee_id
    ).first()
    if not record:
        raise HTTPException(status_code=404, detail="培训记录不存在")

    try:
        if file and getattr(file, "filename", None):
            file_content = await file.read()
            if file_content:
                file_ext = os.path.splitext(file.filename)[1]
                filename = f"{uuid.uuid4()}{file_ext}"
                certificate_url = await save_upload_file(file_content, filename, subfolder="training_certificates")
                record.certificate_url = certificate_url
                record.has_certificate = True

        if name is not None:
            record.name = name
        if category is not None:
            record.category = category
        if completed_date is not None:
            record.completed_date = datetime.fromisoformat(completed_date)

        normalized_status = _normalize_training_status(status)
        if normalized_status is not None:
            record.status = normalized_status

        if score is not None:
            record.score = score
        if has_certificate is not None:
            record.has_certificate = has_certificate
        if certificate_number is not None:
            record.certificate_number = certificate_number
        if certificate_obtained_date is not None:
            if certificate_obtained_date and certificate_obtained_date.strip():
                try:
                    record.certificate_obtained_date = datetime.fromisoformat(certificate_obtained_date)
                except (ValueError, TypeError):
                    record.certificate_obtained_date = None
            else:
                record.certificate_obtained_date = None
        if certificate_expiry_date is not None:
            if certificate_expiry_date and certificate_expiry_date.strip():
                try:
                    record.certificate_expiry_date = datetime.fromisoformat(certificate_expiry_date)
                except (ValueError, TypeError):
                    record.certificate_expiry_date = None
            else:
                record.certificate_expiry_date = None
        if training_institution is not None:
            record.training_institution = training_institution
        if notes is not None:
            record.notes = notes

        db.commit()
        db.refresh(record)
        return record
    except HTTPException:
        db.rollback()
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"更新培训记录失败: employee_id={employee_id}, record_id={record_id}, error={e}")
        raise HTTPException(status_code=500, detail="更新培训记录失败")


@router.post("/{employee_id}/training-records/upload", status_code=status.HTTP_201_CREATED)
async def add_training_record_with_file(
    employee_id: str,
    name: str = Form(...),
    category: str = Form(None),
    completed_date: str = Form(...),
    status: str = Form(None),
    score: str = Form(None),
    has_certificate: bool = Form(False),
    certificate_number: str = Form(None),
    certificate_obtained_date: str = Form(None),
    certificate_expiry_date: str = Form(None),
    training_institution: str = Form(None),
    notes: str = Form(None),
    file: UploadFile = File(None),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """添加培训记录（可选证书文件）"""
    employee = db.query(Employee).filter(Employee.id == employee_id).first()
    if not employee:
        raise HTTPException(status_code=404, detail="员工不存在")

    normalized_status = _normalize_training_status(status) if status is not None else None

    certificate_url = None
    if file:
        file_ext = os.path.splitext(file.filename)[1]
        filename = f"{uuid.uuid4()}{file_ext}"
        file_content = await file.read()
        certificate_url = await save_upload_file(file_content, filename, subfolder="training_certificates")

    completed_dt = datetime.fromisoformat(completed_date)
    if completed_dt.tzinfo is not None:
        completed_dt = completed_dt.replace(tzinfo=None)

    cert_obtained_dt = datetime.fromisoformat(certificate_obtained_date) if certificate_obtained_date else None
    if isinstance(cert_obtained_dt, datetime) and cert_obtained_dt.tzinfo is not None:
        cert_obtained_dt = cert_obtained_dt.replace(tzinfo=None)

    cert_expiry_dt = datetime.fromisoformat(certificate_expiry_date) if certificate_expiry_date else None
    if isinstance(cert_expiry_dt, datetime) and cert_expiry_dt.tzinfo is not None:
        cert_expiry_dt = cert_expiry_dt.replace(tzinfo=None)

    resolved_has_certificate = bool(has_certificate) or bool(file)
    if not resolved_has_certificate:
        certificate_number = None
        certificate_url = None
        cert_obtained_dt = None
        cert_expiry_dt = None

    record = TrainingRecord(
        employee_id=employee_id,
        name=name,
        category=category,
        completed_date=completed_dt,
        status=normalized_status or "completed",
        score=score,
        has_certificate=resolved_has_certificate,
        certificate_number=certificate_number,
        certificate_url=certificate_url,
        certificate_obtained_date=cert_obtained_dt,
        certificate_expiry_date=cert_expiry_dt,
        training_institution=training_institution,
        notes=notes,
        created_by="admin"  # 标识为管理员创建
    )
    db.add(record)
    try:
        db.commit()
    except SQLAlchemyError as e:
        db.rollback()
        err = str(getattr(e, "orig", e))
        raise HTTPException(status_code=500, detail=f"数据库操作失败: {err[:500]}")
    db.refresh(record)
    return record


@router.put("/{employee_id}/training-records/{record_id}/approve", response_model=TrainingRecordResponse)
async def approve_training_record(
    employee_id: str,
    record_id: str,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """审核通过培训记录"""
    record = db.query(TrainingRecord).filter(
        TrainingRecord.id == record_id,
        TrainingRecord.employee_id == employee_id
    ).first()
    
    if not record:
        raise HTTPException(status_code=404, detail="培训记录不存在")
    
    if record.status != 'pending':
        raise HTTPException(status_code=400, detail="只能审核待审核状态的记录")
    
    if record.created_by != 'employee':
        raise HTTPException(status_code=400, detail="只能审核员工提交的记录")
    
    record.status = 'completed'
    db.commit()
    db.refresh(record)
    return record


@router.put("/{employee_id}/training-records/{record_id}/reject", response_model=TrainingRecordResponse)
async def reject_training_record(
    employee_id: str,
    record_id: str,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """审核驳回培训记录"""
    record = db.query(TrainingRecord).filter(
        TrainingRecord.id == record_id,
        TrainingRecord.employee_id == employee_id
    ).first()
    
    if not record:
        raise HTTPException(status_code=404, detail="培训记录不存在")
    
    if record.status != 'pending':
        raise HTTPException(status_code=400, detail="只能审核待审核状态的记录")
    
    if record.created_by != 'employee':
        raise HTTPException(status_code=400, detail="只能审核员工提交的记录")
    
    record.status = 'rejected'
    db.commit()
    db.refresh(record)
    return record


@router.delete("/{employee_id}/training-records/{record_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_training_record(
    employee_id: str,
    record_id: str,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """删除培训记录"""
    record = db.query(TrainingRecord).filter(
        TrainingRecord.id == record_id,
        TrainingRecord.employee_id == employee_id
    ).first()
    if not record:
        raise HTTPException(status_code=404, detail="培训记录不存在")
    
    db.delete(record)
    db.commit()
    return None


@router.get("/{employee_id}/training-records/{record_id}/certificate")
async def get_training_record_certificate(
    employee_id: str,
    record_id: str,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """获取培训记录证书文件"""
    record = db.query(TrainingRecord).filter(
        TrainingRecord.id == record_id,
        TrainingRecord.employee_id == employee_id
    ).first()
    if not record:
        raise HTTPException(status_code=404, detail="培训记录不存在")
    
    # 获取文件路径
    file_path = None
    certificate_url = (record.certificate_url or "").strip()
    if certificate_url:
        file_path = get_file_path(certificate_url)
        if not file_path:
            upload_path = Path(settings.upload_dir).resolve()
            relative_path = certificate_url.replace("uploads/", "").replace("uploads\\", "")
            if str(upload_path) in certificate_url:
                relative_path = certificate_url.replace(str(upload_path), "").lstrip("/\\")
            possible_paths = [
                Path(certificate_url),
                upload_path / relative_path,
                upload_path / "training_certificates" / relative_path,
                Path(relative_path),
            ]
            for p in possible_paths:
                if p.exists() and p.is_file():
                    file_path = p
                    break
    
    if not file_path or not file_path.exists() or not file_path.is_file():
        fallback_qual = None
        if record.certificate_number:
            fallback_qual = db.query(Qualification).filter(
                Qualification.employee_id == employee_id,
                Qualification.certificate_number == record.certificate_number
            ).first()
        if not fallback_qual:
            fallback_qual = db.query(Qualification).filter(
                Qualification.employee_id == employee_id,
                Qualification.name == record.name
            ).first()
        if fallback_qual:
            if getattr(fallback_qual, "certificate_blob", None):
                blob_mime = (fallback_qual.certificate_mime or "application/octet-stream").split(";")[0].strip()
                blob_suffix = mimetypes.guess_extension(blob_mime) or ""
                blob_filename = f"certificate_{record_id}{blob_suffix}"
                return Response(
                    content=fallback_qual.certificate_blob,
                    media_type=blob_mime or "application/octet-stream",
                    headers={
                        "Content-Disposition": f'inline; filename="{blob_filename}"'
                    }
                )
            if getattr(fallback_qual, "certificate_url", None):
                fallback_url = fallback_qual.certificate_url.strip()
                upload_path = Path(settings.upload_dir).resolve()
                relative_path = fallback_url.replace("uploads/", "").replace("uploads\\", "")
                if str(upload_path) in fallback_url:
                    relative_path = fallback_url.replace(str(upload_path), "").lstrip("/\\")
                possible_paths = [
                    Path(fallback_url),
                    get_file_path(fallback_url) or Path(""),
                    upload_path / relative_path,
                    upload_path / "training_certificates" / relative_path,
                    Path(relative_path),
                ]
                for p in possible_paths:
                    if p.exists() and p.is_file():
                        mime_type, _ = mimetypes.guess_type(str(p))
                        if not mime_type:
                            mime_type = "application/octet-stream"
                        return FileResponse(
                            path=str(p),
                            media_type=mime_type,
                            filename=f"certificate_{record_id}{p.suffix}",
                            headers={
                                "Content-Disposition": f"inline; filename=certificate_{record_id}{p.suffix}",
                                "Cache-Control": "public, max-age=3600",
                            }
                        )
        raise HTTPException(status_code=404, detail="证书文件不存在")
    
    # 获取MIME类型
    mime_type, _ = mimetypes.guess_type(str(file_path))
    if not mime_type:
        mime_type = "application/octet-stream"
    
    # 获取文件修改时间用于ETag
    file_stat = file_path.stat()
    etag = f'"{record_id}_{int(file_stat.st_mtime)}"'
    
    return FileResponse(
        path=str(file_path),
        media_type=mime_type,
        filename=f"certificate_{record_id}{file_path.suffix}",
        headers={
            "Content-Disposition": f"inline; filename=certificate_{record_id}{file_path.suffix}",
            "Cache-Control": "public, max-age=3600",  # 缓存1小时
            "ETag": etag  # 基于文件修改时间的ETag，支持条件请求
        }
    )


# 员工文档相关API
@router.get("/{employee_id}/documents", response_model=List[EmployeeDocumentResponse])
async def get_employee_documents(
    employee_id: str,
    document_type: Optional[str] = Query(None, description="文档类型：contract, checklist, code, tracker, handbook, onboarding"),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """获取员工文档列表"""
    employee = db.query(Employee).filter(Employee.id == employee_id).first()
    if not employee:
        raise HTTPException(status_code=404, detail="员工不存在")
    
    query = db.query(EmployeeDocument).filter(EmployeeDocument.employee_id == employee_id)
    if document_type:
        query = query.filter(EmployeeDocument.document_type == document_type)
    
    documents = query.order_by(EmployeeDocument.uploaded_at.desc()).all()
    return documents


@router.post("/{employee_id}/documents/upload", response_model=EmployeeDocumentResponse)
async def upload_employee_document(
    employee_id: str,
    document_type: str = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """上传员工文档"""
    employee = db.query(Employee).filter(Employee.id == employee_id).first()
    if not employee:
        raise HTTPException(status_code=404, detail="员工不存在")
    
    # 验证文档类型
    valid_types = ['contract', 'checklist', 'code', 'tracker', 'handbook', 'onboarding']
    if document_type not in valid_types:
        raise HTTPException(status_code=400, detail=f"无效的文档类型，必须是: {', '.join(valid_types)}")
    
    # 上传文件
    file_ext = os.path.splitext(file.filename)[1]
    filename = f"{uuid.uuid4()}{file_ext}"
    file_content = await file.read()
    
    # 根据文档类型选择子文件夹
    subfolder_map = {
        'contract': 'employee_contracts',
        'checklist': 'employee_checklists',
        'code': 'employee_codes',
        'tracker': 'employee_trackers',
        'handbook': 'employee_handbooks',
        'onboarding': 'employee_onboarding'
    }
    subfolder = subfolder_map.get(document_type, 'employee_documents')
    
    file_url = await save_upload_file(file_content, filename, subfolder=subfolder)
    
    # 获取文件类型
    file_type = file_ext[1:].lower() if file_ext else 'unknown'
    
    # 创建文档记录
    document = EmployeeDocument(
        employee_id=employee_id,
        name=file.filename,
        file_type=file_type,
        file_url=file_url,
        document_type=document_type,
        uploaded_by=str(current_user.id)
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    return document


@router.get("/{employee_id}/documents/{document_id}/preview")
async def preview_employee_document(
    employee_id: str,
    document_id: str,
    format: Optional[str] = Query(default=None),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """预览员工文档（用于iframe显示）"""
    document = db.query(EmployeeDocument).filter(
        EmployeeDocument.id == document_id,
        EmployeeDocument.employee_id == employee_id
    ).first()
    if not document:
        raise HTTPException(status_code=404, detail="文档不存在")
    
    # 获取文件路径
    file_path = get_file_path(document.file_url)
    if not file_path:
        from core.config import settings
        upload_path = Path(settings.upload_dir)
        url_path = document.file_url
        if url_path.startswith("uploads/"):
            url_path = url_path.replace("uploads/", "", 1)
        elif url_path.startswith("uploads\\"):
            url_path = url_path.replace("uploads\\", "", 1)
        file_path = upload_path / url_path
        if not file_path.exists():
            file_path = Path(document.file_url)
    
    if not file_path or not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在")
    
    preview_format = (format or "").lower()
    filename = document.name
    should_force_pdf = (
        preview_format == "pdf"
        or (
            not preview_format
            and document.document_type == "contract"
            and file_path.suffix.lower() in [".doc", ".docx"]
        )
    )
    if should_force_pdf and file_path.suffix.lower() in [".doc", ".docx"]:
        try:
            pdf_path = _convert_docx_to_pdf_in_place(file_path)
            if not pdf_path or not pdf_path.exists():
                raise RuntimeError("未找到转换后的PDF文件")
            file_path = pdf_path
            document.file_url = str(pdf_path)
            document.file_type = "pdf"
            db.commit()
        except Exception as e:
            logger.error(f"docx转PDF失败: {e}", exc_info=True)
            raise HTTPException(status_code=500, detail=f"docx转PDF失败: {e}")
        mime_type = "application/pdf"
        filename = f"{Path(document.name).stem}.pdf"
    else:
        # 获取MIME类型
        mime_type, _ = mimetypes.guess_type(str(file_path))
        if not mime_type:
            mime_type = "application/octet-stream"
    
    ascii_filename = to_ascii_filename(filename)
    return FileResponse(
        path=str(file_path),
        media_type=mime_type,
        filename=ascii_filename,
        headers={
            "Content-Disposition": build_content_disposition(filename, "inline"),
            "X-Frame-Options": "ALLOWALL",
            "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
            "Pragma": "no-cache",
        }
    )


@router.get("/{employee_id}/documents/{document_id}/download")
async def download_employee_document(
    employee_id: str,
    document_id: str,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """下载员工文档"""
    document = db.query(EmployeeDocument).filter(
        EmployeeDocument.id == document_id,
        EmployeeDocument.employee_id == employee_id
    ).first()
    if not document:
        raise HTTPException(status_code=404, detail="文档不存在")
    
    # 获取文件路径
    file_path = get_file_path(document.file_url)
    if not file_path:
        from core.config import settings
        upload_path = Path(settings.upload_dir)
        file_path = upload_path / document.file_url.replace("uploads/", "").replace("uploads\\", "")
        if not file_path.exists():
            file_path = Path(document.file_url)
    
    if not file_path or not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在")
    
    # 获取MIME类型
    mime_type, _ = mimetypes.guess_type(str(file_path))
    if not mime_type:
        mime_type = "application/octet-stream"
    
    ascii_filename = to_ascii_filename(document.name)
    return FileResponse(
        path=str(file_path),
        media_type=mime_type,
        filename=ascii_filename,
        headers={
            "Content-Disposition": build_content_disposition(document.name, "attachment")
        }
    )


@router.delete("/{employee_id}/documents/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_employee_document(
    employee_id: str,
    document_id: str,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """删除员工文档"""
    document = db.query(EmployeeDocument).filter(
        EmployeeDocument.id == document_id,
        EmployeeDocument.employee_id == employee_id
    ).first()
    if not document:
        raise HTTPException(status_code=404, detail="文档不存在")
    
    # 删除文件
    try:
        file_path = get_file_path(document.file_url)
        if file_path and file_path.exists():
            file_path.unlink()
    except Exception as e:
        # 记录错误但不阻止删除数据库记录
        print(f"删除文件失败: {e}")
    
    db.delete(document)
    db.commit()
    return None


@router.post("/documents/bulk-upload", response_model=List[EmployeeDocumentResponse])
async def bulk_upload_employee_documents(
    document_type: str = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """批量上传文档给所有员工"""
    # 验证文档类型
    valid_types = ['checklist', 'code', 'tracker', 'handbook', 'onboarding']
    if document_type not in valid_types:
        raise HTTPException(status_code=400, detail=f"无效的文档类型，必须是: {', '.join(valid_types)}")
    
    # 获取所有员工
    employees = db.query(Employee).all()
    if not employees:
        raise HTTPException(status_code=404, detail="没有找到任何员工")
    
    # 上传文件
    file_ext = os.path.splitext(file.filename)[1]
    filename = f"{uuid.uuid4()}{file_ext}"
    file_content = await file.read()
    
    # 根据文档类型选择子文件夹
    subfolder_map = {
        'checklist': 'employee_checklists',
        'code': 'employee_codes',
        'tracker': 'employee_trackers',
        'handbook': 'employee_handbooks',
        'onboarding': 'employee_onboarding'
    }
    subfolder = subfolder_map.get(document_type, 'employee_documents')
    
    file_url = await save_upload_file(file_content, filename, subfolder=subfolder)
    
    # 获取文件类型
    file_type = file_ext[1:].lower() if file_ext else 'unknown'
    
    # 为所有员工创建文档记录
    documents = []
    for employee in employees:
        document = EmployeeDocument(
            employee_id=employee.id,
            name=file.filename,
            file_type=file_type,
            file_url=file_url,
            document_type=document_type,
            uploaded_by=str(current_user.id)
        )
        db.add(document)
        documents.append(document)
    
    db.commit()
    
    # 刷新所有文档记录
    for document in documents:
        db.refresh(document)
    
    return documents


@router.post("/{employee_id}/contracts/generate", response_model=EmployeeDocumentResponse)
async def generate_employee_contract(
    employee_id: str,
    contract_data: ContractGenerateRequest,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """生成员工合同"""
    try:
        from docx import Document
        from docx.shared import RGBColor
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx 库未安装，无法生成合同")
    
    employee = db.query(Employee).filter(Employee.id == employee_id).first()
    if not employee:
        raise HTTPException(status_code=404, detail="员工不存在")
    
    # 获取模板文件路径
    template_path = Path(__file__).parent.parent.parent.parent / "static" / "Staff Contract Sample.docx"
    if not template_path.exists():
        raise HTTPException(status_code=404, detail="合同模板文件不存在")
    
    try:
        # 读取模板
        doc = Document(str(template_path))
    except Exception as e:
        logger.error(f"读取合同模板失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"读取合同模板失败: {str(e)}")

    # 格式化日期
    def format_date(date_str):
        if not date_str:
            return ''
        try:
            from datetime import datetime
            date_obj = datetime.strptime(date_str, '%Y-%m-%d')
            return date_obj.strftime('%d/%m/%Y')
        except Exception:
            return date_str

    # 格式化薪资（不带 $，避免模板内已有 $ 时出现双 $）
    def format_salary(salary):
        if salary is None:
            return ''
        return f"{float(salary):,.2f}"

    # 准备替换数据
    start_date = format_date(contract_data.start_date)
    employment_type = contract_data.employment_type
    position = contract_data.position
    hours_per_week = str(contract_data.hours_per_week) if contract_data.hours_per_week else ''
    work_hours = contract_data.work_hours or ''
    gross_salary = format_salary(contract_data.gross_salary)
    superior_first_name = contract_data.superior_first_name or ''
    superior_last_name = contract_data.superior_last_name or ''
    superior_title = contract_data.superior_title or ''

    # 构建替换映射（只替换模板中的红字占位符）
    replacements = []

    # Start Date
    if start_date:
        replacements.append(('Start Date', start_date))

    # permanent full-time（按表单选择的雇佣类型文本替换）
    employment_type_text = ''
    if employment_type == 'full-time':
        employment_type_text = 'Permanent Full-time'
    elif employment_type == 'part-time':
        employment_type_text = 'Part time'
    elif employment_type == 'casual':
        employment_type_text = 'Casual'
    if employment_type_text:
        replacements.append(('permanent full-time', employment_type_text))

    # Position（按表单选择的职位文本替换）
    position_text = ''
    if position == 'support-worker':
        position_text = 'Support Worker'
    elif position == 'admin':
        position_text = 'Admin'
    elif position == 'office-staff':
        position_text = 'Office Staff'
    if position_text:
        replacements.append(('Position', position_text))

    # Superior first name Superior last name, Superior title（包含 fi 连字）
    superior_combined = ''
    if superior_first_name or superior_last_name or superior_title:
        full_name = f"{superior_first_name} {superior_last_name}".strip()
        superior_combined = f"{full_name}, {superior_title}".strip(', ')
    if superior_combined:
        replacements.append(('Superior first name Superior last name, Superior title', superior_combined))
        replacements.append(('Superior ﬁrst name Superior last name, Superior title', superior_combined))

    # 37.5 hours per week (for a full-time employee)
    if hours_per_week:
        replacements.append(('37.5 hours per week (for a full-time employee)', f'{hours_per_week} hours per week (for a full-time employee)'))

    # 9 am to 5 pm
    if work_hours:
        replacements.append(('9 am to 5 pm', work_hours))

    # Insert gross salary
    if gross_salary:
        replacements.append(('Insert gross salary', gross_salary))
        replacements.append(('$Insert gross salary', f'${gross_salary}'))

    def replace_text_in_paragraph(paragraph):
        import re
        runs = list(paragraph.runs)
        punctuation_pattern = re.compile(r'^[\s,.:;]+$')

        def replace_match_in_runs(match_start, match_end, new_text):
            pos = 0
            start_idx = end_idx = None
            start_offset = end_offset = 0
            for idx, run in enumerate(runs):
                run_text = run.text
                run_len = len(run_text)
                if start_idx is None and pos + run_len > match_start:
                    start_idx = idx
                    start_offset = match_start - pos
                if pos + run_len >= match_end:
                    end_idx = idx
                    end_offset = match_end - pos
                    break
                pos += run_len

            if start_idx is None or end_idx is None:
                return

            def set_black(run):
                try:
                    # 尝试直接设置颜色
                    run.font.color.rgb = RGBColor(0, 0, 0)
                except (AttributeError, TypeError):
                    # 如果color为None或无法设置，跳过颜色设置
                    # 这不会影响文本替换功能
                    pass

            if start_idx == end_idx:
                run = runs[start_idx]
                run.text = run.text[:start_offset] + new_text + run.text[end_offset:]
                set_black(run)
                # 邻近标点可能是红色，统一改黑
                if start_idx + 1 < len(runs) and punctuation_pattern.match(runs[start_idx + 1].text):
                    set_black(runs[start_idx + 1])
                if start_idx - 1 >= 0 and punctuation_pattern.match(runs[start_idx - 1].text):
                    set_black(runs[start_idx - 1])
                return

            first = runs[start_idx]
            last = runs[end_idx]
            prefix = first.text[:start_offset]
            suffix = last.text[end_offset:]
            first.text = prefix + new_text
            set_black(first)
            for idx in range(start_idx + 1, end_idx):
                runs[idx].text = ''
                set_black(runs[idx])
            last.text = suffix
            set_black(last)
            # 邻近标点可能是红色，统一改黑
            if end_idx + 1 < len(runs) and punctuation_pattern.match(runs[end_idx + 1].text):
                set_black(runs[end_idx + 1])
            if start_idx - 1 >= 0 and punctuation_pattern.match(runs[start_idx - 1].text):
                set_black(runs[start_idx - 1])

        for old_text, new_text in replacements:
            if not old_text:
                continue
            try:
                pattern = re.compile(re.escape(old_text), re.IGNORECASE)
                max_iterations = 100  # 防止无限循环
                iteration = 0
                while iteration < max_iterations:
                    combined_text = ''.join(run.text for run in runs)
                    match = pattern.search(combined_text)
                    if not match:
                        break
                    replace_match_in_runs(match.start(), match.end(), new_text)
                    iteration += 1
                if iteration >= max_iterations:
                    logger.warning(f"替换文本 '{old_text}' 达到最大迭代次数，可能存在问题")
            except Exception as e:
                logger.error(f"替换文本 '{old_text}' 时出错: {e}", exc_info=True)
                # 继续处理其他替换，不中断整个流程

    try:
        # 处理段落
        for paragraph in doc.paragraphs:
            try:
                replace_text_in_paragraph(paragraph)
            except Exception as e:
                logger.error(f"处理段落时出错: {e}", exc_info=True)
                # 继续处理其他段落

        # 处理表格
        for table in doc.tables:
            try:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            try:
                                replace_text_in_paragraph(paragraph)
                            except Exception as e:
                                logger.error(f"处理表格单元格段落时出错: {e}", exc_info=True)
                                # 继续处理其他单元格
            except Exception as e:
                logger.error(f"处理表格时出错: {e}", exc_info=True)
                # 继续处理其他表格
    except Exception as e:
        logger.error(f"处理Word文档时出错: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"处理Word文档失败: {str(e)}")
    
    # 保存生成的合同
    output_filename = f"contract_{employee_id}_{uuid.uuid4()}.docx"
    subfolder = 'employee_contracts'
    tmp_file_path = None
    
    try:
        # 创建临时文件保存生成的合同
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            doc.save(tmp_file.name)
            tmp_file_path = tmp_file.name
        
        # 读取临时文件内容
        with open(tmp_file_path, 'rb') as f:
            file_content = f.read()
        
        # 删除临时文件
        try:
            os.unlink(tmp_file_path)
            tmp_file_path = None
        except Exception as e:
            logger.warning(f"删除临时文件失败: {e}")
        
        # 上传文件
        file_url = await save_upload_file(file_content, output_filename, subfolder=subfolder)
        
        # 创建文档记录
        document = EmployeeDocument(
            employee_id=employee_id,
            name=output_filename,
            file_type='docx',
            file_url=file_url,
            document_type='contract',
            uploaded_by=str(current_user.id)
        )
        db.add(document)
        db.commit()
        db.refresh(document)
        
        return document
    except Exception as e:
        # 如果出错，回滚数据库事务
        db.rollback()
        logger.error(f"保存合同文件失败: {e}", exc_info=True)
        
        # 清理临时文件
        if tmp_file_path and os.path.exists(tmp_file_path):
            try:
                os.unlink(tmp_file_path)
            except Exception:
                pass
        
        raise HTTPException(status_code=500, detail=f"保存合同文件失败: {str(e)}")



# ==================== 合同签署链接（无需登录） ====================

@router.post("/{employee_id}/contracts/{contract_id}/create-sign-link")
async def create_employee_contract_sign_link(
    employee_id: str,
    contract_id: str,
    request: Request,
    language: str = Query("en", description="邮件语言：zh / en"),
    payload: dict | None = Body(default=None),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user),
):
    """生成合同签署链接并发送到员工邮箱（默认 7 天有效）"""
    employee = db.query(Employee).filter(Employee.id == employee_id).first()
    if not employee:
        raise HTTPException(status_code=404, detail="员工不存在")
    if not employee.email:
        raise HTTPException(status_code=400, detail="员工邮箱不存在，无法发送签署链接")

    document = db.query(EmployeeDocument).filter(
        EmployeeDocument.id == contract_id,
        EmployeeDocument.employee_id == employee_id,
        EmployeeDocument.document_type == "contract",
    ).first()
    if not document:
        raise HTTPException(status_code=404, detail="合同不存在")
    if document.employee_signed_at:
        raise HTTPException(status_code=400, detail="该合同已由员工签字，无需再次发送")
    if not document.file_url:
        raise HTTPException(status_code=400, detail="合同文件未就绪")

    db.query(EmployeeContractSignRequest).filter(
        EmployeeContractSignRequest.employee_id == employee_id,
        EmployeeContractSignRequest.contract_id == contract_id,
        EmployeeContractSignRequest.status == "pending",
    ).update({"status": "expired"})

    token = str(uuid.uuid4())
    expires_at = datetime.utcnow() + timedelta(days=7)
    req = EmployeeContractSignRequest(
        token=token,
        employee_id=employee_id,
        contract_id=contract_id,
        status="pending",
        expires_at=expires_at,
    )
    db.add(req)
    db.commit()

    try:
        touch_business_unread(
            db,
            business_code="employee_document",
            receiver_user_id=str(employee_id),
            data_id=str(contract_id),
            scope_id=str(employee_id),
            trigger_user_id=str(current_user.id),
        )
        db.commit()
    except Exception:
        db.rollback()

    if settings.sign_frontend_base_url:
        raw_base_url = settings.sign_frontend_base_url.strip()
        base_url = raw_base_url.split("#", 1)[0].rstrip("/")
        sign_path = f"/sign/document/{token}" if base_url.lower().endswith("/admin") else f"/admin/sign/document/{token}"
        sign_url = f"{base_url}{sign_path}"
    else:
        raw_origin = (request.headers.get("origin") or str(request.base_url)).strip()
        origin = raw_origin.split("#", 1)[0].rstrip("/")
        sign_path = f"/admin/sign/document/{token}"
        sign_url = f"{origin}{sign_path}"

    def _normalize_lang(value: str | None, accept: str | None) -> str:
        raw = (value or "").strip().lower()
        if not raw and accept:
            raw = str(accept).split(",", 1)[0].strip().lower()
        if raw in ("zh", "zh-cn", "zh_hans", "cn", "chinese", "中文", "简体中文"):
            return "zh"
        if raw in ("en", "en-us", "en-gb", "english", "英文"):
            return "en"
        if raw.startswith("zh"):
            return "zh"
        if raw.startswith("en"):
            return "en"
        return "en"

    body_lang = None
    if isinstance(payload, dict):
        body_lang = payload.get("language") or payload.get("lang")
    lang = _normalize_lang(language or body_lang, request.headers.get("accept-language"))

    sep = "&" if "?" in sign_url else "?"
    sign_url_with_lang = f"{sign_url}{sep}lang={lang}"

    if lang == "en":
        subject = "Employee Contract Signing Link"
        plain_body = (
            f"Dear {employee.name},\n\n"
            f"Please click the link below to sign your contract (valid until: {expires_at.isoformat()}):\n"
            f"{sign_url_with_lang}\n\n"
            "If the link has expired, please contact the administrator to resend it.\n"
        )
        html_body = f"""
<p>Dear {employee.name},</p>
<p>Please click the link below to sign your contract (valid until: {expires_at.isoformat()}):</p>
<p><a href="{sign_url_with_lang}" target="_blank" rel="noopener noreferrer">{sign_url_with_lang}</a></p>
<p>If the link has expired, please contact the administrator to resend it.</p>
""".strip()
    else:
        subject = "员工合同签署链接"
        plain_body = (
            f"{employee.name}，您好：\n\n"
            f"请点击以下链接完成合同签署（有效期至：{expires_at.isoformat()}）：\n"
            f"{sign_url_with_lang}\n\n"
            "如链接已过期，请联系管理员重新发送。"
        )
        html_body = f"""
<p>{employee.name}，您好：</p>
<p>请点击以下链接完成合同签署（有效期至：{expires_at.isoformat()}）：</p>
<p><a href="{sign_url_with_lang}" target="_blank" rel="noopener noreferrer">{sign_url_with_lang}</a></p>
<p>如链接已过期，请联系管理员重新发送。</p>
""".strip()

    try:
        send_contact_email(subject=subject, html_body=html_body, plain_body=plain_body, to_emails=[employee.email])
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"发送签署链接邮件失败: {str(e)}")

    return {
        "token": token,
        "sign_path": sign_path,
        "sign_url": sign_url_with_lang,
        "expires_at": expires_at.isoformat(),
    }


# ==================== 合同签字 API ====================

@router.post("/{employee_id}/contracts/{contract_id}/sign")
async def sign_employee_contract(
    employee_id: str,
    contract_id: str,
    signature_data: str = Body(..., embed=True),
    lang: str | None = Query(default=None),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """管理员提交合同签名"""
    try:
        resolved_lang = "en" if str(lang or "").strip().lower().startswith("en") else "zh"
        # 查找合同文档
        document = db.query(EmployeeDocument).filter(
            EmployeeDocument.id == contract_id,
            EmployeeDocument.employee_id == employee_id,
            EmployeeDocument.document_type == "contract"
        ).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Contract not found" if resolved_lang == "en" else "合同不存在")

        if document.admin_signed_at or document.admin_signature_blob:
            raise HTTPException(status_code=400, detail="Already signed. Please do not sign again." if resolved_lang == "en" else "您已签字请勿重复签名")
        
        # 解析签名数据
        mime_type, blob = _parse_data_url(signature_data)
        if not blob:
            raise HTTPException(status_code=400, detail="Invalid signature data" if resolved_lang == "en" else "无效的签名数据")
        
        # 保存签名到数据库
        document.admin_signature_blob = blob
        document.admin_signature_mime = mime_type or "image/png"
        document.admin_signed_at = datetime.utcnow()
        document.admin_signed_by = str(current_user.id)
        
        # 将签名嵌入到文档中（优先使用管理员坐标，若无则使用员工坐标）
        if document.file_url:
            contract_path = get_file_path(document.file_url)
            if contract_path and contract_path.exists():
                try:
                    x = document.admin_signature_x or document.employee_signature_x
                    y = document.admin_signature_y or document.employee_signature_y
                    width = document.admin_signature_width or document.employee_signature_width
                    height = document.admin_signature_height or document.employee_signature_height
                    page = document.admin_signature_page
                    if page is None:
                        page = document.employee_signature_page or 0

                    # doc/docx 多页签字：必须先转为PDF再嵌入，避免总是落在第一页
                    if contract_path.suffix.lower() in [".doc", ".docx"]:
                        logger.info(f"检测到docx文件，开始转换为PDF: {contract_path}")
                        try:
                            pdf_path = _convert_docx_to_pdf_in_place(contract_path)
                            if not pdf_path or not pdf_path.exists():
                                raise RuntimeError(f"PDF转换失败：转换后的文件不存在: {pdf_path}")
                            logger.info(f"docx转换为PDF成功: {pdf_path}")
                            contract_path = pdf_path
                            # 更新数据库中的文件URL和类型
                            document.file_url = str(pdf_path)
                            document.file_type = "pdf"
                            logger.info(f"已更新文档file_url为PDF路径: {document.file_url}, file_type: {document.file_type}")
                        except Exception as convert_error:
                            logger.error(f"docx转PDF失败: {convert_error}", exc_info=True)
                            raise HTTPException(status_code=500, detail=f"docx转PDF失败: {str(convert_error)}")
                    
                    # 验证转换后的文件必须是PDF
                    if contract_path.suffix.lower() != ".pdf":
                        error_msg = f"文件类型错误：转换后应该是PDF，但实际是 {contract_path.suffix}。文件路径: {contract_path}"
                        logger.error(error_msg)
                        raise RuntimeError(error_msg)
                    
                    # 嵌入签名到PDF
                    logger.info(f"开始嵌入签名到PDF: {contract_path}, 坐标: x={x}, y={y}, width={width}, height={height}, page={page}")
                    embed_success = _embed_signature_to_pdf(
                        contract_path,
                        blob,
                        "admin",
                        x=x,
                        y=y,
                        width=width,
                        height=height,
                        page_index=page,
                        date_text=datetime.now().date().isoformat(),
                    )
                    if embed_success is False:
                        raise RuntimeError("嵌入签名失败")
                except Exception as e:
                    logger.error(f"嵌入签名到合同失败: {e}")
                    raise HTTPException(status_code=500, detail=f"嵌入签名失败: {e}")
        
        db.commit()
        db.refresh(document)
        
        logger.info(f"Contract signed by admin: {document.id}, employee: {employee_id}, admin: {current_user.id}")
        
        return {
            "id": document.id,
            "admin_signed_at": document.admin_signed_at.isoformat() if document.admin_signed_at else None,
            "admin_signed_by": document.admin_signed_by,
            "lang": resolved_lang,
            "message": "Contract signing submitted successfully" if resolved_lang == "en" else "合同签名提交成功"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to sign contract: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to submit contract signing: {str(e)}" if (str(lang or '').strip().lower().startswith('en')) else f"提交合同签名失败: {str(e)}")


@router.post("/{employee_id}/contracts/{contract_id}/signature-position")
async def save_admin_signature_position(
    employee_id: str,
    contract_id: str,
    x: float = Body(...),
    y: float = Body(...),
    width: float = Body(...),
    height: float = Body(...),
    page: int = Body(0),
    lang: str | None = Query(default=None),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """管理员保存合同签字坐标"""
    try:
        resolved_lang = "en" if str(lang or "").strip().lower().startswith("en") else "zh"
        document = db.query(EmployeeDocument).filter(
            EmployeeDocument.id == contract_id,
            EmployeeDocument.employee_id == employee_id,
            EmployeeDocument.document_type == "contract"
        ).first()

        if not document:
            raise HTTPException(status_code=404, detail="Contract not found" if resolved_lang == "en" else "合同不存在")

        if document.admin_signed_at or document.admin_signature_blob:
            raise HTTPException(status_code=400, detail="Already signed. Please do not sign again." if resolved_lang == "en" else "您已签字请勿重复签名")

        document.admin_signature_x = x
        document.admin_signature_y = y
        document.admin_signature_width = width
        document.admin_signature_height = height
        document.admin_signature_page = page

        db.commit()
        db.refresh(document)

        return {
            "id": document.id,
            "admin_signature_x": document.admin_signature_x,
            "admin_signature_y": document.admin_signature_y,
            "admin_signature_width": document.admin_signature_width,
            "admin_signature_height": document.admin_signature_height,
            "admin_signature_page": document.admin_signature_page,
            "lang": resolved_lang,
            "message": "Admin signature position saved successfully" if resolved_lang == "en" else "管理员签字坐标保存成功"
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"保存管理员签字坐标失败: {e}")
        raise HTTPException(status_code=500, detail=f"保存签字坐标失败: {str(e)}")


@router.get("/{employee_id}/contracts/{contract_id}/employee-signature/image")
async def get_contract_employee_signature_image(
    employee_id: str,
    contract_id: str,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """获取员工合同签名图片"""
    document = db.query(EmployeeDocument).filter(
        EmployeeDocument.id == contract_id,
        EmployeeDocument.employee_id == employee_id,
        EmployeeDocument.document_type == "contract"
    ).first()
    
    if not document or not document.employee_signature_blob:
        raise HTTPException(status_code=404, detail="签名不存在")
    
    return Response(
        content=document.employee_signature_blob,
        media_type=document.employee_signature_mime or "image/png"
    )


@router.get("/{employee_id}/contracts/{contract_id}/admin-signature/image")
async def get_contract_admin_signature_image(
    employee_id: str,
    contract_id: str,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """获取管理员合同签名图片"""
    document = db.query(EmployeeDocument).filter(
        EmployeeDocument.id == contract_id,
        EmployeeDocument.employee_id == employee_id,
        EmployeeDocument.document_type == "contract"
    ).first()
    
    if not document or not document.admin_signature_blob:
        raise HTTPException(status_code=404, detail="签名不存在")
    
    return Response(
        content=document.admin_signature_blob,
        media_type=document.admin_signature_mime or "image/png"
    )
